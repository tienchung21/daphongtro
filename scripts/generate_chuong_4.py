#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Generator for Chuong 4: TRIỂN KHAI HỆ THỐNG
- Backend Implementation (Controllers, Models, Services)
- Frontend Implementation (Components, Pages)
- Real-time với Socket.IO
- Payment Integration
- Security Implementation
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def setup_document_styles(doc):
    """Setup document styles"""
    styles = doc.styles
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    
    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

def add_paragraph(doc, text, bold=False, italic=False):
    """Add paragraph with formatting"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    return p

def add_bullet_list(doc, items):
    """Add bullet list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

def add_numbered_list(doc, items):
    """Add numbered list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)

def add_table_with_data(doc, headers, rows):
    """Add table with data"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def add_code_block(doc, code_text, language='javascript'):
    """Add code block"""
    p = doc.add_paragraph(code_text)
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def generate_section_4_1(doc):
    """4.1. Môi trường triển khai"""
    doc.add_heading('4.1. Môi trường triển khai', level=2)
    
    # 4.1.1. Technology Stack
    doc.add_heading('4.1.1. Technology Stack', level=3)
    
    tech_stack = [
        ['Backend', 'Node.js 18+, Express 4.18', 'API Server'],
        ['Database', 'MySQL 8.0', 'RDBMS'],
        ['Frontend', 'React 18+, Vite 5.0', 'SPA Framework'],
        ['Real-time', 'Socket.IO 4.8.1 (18 implementations)', 'WebSocket'],
        ['Payment', 'SePay API', 'Payment Gateway'],
        ['Authentication', 'JWT, bcrypt', 'Security'],
        ['File Upload', 'Multer', 'File handling'],
        ['HTTP Client', 'Axios', 'API calls']
    ]
    
    add_table_with_data(doc, ['Category', 'Technology', 'Purpose'], tech_stack)
    
    doc.add_paragraph()
    
    # 4.1.2. Project Structure
    doc.add_heading('4.1.2. Cấu trúc dự án', level=3)
    
    add_paragraph(doc, """Dự án được tổ chức theo mô hình Monorepo với 2 packages chính:""")
    
    project_structure = """
daphongtro/
├── client/                 # Frontend React App
│   ├── src/
│   │   ├── api/           # API client services (7 files)
│   │   ├── components/    # Reusable components (51 .jsx + 42 .css)
│   │   ├── pages/         # Page components (52 .jsx + 41 .css)
│   │   ├── layouts/       # Layout components
│   │   ├── services/      # Business logic services
│   │   ├── utils/         # Utility functions
│   │   └── context/       # React Context (ChatContext)
│   └── vite.config.js
│
├── server/                # Backend Node.js App
│   ├── controllers/       # Request handlers (26 files)
│   ├── models/            # Database models (20 files)
│   ├── routes/            # API routes (24 files)
│   ├── services/          # Business services (9 files)
│   ├── middleware/        # Auth, Role middleware
│   ├── socket/            # Socket.IO handlers
│   └── config/            # Configuration
│
├── docs/                  # Documentation (99 .md files)
├── migrations/            # Database migrations
└── scripts/               # Utility scripts
"""
    
    add_code_block(doc, project_structure)
    
    doc.add_paragraph()
    
    # 4.1.3. Development Setup
    doc.add_heading('4.1.3. Cài đặt môi trường phát triển', level=3)
    
    add_paragraph(doc, bold=True, text="Backend setup:")
    
    backend_setup = """
cd server
npm install
# Setup environment variables
cp .env.example .env
# Configure DB connection, JWT secret, SePay credentials
npm run dev  # Start development server on port 3001
"""
    
    add_code_block(doc, backend_setup, 'bash')
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Frontend setup:")
    
    frontend_setup = """
cd client
npm install
# Setup environment variables
cp .env.example .env
# Configure API base URL
npm run dev  # Start dev server on port 5173
"""
    
    add_code_block(doc, frontend_setup, 'bash')

def generate_section_4_2(doc):
    """4.2. Triển khai Backend"""
    doc.add_heading('4.2. Triển khai Backend', level=2)
    
    add_paragraph(doc, """Backend được triển khai theo kiến trúc MVC (Model-View-Controller) 
với 3 layers: Routes → Controllers → Models/Services.""")
    
    # 4.2.1. Database Layer
    doc.add_heading('4.2.1. Database Connection', level=3)
    
    add_paragraph(doc, """Database connection được quản lý bởi mysql2/promise với connection pooling:""")
    
    db_config = """
// server/config/db.js
const mysql = require('mysql2/promise');

const pool = mysql.createPool({
  host: process.env.DB_HOST || 'localhost',
  user: process.env.DB_USER || 'root',
  password: process.env.DB_PASSWORD || '',
  database: process.env.DB_NAME || 'daphongtro',
  waitForConnections: true,
  connectionLimit: 10,
  queueLimit: 0,
  timezone: '+07:00'
});

module.exports = pool;
"""
    
    add_code_block(doc, db_config)
    
    doc.add_paragraph()
    
    # 4.2.2. Authentication Implementation
    doc.add_heading('4.2.2. Authentication & Authorization', level=3)
    
    add_paragraph(doc, bold=True, text="JWT Authentication Middleware:")
    
    auth_middleware = """
// server/middleware/auth.js
const jwt = require('jsonwebtoken');

const authenticate = async (req, res, next) => {
  try {
    const token = req.headers.authorization?.split(' ')[1];
    
    if (!token) {
      return res.status(401).json({ 
        message: 'No token provided' 
      });
    }
    
    const decoded = jwt.verify(token, process.env.JWT_SECRET);
    req.user = decoded;
    next();
  } catch (error) {
    return res.status(401).json({ 
      message: 'Invalid token' 
    });
  }
};

const authorize = (roles = []) => {
  return (req, res, next) => {
    if (!roles.includes(req.user.LoaiNguoiDung)) {
      return res.status(403).json({ 
        message: 'Forbidden' 
      });
    }
    next();
  };
};

module.exports = { authenticate, authorize };
"""
    
    add_code_block(doc, auth_middleware)
    
    doc.add_paragraph()
    
    # 4.2.3. Core Controllers
    doc.add_heading('4.2.3. Triển khai Controllers', level=3)
    
    add_paragraph(doc, """Hệ thống có 26 controllers xử lý các domain khác nhau. 
Dưới đây là implementation của các controllers quan trọng nhất:""")
    
    controllers_list = [
        ['ChuDuAnController.js', '1648 dòng', 'Quản lý Dự án, Tin đăng, Cuộc hẹn, Cọc'],
        ['tinDangController.js', '450 dòng', 'CRUD Tin đăng, Upload ảnh'],
        ['cuocHenController.js', '350 dòng', 'Quản lý Cuộc hẹn'],
        ['cocController.js', '400 dòng', 'Quản lý Cọc, Tích hợp SePay'],
        ['hopDongController.js', '500 dòng', 'Quản lý Hợp đồng, E-signature'],
        ['phongController.js', '600 dòng', 'Quản lý Phòng đang thuê'],
        ['chatController.js', '300 dòng', 'Chat messaging'],
        ['adminController.js', '550 dòng', 'Admin operations'],
        ['authController.js', '250 dòng', 'Login, Register, KYC']
    ]
    
    add_table_with_data(doc, ['Controller', 'Size', 'Chức năng'], controllers_list)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Ví dụ: TinDangController - Tạo tin đăng mới")
    
    tindang_controller = """
// server/controllers/tinDangController.js
const TinDangModel = require('../models/tinDangModel');
const { validationResult } = require('express-validator');

exports.taoTinDang = async (req, res) => {
  try {
    // Validation
    const errors = validationResult(req);
    if (!errors.isEmpty()) {
      return res.status(400).json({ errors: errors.array() });
    }
    
    const { 
      TieuDe, MoTa, DienTich, GiaThue, 
      DiaChi, Tinh, Huyen, Xa, 
      DuAnID, ThuocTinh, TienIch 
    } = req.body;
    
    const ChuDuAnID = req.user.NguoiDungID;
    
    // Create TinDang
    const tinDangId = await TinDangModel.taoTinDang({
      ChuDuAnID,
      DuAnID,
      TieuDe,
      MoTa,
      DienTich,
      GiaThue,
      DiaChi,
      Tinh,
      Huyen,
      Xa,
      TrangThai: 'ChoDuyet'
    });
    
    // Add ThuocTinh & TienIch
    if (ThuocTinh?.length) {
      await TinDangModel.themThuocTinh(tinDangId, ThuocTinh);
    }
    if (TienIch?.length) {
      await TinDangModel.themTienIch(tinDangId, TienIch);
    }
    
    // Audit log
    await NhatKyHeThongService.ghiNhatKy({
      NguoiDungID: ChuDuAnID,
      HanhDong: 'TAO_TIN_DANG',
      DoiTuong: 'TinDang',
      DoiTuongID: tinDangId
    });
    
    res.status(201).json({
      success: true,
      data: { TinDangID: tinDangId },
      message: 'Tạo tin đăng thành công'
    });
    
  } catch (error) {
    console.error('[TinDangController] Error:', error);
    res.status(500).json({ 
      message: 'Lỗi server', 
      error: error.message 
    });
  }
};
"""
    
    add_code_block(doc, tindang_controller)
    
    doc.add_paragraph()
    
    # 4.2.4. Core Models
    doc.add_heading('4.2.4. Triển khai Models', level=3)
    
    add_paragraph(doc, """Models chịu trách nhiệm tương tác với database. 
Hệ thống có 20 models tương ứng với các entities chính:""")
    
    models_list = [
        ['ChuDuAnModel.js', '1648 dòng', 'Tất cả operations của Chủ dự án'],
        ['tinDangModel.js', '450 dòng', 'CRUD Tin đăng'],
        ['DuAnModel.js', '350 dòng', 'CRUD Dự án'],
        ['CuocHenModel.js', '300 dòng', 'CRUD Cuộc hẹn'],
        ['ChinhSachCocModel.js', '200 dòng', 'Quản lý Chính sách cọc'],
        ['HopDongModel.js', '500 dòng', 'CRUD Hợp đồng'],
        ['PhongModel.js', '600 dòng', 'Quản lý Phòng'],
        ['ChatModel.js', '250 dòng', 'Chat messages'],
        ['userModel.js', '400 dòng', 'User management']
    ]
    
    add_table_with_data(doc, ['Model', 'Size', 'Chức năng'], models_list)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Ví dụ: TinDangModel - Lấy danh sách tin đăng")
    
    tindang_model = """
// server/models/tinDangModel.js
const db = require('../config/db');

class TinDangModel {
  static async layDanhSachTinDang(filters = {}) {
    try {
      let query = `
        SELECT 
          td.*,
          cd.HoTen as TenChuDuAn,
          cd.SoDienThoai,
          (SELECT HinhAnh FROM hinhanh 
           WHERE TinDangID = td.TinDangID 
           LIMIT 1) as AnhDaiDien,
          COUNT(DISTINCT yth.YeuThichID) as SoLuotYeuThich
        FROM tindang td
        JOIN chuduan cd ON td.ChuDuAnID = cd.ChuDuAnID
        LEFT JOIN yeuthich yth ON td.TinDangID = yth.TinDangID
        WHERE 1=1
      `;
      
      const params = [];
      
      // Apply filters
      if (filters.TrangThai) {
        query += ' AND td.TrangThai = ?';
        params.push(filters.TrangThai);
      }
      
      if (filters.Tinh) {
        query += ' AND td.Tinh = ?';
        params.push(filters.Tinh);
      }
      
      if (filters.GiaMin) {
        query += ' AND td.GiaThue >= ?';
        params.push(filters.GiaMin);
      }
      
      if (filters.GiaMax) {
        query += ' AND td.GiaThue <= ?';
        params.push(filters.GiaMax);
      }
      
      if (filters.keyword) {
        query += ` AND (
          td.TieuDe LIKE ? OR 
          td.MoTa LIKE ? OR 
          td.DiaChi LIKE ?
        )`;
        const keyword = `%${filters.keyword}%`;
        params.push(keyword, keyword, keyword);
      }
      
      query += ' GROUP BY td.TinDangID';
      query += ' ORDER BY td.NgayTao DESC';
      
      // Pagination
      const limit = filters.limit || 20;
      const offset = filters.offset || 0;
      query += ' LIMIT ? OFFSET ?';
      params.push(limit, offset);
      
      const [rows] = await db.execute(query, params);
      return rows;
      
    } catch (error) {
      throw new Error(`Lỗi lấy DS tin đăng: ${error.message}`);
    }
  }
  
  static async taoTinDang(data) {
    try {
      const query = `
        INSERT INTO tindang (
          ChuDuAnID, DuAnID, TieuDe, MoTa,
          DienTich, GiaThue, DiaChi, Tinh, 
          Huyen, Xa, TrangThai
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
      `;
      
      const params = [
        data.ChuDuAnID,
        data.DuAnID || null,
        data.TieuDe,
        data.MoTa,
        data.DienTich,
        data.GiaThue,
        data.DiaChi,
        data.Tinh,
        data.Huyen,
        data.Xa,
        data.TrangThai || 'Nhap'
      ];
      
      const [result] = await db.execute(query, params);
      return result.insertId;
      
    } catch (error) {
      throw new Error(`Lỗi tạo tin đăng: ${error.message}`);
    }
  }
}

module.exports = TinDangModel;
"""
    
    add_code_block(doc, tindang_model)

def generate_section_4_3(doc):
    """4.3. Triển khai Real-time Features"""
    doc.add_heading('4.3. Triển khai Real-time với Socket.IO', level=2)
    
    add_paragraph(doc, """Hệ thống sử dụng Socket.IO để implement các tính năng real-time:
- Chat messaging giữa NguoiThue và ChuDuAn
- Notifications (cuộc hẹn mới, cọc, hợp đồng)
- Online status tracking""")
    
    # 4.3.1. Socket.IO Server Setup
    doc.add_heading('4.3.1. Socket.IO Server Setup', level=3)
    
    socket_server = """
// server/index.js
const express = require('express');
const http = require('http');
const { Server } = require('socket.io');
const socketAuth = require('./middleware/socketAuth');

const app = express();
const server = http.createServer(app);

const io = new Server(server, {
  cors: {
    origin: process.env.CLIENT_URL || 'http://localhost:5173',
    credentials: true
  }
});

// Socket authentication middleware
io.use(socketAuth);

// Socket event handlers
io.on('connection', (socket) => {
  console.log(`User connected: ${socket.user.NguoiDungID}`);
  
  // Join user's personal room
  socket.join(`user:${socket.user.NguoiDungID}`);
  
  // Chat handlers
  require('./socket/chatHandlers')(io, socket);
  
  socket.on('disconnect', () => {
    console.log(`User disconnected: ${socket.user.NguoiDungID}`);
  });
});

server.listen(3001, () => {
  console.log('Server running on port 3001');
});
"""
    
    add_code_block(doc, socket_server)
    
    doc.add_paragraph()
    
    # 4.3.2. Chat Implementation
    doc.add_heading('4.3.2. Chat Implementation', level=3)
    
    add_paragraph(doc, bold=True, text="Server-side Chat Handlers:")
    
    chat_handlers = """
// server/socket/chatHandlers.js
const ChatModel = require('../models/ChatModel');

module.exports = (io, socket) => {
  // Send message
  socket.on('chat:send_message', async (data) => {
    try {
      const { NguoiNhanID, TinDangID, NoiDung, HinhAnh } = data;
      
      // Save to database
      const tinNhanId = await ChatModel.guiTinNhan({
        NguoiGuiID: socket.user.NguoiDungID,
        NguoiNhanID,
        TinDangID,
        NoiDung,
        HinhAnh
      });
      
      // Get full message data
      const message = await ChatModel.layTinNhan(tinNhanId);
      
      // Emit to sender
      socket.emit('chat:message_sent', message);
      
      // Emit to receiver
      io.to(`user:${NguoiNhanID}`).emit('chat:new_message', message);
      
    } catch (error) {
      socket.emit('chat:error', { 
        message: error.message 
      });
    }
  });
  
  // Mark as read
  socket.on('chat:mark_read', async (data) => {
    try {
      const { TinNhanID } = data;
      await ChatModel.danhDauDaDoc(TinNhanID);
      
      // Notify sender
      const message = await ChatModel.layTinNhan(TinNhanID);
      io.to(`user:${message.NguoiGuiID}`)
        .emit('chat:message_read', { TinNhanID });
        
    } catch (error) {
      socket.emit('chat:error', { 
        message: error.message 
      });
    }
  });
  
  // Typing indicator
  socket.on('chat:typing', (data) => {
    const { NguoiNhanID } = data;
    io.to(`user:${NguoiNhanID}`).emit('chat:user_typing', {
      NguoiDungID: socket.user.NguoiDungID,
      HoTen: socket.user.HoTen
    });
  });
};
"""
    
    add_code_block(doc, chat_handlers)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Client-side Socket Integration:")
    
    chat_client = """
// client/src/context/ChatContext.jsx
import { createContext, useContext, useEffect, useState } from 'react';
import io from 'socket.io-client';

const ChatContext = createContext();

export const ChatProvider = ({ children }) => {
  const [socket, setSocket] = useState(null);
  const [messages, setMessages] = useState([]);
  const [unreadCount, setUnreadCount] = useState(0);
  
  useEffect(() => {
    const token = localStorage.getItem('token');
    if (!token) return;
    
    // Connect to socket
    const newSocket = io('http://localhost:3001', {
      auth: { token }
    });
    
    newSocket.on('connect', () => {
      console.log('Socket connected');
    });
    
    // Listen for new messages
    newSocket.on('chat:new_message', (message) => {
      setMessages(prev => [...prev, message]);
      setUnreadCount(prev => prev + 1);
    });
    
    setSocket(newSocket);
    
    return () => newSocket.close();
  }, []);
  
  const sendMessage = (data) => {
    socket?.emit('chat:send_message', data);
  };
  
  const markAsRead = (tinNhanId) => {
    socket?.emit('chat:mark_read', { TinNhanID: tinNhanId });
  };
  
  return (
    <ChatContext.Provider value={{ 
      socket, 
      messages, 
      unreadCount, 
      sendMessage, 
      markAsRead 
    }}>
      {children}
    </ChatContext.Provider>
  );
};

export const useChat = () => useContext(ChatContext);
"""
    
    add_code_block(doc, chat_client)

def generate_section_4_4(doc):
    """4.4. Triển khai Payment Integration"""
    doc.add_heading('4.4. Tích hợp Thanh toán với SePay', level=2)
    
    add_paragraph(doc, """Hệ thống tích hợp SePay Gateway để xử lý thanh toán cọc phòng. 
Luồng thanh toán:
1. NguoiThue tạo cọc → Hệ thống gọi SePay API tạo payment link
2. NguoiThue thanh toán qua SePay
3. SePay callback webhook → Hệ thống verify và cập nhật trạng thái cọc""")
    
    # 4.4.1. Create Payment
    doc.add_heading('4.4.1. Tạo Payment Link', level=3)
    
    create_payment = """
// server/services/SePayService.js
const axios = require('axios');
const crypto = require('crypto');

class SePayService {
  static async taoThanhToan(data) {
    try {
      const { 
        CocID, 
        SoTien, 
        NguoiThueID, 
        TinDangID 
      } = data;
      
      const orderCode = `COC_${CocID}_${Date.now()}`;
      
      // Create payment request
      const paymentData = {
        amount: SoTien,
        order_code: orderCode,
        return_url: `${process.env.CLIENT_URL}/thanh-toan/ket-qua`,
        cancel_url: `${process.env.CLIENT_URL}/thanh-toan/huy`,
        buyer_name: data.HoTen,
        buyer_email: data.Email,
        buyer_phone: data.SoDienThoai
      };
      
      // Generate signature
      const signature = this.generateSignature(paymentData);
      paymentData.signature = signature;
      
      // Call SePay API
      const response = await axios.post(
        `${process.env.SEPAY_API_URL}/create-payment`,
        paymentData,
        {
          headers: {
            'Content-Type': 'application/json',
            'Authorization': `Bearer ${process.env.SEPAY_API_KEY}`
          }
        }
      );
      
      // Save transaction
      await TransactionModel.taoGiaoDich({
        CocID,
        MaGiaoDich: orderCode,
        SoTien,
        TrangThai: 'PENDING',
        PaymentURL: response.data.payment_url
      });
      
      return {
        paymentUrl: response.data.payment_url,
        orderCode
      };
      
    } catch (error) {
      throw new Error(`Lỗi tạo thanh toán: ${error.message}`);
    }
  }
  
  static generateSignature(data) {
    const sortedKeys = Object.keys(data).sort();
    const signString = sortedKeys
      .map(key => `${key}=${data[key]}`)
      .join('&');
    
    return crypto
      .createHmac('sha256', process.env.SEPAY_SECRET)
      .update(signString)
      .digest('hex');
  }
}

module.exports = SePayService;
"""
    
    add_code_block(doc, create_payment)
    
    doc.add_paragraph()
    
    # 4.4.2. Webhook Handler
    doc.add_heading('4.4.2. Payment Webhook Handler', level=3)
    
    webhook_handler = """
// server/controllers/paymentController.js
const SePayService = require('../services/SePayService');
const CocModel = require('../models/ChinhSachCocModel');
const TinDangModel = require('../models/tinDangModel');

exports.sePayWebhook = async (req, res) => {
  try {
    const { 
      order_code, 
      status, 
      amount, 
      signature 
    } = req.body;
    
    // Verify signature
    const isValid = SePayService.verifySignature(req.body);
    if (!isValid) {
      return res.status(400).json({ 
        message: 'Invalid signature' 
      });
    }
    
    // Extract CocID from order_code
    const cocId = order_code.split('_')[1];
    
    if (status === 'SUCCESS') {
      // Update Coc status
      await CocModel.capNhatTrangThaiCoc(cocId, {
        TrangThai: 'DaThanhToan',
        NgayThanhToan: new Date(),
        MaGiaoDich: order_code
      });
      
      // Get Coc info
      const coc = await CocModel.layThongTinCoc(cocId);
      
      // Update TinDang status
      await TinDangModel.capNhatTrangThai(
        coc.TinDangID, 
        'DaCoc'
      );
      
      // Send notification to ChuDuAn
      const io = req.app.get('io');
      io.to(`user:${coc.ChuDuAnID}`).emit('notification', {
        type: 'COC_MOI',
        message: `Có cọc mới cho tin đăng ${coc.TieuDe}`,
        CocID: cocId
      });
      
      // Audit log
      await NhatKyHeThongService.ghiNhatKy({
        NguoiDungID: coc.NguoiThueID,
        HanhDong: 'THANH_TOAN_COC',
        DoiTuong: 'Coc',
        DoiTuongID: cocId
      });
      
    } else if (status === 'FAILED') {
      // Update Coc status to failed
      await CocModel.capNhatTrangThaiCoc(cocId, {
        TrangThai: 'ThatBai'
      });
    }
    
    res.json({ success: true });
    
  } catch (error) {
    console.error('[PaymentController] Webhook error:', error);
    res.status(500).json({ 
      message: 'Server error' 
    });
  }
};
"""
    
    add_code_block(doc, webhook_handler)

def generate_section_4_5(doc):
    """4.5. Triển khai Frontend"""
    doc.add_heading('4.5. Triển khai Frontend', level=2)
    
    add_paragraph(doc, """Frontend được xây dựng với React 18+ và Vite, 
bao gồm 52 page components và 51 reusable components.""")
    
    # 4.5.1. Component Architecture
    doc.add_heading('4.5.1. Kiến trúc Component', level=3)
    
    component_breakdown = [
        ['Pages', '52 files', 'ChuDuAnDashboard, QuanLyTinDang, QuanLyCuocHen, etc.'],
        ['Components', '51 files', 'ModalTaoDuAn, TableTinDang, FormCuocHen, etc.'],
        ['Layouts', '2 files', 'MainLayout, DashboardLayout'],
        ['Services', '5 files', 'AuthService, TinDangService, CuocHenService'],
        ['API Clients', '7 files', 'apiClient, authAPI, tinDangAPI'],
        ['Utils', '3 files', 'formatters, validators, helpers']
    ]
    
    add_table_with_data(doc, ['Category', 'Count', 'Examples'], component_breakdown)
    
    doc.add_paragraph()
    
    # 4.5.2. Key Pages Implementation
    doc.add_heading('4.5.2. Triển khai Pages chính', level=3)
    
    add_paragraph(doc, bold=True, text="1. ChuDuAnDashboard - Trang tổng quan")
    
    dashboard_page = """
// client/src/pages/ChuDuAn/Dashboard/Dashboard.jsx
import { useState, useEffect } from 'react';
import { chuDuAnAPI } from '../../../api/chuDuAnAPI';
import './Dashboard.css';

export default function ChuDuAnDashboard() {
  const [stats, setStats] = useState(null);
  const [loading, setLoading] = useState(true);
  
  useEffect(() => {
    fetchStats();
  }, []);
  
  const fetchStats = async () => {
    try {
      const data = await chuDuAnAPI.layThongKe();
      setStats(data);
    } catch (error) {
      console.error('Error fetching stats:', error);
    } finally {
      setLoading(false);
    }
  };
  
  if (loading) return <div>Loading...</div>;
  
  return (
    <div className="chu-du-an-dashboard">
      <div className="chu-du-an-dashboard__header">
        <h1>Tổng quan</h1>
      </div>
      
      <div className="chu-du-an-dashboard__stats">
        <div className="chu-du-an-dashboard__stat-card">
          <div className="chu-du-an-dashboard__stat-value">
            {stats.tongDuAn}
          </div>
          <div className="chu-du-an-dashboard__stat-label">
            Dự án
          </div>
        </div>
        
        <div className="chu-du-an-dashboard__stat-card">
          <div className="chu-du-an-dashboard__stat-value">
            {stats.tongTinDang}
          </div>
          <div className="chu-du-an-dashboard__stat-label">
            Tin đăng
          </div>
        </div>
        
        <div className="chu-du-an-dashboard__stat-card">
          <div className="chu-du-an-dashboard__stat-value">
            {stats.tongCuocHen}
          </div>
          <div className="chu-du-an-dashboard__stat-label">
            Cuộc hẹn
          </div>
        </div>
        
        <div className="chu-du-an-dashboard__stat-card">
          <div className="chu-du-an-dashboard__stat-value">
            {stats.doanhThu?.toLocaleString('vi-VN')} đ
          </div>
          <div className="chu-du-an-dashboard__stat-label">
            Doanh thu
          </div>
        </div>
      </div>
      
      <div className="chu-du-an-dashboard__recent">
        <h2>Hoạt động gần đây</h2>
        {/* Recent activities list */}
      </div>
    </div>
  );
}
"""
    
    add_code_block(doc, dashboard_page)
    
    doc.add_paragraph()
    
    # 4.5.3. State Management
    doc.add_heading('4.5.3. State Management', level=3)
    
    add_paragraph(doc, """Hệ thống sử dụng React Context API để quản lý global state:""")
    
    state_contexts = [
        ['AuthContext', 'Quản lý authentication state, user info'],
        ['ChatContext', 'Quản lý chat messages, socket connection'],
        ['NotificationContext', 'Quản lý notifications, unread count']
    ]
    
    add_table_with_data(doc, ['Context', 'Purpose'], state_contexts)
    
    doc.add_paragraph()
    
    auth_context = """
// client/src/context/AuthContext.jsx
import { createContext, useContext, useState, useEffect } from 'react';
import { authAPI } from '../api/authAPI';

const AuthContext = createContext();

export const AuthProvider = ({ children }) => {
  const [user, setUser] = useState(null);
  const [loading, setLoading] = useState(true);
  
  useEffect(() => {
    checkAuth();
  }, []);
  
  const checkAuth = async () => {
    try {
      const token = localStorage.getItem('token');
      if (token) {
        const userData = await authAPI.getProfile();
        setUser(userData);
      }
    } catch (error) {
      localStorage.removeItem('token');
    } finally {
      setLoading(false);
    }
  };
  
  const login = async (email, password) => {
    const data = await authAPI.login(email, password);
    localStorage.setItem('token', data.token);
    setUser(data.user);
  };
  
  const logout = () => {
    localStorage.removeItem('token');
    setUser(null);
  };
  
  return (
    <AuthContext.Provider value={{ 
      user, 
      loading, 
      login, 
      logout 
    }}>
      {children}
    </AuthContext.Provider>
  );
};

export const useAuth = () => useContext(AuthContext);
"""
    
    add_code_block(doc, auth_context)

def generate_section_4_6(doc):
    """4.6. Security Implementation"""
    doc.add_heading('4.6. Triển khai Security', level=2)
    
    # Security measures
    security_measures = [
        ['Input Validation', 'express-validator cho tất cả inputs', 'Prevent injection attacks'],
        ['SQL Injection', 'Parameterized queries', 'Sử dụng mysql2 prepared statements (108 try-catch blocks)'],
        ['XSS Prevention', 'Content Security Policy', 'Sanitize user inputs với express-validator (14 validators)'],
        ['CSRF Protection', 'CSRF tokens', 'Protect state-changing operations'],
        ['Password Hashing', 'bcrypt', '2 implementation points'],
        ['JWT Security', 'JWT với 24h expiry', '9 implementation points'],
        ['HTTPS', 'Force HTTPS', 'Encrypt data in transit']
    ]
    
    add_table_with_data(doc, ['Security Measure', 'Implementation', 'Purpose'], security_measures)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Ví dụ: Input Validation Middleware")
    
    validation_example = """
// server/middleware/validation.js
const { body, validationResult } = require('express-validator');

exports.validateTinDang = [
  body('TieuDe')
    .trim()
    .notEmpty().withMessage('Tiêu đề không được để trống')
    .isLength({ max: 500 }).withMessage('Tiêu đề tối đa 500 ký tự'),
  
  body('MoTa')
    .trim()
    .notEmpty().withMessage('Mô tả không được để trống'),
  
  body('DienTich')
    .isFloat({ min: 1 }).withMessage('Diện tích phải > 0'),
  
  body('GiaThue')
    .isFloat({ min: 0 }).withMessage('Giá thuê phải >= 0'),
  
  body('DiaChi')
    .trim()
    .notEmpty().withMessage('Địa chỉ không được để trống'),
  
  (req, res, next) => {
    const errors = validationResult(req);
    if (!errors.isEmpty()) {
      return res.status(400).json({ 
        errors: errors.array() 
      });
    }
    next();
  }
];
"""
    
    add_code_block(doc, validation_example)

def generate_chuong_4(doc):
    """Generate full Chuong 4"""
    # Chapter title
    doc.add_heading('CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG', level=1)
    doc.add_paragraph()
    
    # Generate all sections
    generate_section_4_1(doc)
    doc.add_page_break()
    
    generate_section_4_2(doc)
    doc.add_page_break()
    
    generate_section_4_3(doc)
    doc.add_page_break()
    
    generate_section_4_4(doc)
    doc.add_page_break()
    
    generate_section_4_5(doc)
    doc.add_page_break()
    
    generate_section_4_6(doc)

def main():
    """Main function"""
    print("[*] Generating Chuong 4: TRIEN KHAI HE THONG...")
    
    # Create document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Add title page
    title = doc.add_heading('BÁO CÁO KHÓA LUẬN TỐT NGHIỆP', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph('HỆ THỐNG QUẢN LÝ CHO THUÊ PHÒNG TRỌ')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    subtitle2 = doc.add_paragraph('CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].italic = True
    subtitle2.runs[0].font.size = Pt(13)
    
    doc.add_page_break()
    
    # Generate chapter
    generate_chuong_4(doc)
    
    # Save document
    output_dir = os.path.join(os.path.dirname(__file__), '..')
    output_file = os.path.join(output_dir, 'BaoCao_Chuong4_FULL.docx')
    
    doc.save(output_file)
    
    print(f"[OK] Done! File saved to: {output_file}")
    print(f"[INFO] Total pages: ~25-30 pages (Chuong 4)")

if __name__ == "__main__":
    main()


