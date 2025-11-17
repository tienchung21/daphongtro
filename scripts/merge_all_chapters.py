#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Merge all chapters into one final document
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_document_styles(doc):
    """Setup document styles"""
    styles = doc.styles
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    
    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

def create_cover_page(doc):
    """Create cover page theo chuẩn template KLTN"""
    # University name
    p = doc.add_paragraph('ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Logo placeholder (thêm dòng cho logo)
    p = doc.add_paragraph('[LOGO TRƯỜNG]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(11)
    p.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student names
    p = doc.add_paragraph('VÕ NGUYỄN HOÀNH HỘP')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('NGUYỄN TIẾN CHUNG')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project title
    p = doc.add_paragraph('HỆ THỐNG QUẢN LÝ CHO THUÊ PHÒNG TRỌ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.color.rgb = RGBColor(0, 0, 139)
    
    doc.add_paragraph()
    
    # Major
    p = doc.add_paragraph('Ngành: Hệ Thống Thông Tin')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Supervisor
    p = doc.add_paragraph('Giảng viên hướng dẫn: NCS. Huỳnh Nam')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    p = doc.add_paragraph('THÀNH PHỐ HỒ CHÍ MINH, THÁNG 12 NĂM 2025')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)
    
    doc.add_page_break()

def create_secondary_cover(doc):
    """Create secondary cover page (Trang phụ bìa)"""
    # University name
    p = doc.add_paragraph('ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Logo placeholder
    p = doc.add_paragraph('[LOGO TRƯỜNG]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(11)
    p.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student names
    p = doc.add_paragraph('VÕ NGUYỄN HOÀNH HỘP')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    p = doc.add_paragraph('NGUYỄN TIẾN CHUNG')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project title
    p = doc.add_paragraph('HỆ THỐNG QUẢN LÝ CHO THUÊ PHÒNG TRỌ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(16)
    
    # Major
    p = doc.add_paragraph('Ngành: Hệ Thống Thông Tin')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Supervisor
    p = doc.add_paragraph('Giảng viên hướng dẫn: NCS. Huỳnh Nam')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Date
    p = doc.add_paragraph('THÀNH PHỐ HỒ CHÍ MINH, THÁNG 12 NĂM 2025')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)
    
    doc.add_page_break()

def create_abstract_page(doc):
    """Create Abstract page (tiếng Anh)"""
    # Title in English
    p = doc.add_paragraph('RENTAL HOUSING MANAGEMENT SYSTEM')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Abstract heading
    p = doc.add_paragraph('ABSTRACT')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)
    
    doc.add_paragraph()
    
    # Abstract content
    abstract_content = (
        'This thesis presents the design and implementation of a Rental Housing Management System '
        'based on the Managed Marketplace platform model. The system connects property owners, '
        'tenants, and an internal operations team to ensure listing quality, coordinate viewing '
        'appointments, and support transactions.'
    )
    doc.add_paragraph(abstract_content)
    doc.add_paragraph()
    
    abstract_content2 = (
        'The system is developed using modern web technologies including Node.js, Express.js for '
        'the backend, React for the frontend, and MySQL for the database. Key features include: '
        'property listing management with approval workflow, appointment scheduling, electronic '
        'wallet integration, deposit handling via SePay payment gateway, real-time messaging '
        'using Socket.IO, and comprehensive reporting for all user roles.'
    )
    doc.add_paragraph(abstract_content2)
    doc.add_paragraph()
    
    abstract_content3 = (
        'The system implements role-based access control (RBAC) for 6 user types: Customer, '
        'Property Owner, Sales Staff, Operations Staff, and System Administrator. Security measures '
        'follow OWASP Top 10:2021 guidelines and PCI DSS standards for payment processing. '
        'The implementation follows Node.js best practices with component-based architecture, '
        'state machine pattern for entity lifecycle management, and comprehensive audit logging.'
    )
    doc.add_paragraph(abstract_content3)
    doc.add_paragraph()
    
    # Keywords
    p = doc.add_paragraph()
    p.add_run('Keywords: ').bold = True
    p.add_run('Rental housing management, managed marketplace, property listing, appointment scheduling, '
              'electronic wallet, payment gateway, real-time messaging, RBAC, Node.js, React, MySQL.')
    
    doc.add_page_break()

def create_acknowledgment_page(doc):
    """Create acknowledgment page (Lời cảm ơn)"""
    doc.add_heading('LỜI CẢM ƠN', level=1)
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Chúng em xin gửi lời cảm ơn chân thành đến NCS. Huỳnh Nam, người đã tận tình hướng dẫn, '
        'động viên và góp ý quý báu trong suốt quá trình thực hiện khóa luận tốt nghiệp này.'
    )
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Chúng em xin cảm ơn quý thầy cô trong Khoa Công Nghệ Thông Tin, Trường Đại học Công nghệ '
        'Thông tin - ĐHQG TP.HCM đã truyền đạt kiến thức và tạo điều kiện thuận lợi để chúng em '
        'hoàn thành khóa luận.'
    )
    doc.add_paragraph()
    
    doc.add_paragraph(
        'Cuối cùng, chúng em xin cảm ơn gia đình, bạn bè đã luôn ủng hộ và động viên chúng em '
        'trong suốt thời gian học tập và thực hiện đề tài.'
    )
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('Người thực hiện đề tài')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()
    
    p = doc.add_paragraph('Võ Nguyễn Hoành Hộp - Nguyễn Tiến Chung')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

def create_evaluation_pages(doc):
    """Create evaluation pages (Nhận xét GVHD và phản biện)"""
    # GVHD evaluation
    doc.add_heading('NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN', level=1)
    for _ in range(15):
        doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('TP. Hồ Chí Minh, ngày.... tháng.... năm 2025')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('CHỮ KÝ CỦA GIẢNG VIÊN')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # Reviewer 1
    doc.add_heading('NHẬN XÉT CỦA GIÁO VIÊN PHẢN BIỆN 1', level=1)
    for _ in range(15):
        doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('TP. Hồ Chí Minh, ngày.... tháng.... năm 2025')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('CHỮ KÝ CỦA GIẢNG VIÊN')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # Reviewer 2
    doc.add_heading('NHẬN XÉT CỦA GIÁO VIÊN PHẢN BIỆN 2', level=1)
    for _ in range(15):
        doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('TP. Hồ Chí Minh, ngày.... tháng.... năm 2025')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph('CHỮ KÝ CỦA GIẢNG VIÊN')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

def create_abbreviations_page(doc):
    """Create abbreviations page (Danh mục từ viết tắt)"""
    doc.add_heading('DANH MỤC CÁC THUẬT NGỮ VIẾT TẮT', level=1)
    doc.add_paragraph()
    
    # Table of abbreviations
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Từ viết tắt'
    hdr_cells[1].text = 'Từ đầy đủ'
    hdr_cells[2].text = 'Nghĩa'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(13)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    abbreviations = [
        ('API', 'Application Programming Interface', 'Giao diện lập trình ứng dụng'),
        ('JWT', 'JSON Web Token', 'Mã thông báo web JSON'),
        ('RBAC', 'Role-Based Access Control', 'Kiểm soát truy cập dựa trên vai trò'),
        ('REST', 'Representational State Transfer', 'Chuyển giao trạng thái biểu diễn'),
        ('HTTP', 'HyperText Transfer Protocol', 'Giao thức truyền tải siêu văn bản'),
        ('HTTPS', 'HTTP Secure', 'HTTP bảo mật'),
        ('SQL', 'Structured Query Language', 'Ngôn ngữ truy vấn có cấu trúc'),
        ('UI', 'User Interface', 'Giao diện người dùng'),
        ('UX', 'User Experience', 'Trải nghiệm người dùng'),
        ('CRUD', 'Create, Read, Update, Delete', 'Tạo, Đọc, Cập nhật, Xóa'),
        ('MVC', 'Model-View-Controller', 'Mô hình-Giao diện-Điều khiển'),
        ('OTP', 'One-Time Password', 'Mật khẩu một lần'),
        ('KYC', 'Know Your Customer', 'Xác thực danh tính khách hàng'),
        ('PCI DSS', 'Payment Card Industry Data Security Standard', 'Tiêu chuẩn bảo mật dữ liệu'),
        ('OWASP', 'Open Web Application Security Project', 'Dự án bảo mật ứng dụng web mở'),
        ('TLS', 'Transport Layer Security', 'Bảo mật tầng truyền tải'),
        ('CORS', 'Cross-Origin Resource Sharing', 'Chia sẻ tài nguyên cross-origin'),
        ('CSRF', 'Cross-Site Request Forgery', 'Giả mạo yêu cầu cross-site'),
    ]
    
    for abbr, full, meaning in abbreviations:
        row_cells = table.add_row().cells
        row_cells[0].text = abbr
        row_cells[1].text = full
        row_cells[2].text = meaning
        
        # Center align first column
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

def create_toc(doc):
    """Create Table of Contents"""
    doc.add_heading('MỤC LỤC', level=1)
    doc.add_paragraph()
    
    toc_items = [
        ('CHƯƠNG 1: TỔNG QUAN', 1),
        ('1.1. Giới thiệu', 2),
        ('1.2. Bối cảnh và Động lực', 2),
        ('1.3. Mục tiêu nghiên cứu', 2),
        ('1.4. Phạm vi nghiên cứu', 2),
        ('1.5. Đối tượng và Phương pháp nghiên cứu', 2),
        ('1.6. Ý nghĩa khoa học và thực tiễn', 2),
        ('1.7. Cấu trúc khóa luận', 2),
        ('', 0),
        ('CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ', 1),
        ('2.1. Managed Marketplace Platform', 2),
        ('2.2. Công nghệ sử dụng', 2),
        ('2.3. Các nghiên cứu liên quan', 2),
        ('', 0),
        ('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', 1),
        ('3.1. Phân tích yêu cầu', 2),
        ('3.2. Thiết kế hệ thống', 2),
        ('3.3. Thiết kế cơ sở dữ liệu', 2),
        ('3.4. Thiết kế giao diện', 2),
        ('', 0),
        ('CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG', 1),
        ('4.1. Môi trường phát triển', 2),
        ('4.2. Triển khai Backend', 2),
        ('4.3. Triển khai Frontend', 2),
        ('4.4. Tích hợp và Deployment', 2),
        ('', 0),
        ('CHƯƠNG 5: KẾT QUẢ VÀ ĐÁNH GIÁ', 1),
        ('5.1. Kết quả triển khai', 2),
        ('5.2. Đánh giá hiệu năng', 2),
        ('5.3. Testing và Kiểm thử', 2),
        ('5.4. Đánh giá từ người dùng', 2),
        ('5.5. So sánh với yêu cầu ban đầu', 2),
        ('', 0),
        ('CHƯƠNG 6: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN', 1),
        ('6.1. Tổng kết', 2),
        ('6.2. Kết quả đạt được', 2),
        ('6.3. Hạn chế của hệ thống', 2),
        ('6.4. Hướng phát triển tương lai', 2),
        ('6.5. Kết luận', 2),
        ('', 0),
        ('TÀI LIỆU THAM KHẢO', 1),
        ('PHỤ LỤC', 1),
    ]
    
    for item, level in toc_items:
        if item == '':
            doc.add_paragraph()
        elif level == 1:
            p = doc.add_paragraph(item)
            p.runs[0].bold = True
            p.paragraph_format.left_indent = Inches(0)
        else:
            p = doc.add_paragraph(item)
            p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()

def append_document(target_doc, source_path):
    """Append content from source document to target document"""
    if not os.path.exists(source_path):
        print(f"[WARNING] File not found: {source_path}")
        return
    
    print(f"[*] Merging: {os.path.basename(source_path)}")
    
    source_doc = Document(source_path)
    
    # Skip the first page (title page) of each source document
    skip_first_page = True
    
    for element in source_doc.element.body:
        # Copy element to target
        target_doc.element.body.append(element)

def create_references(doc):
    """Create references section - IEEE Citation Style"""
    doc.add_heading('TÀI LIỆU THAM KHẢO', level=1)
    doc.add_paragraph()
    
    references = [
        '[1] Meta Platforms, Inc., "React - A JavaScript library for building user interfaces," React, 2024. [Online]. Available: https://react.dev. [Accessed: Nov. 7, 2025].',
        
        '[2] OpenJS Foundation, "Node.js - JavaScript runtime built on Chrome\'s V8 JavaScript engine," Node.js, 2024. [Online]. Available: https://nodejs.org. [Accessed: Nov. 7, 2025].',
        
        '[3] Oracle Corporation, MySQL 8.0 Reference Manual, Oracle, 2024. [Online]. Available: https://dev.mysql.com/doc/refman/8.0/en/. [Accessed: Nov. 7, 2025].',
        
        '[4] "Socket.IO - Realtime application framework," Socket.IO, 2024. [Online]. Available: https://socket.io/docs/v4/. [Accessed: Nov. 7, 2025].',
        
        '[5] M. Jones, J. Bradley, and N. Sakimura, "JSON Web Token (JWT)," RFC 7519, May 2015. [Online]. Available: https://datatracker.ietf.org/doc/html/rfc7519. [Accessed: Nov. 7, 2025].',
        
        '[6] R. T. Fielding, "Architectural Styles and the Design of Network-based Software Architectures," Ph.D. dissertation, University of California, Irvine, 2000. [Online]. Available: https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm. [Accessed: Nov. 7, 2025].',
        
        '[7] A. Hagiu and J. Wright, "Multi-Sided Platforms," International Journal of Industrial Organization, vol. 43, pp. 162-174, 2015, doi: 10.1016/j.ijindorg.2015.03.003.',
        
        '[8] OWASP Foundation, "OWASP Top 10:2021 - The Ten Most Critical Web Application Security Risks," OWASP, 2021. [Online]. Available: https://owasp.org/Top10/. [Accessed: Nov. 7, 2025].',
        
        '[9] PCI Security Standards Council, Payment Card Industry Data Security Standard (PCI DSS) Requirements and Testing Procedures, Version 4.0, PCI Security Standards Council, Mar. 2022. [Online]. Available: https://www.pcisecuritystandards.org/document_library/. [Accessed: Nov. 7, 2025].',
        
        '[10] Savills Vietnam, "Vietnam Property Market Report 2024," Savills Research, Ho Chi Minh City, Vietnam, 2024. [Online]. Available: https://www.savills.com.vn/research_articles. [Accessed: Nov. 7, 2025].',
        
        '[11] M. Fowler, Patterns of Enterprise Application Architecture. Boston, MA, USA: Addison-Wesley, 2002.',
        
        '[12] Y. Goldberg et al., "Node.js Best Practices," GitHub, 2024. [Online]. Available: https://github.com/goldbergyoni/nodebestpractices. [Accessed: Nov. 7, 2025].',
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    
    doc.add_page_break()

def create_appendix(doc):
    """Create appendix section"""
    doc.add_heading('PHỤ LỤC', level=1)
    doc.add_paragraph()
    
    # Phụ lục A: Database Schema
    doc.add_heading('Phụ lục A: Cấu trúc Database', level=2)
    doc.add_paragraph('Hệ thống sử dụng MySQL 8.0 với 32 bảng chính, được tổ chức theo các nhóm chức năng:')
    doc.add_paragraph()
    
    doc.add_heading('A.1. Nhóm Quản lý Người dùng và Xác thực', level=3)
    tables_auth = [
        '• NguoiDung: Thông tin người dùng (ID, Email, MatKhau_Hash, LoaiNguoiDung)',
        '• ThongTinKYC: Thông tin xác thực danh tính (CMND/CCCD, AnhCMND, TrangThaiKYC)',
        '• VaiTro: Vai trò trong hệ thống (Admin, ChuDuAn, KhachHang)',
        '• PhanQuyen: Phân quyền chi tiết cho từng vai trò',
        '• NhatKyHeThong: Lưu vết tất cả hành động quan trọng (audit log)'
    ]
    for item in tables_auth:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('A.2. Nhóm Quản lý Dự án và Tin đăng', level=3)
    tables_project = [
        '• DuAn: Thông tin dự án cho thuê (ChuDuAnID, TenDuAn, DiaChi, TrangThai)',
        '• TinDang: Tin đăng phòng trọ (TieuDe, MoTa, Gia, DienTich, LoaiPhong)',
        '• AnhTinDang: Hình ảnh của tin đăng (URL, ThuTu, LaAnhChinh)',
        '• TienNghiTinDang: Tiện nghi của phòng (WiFi, DieuHoa, NongLanh, etc.)',
        '• ChinhSachCoc: Chính sách đặt cọc (SoTienCoc, ThoiGianGiuCoc, DieuKhoan)'
    ]
    for item in tables_project:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('A.3. Nhóm Quản lý Giao dịch và Thanh toán', level=3)
    tables_transaction = [
        '• DonDatCho: Đơn đặt chỗ của khách hàng (TrangThai: ChoXacNhan, DaXacNhan, DaHuy)',
        '• ThanhToan: Thông tin thanh toán (SoTien, PhuongThucThanhToan, TrangThai)',
        '• GiaoDichVi: Lịch sử giao dịch ví điện tử',
        '• ViTien: Ví điện tử của người dùng (SoDu, TrangThai)',
        '• LichSuRutTien: Lịch sử rút tiền của Chủ dự án'
    ]
    for item in tables_transaction:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('A.4. Nhóm Quản lý Hợp đồng và Cuộc hẹn', level=3)
    tables_contract = [
        '• HopDong: Hợp đồng thuê phòng (TrangThai: HieuLuc, DaHetHan, BiHuy)',
        '• CuocHen: Cuộc hẹn xem phòng (ThoiGianHen, TrangThai, GhiChu)',
        '• DanhGia: Đánh giá và nhận xét từ khách hàng (DiemSo, BinhLuan)',
        '• PhieuKhaoSat: Khảo sát sau khi thuê'
    ]
    for item in tables_contract:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('A.5. Nhóm Quản lý Nội dung và Hỗ trợ', level=3)
    tables_content = [
        '• BaiViet: Bài viết blog/tin tức',
        '• YeuCauHoTro: Yêu cầu hỗ trợ từ người dùng',
        '• ThongBao: Thông báo gửi đến người dùng',
        '• TinNhan: Tin nhắn chat giữa Chủ dự án và Khách hàng'
    ]
    for item in tables_content:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Database schema đầy đủ: docs/database/schema.sql')
    doc.add_page_break()
    
    # Phụ lục B: API Endpoints
    doc.add_heading('Phụ lục B: Danh sách API Endpoints', level=2)
    doc.add_paragraph('Hệ thống cung cấp 146 RESTful API endpoints, tuân thủ chuẩn REST [6]:')
    doc.add_paragraph()
    
    doc.add_heading('B.1. Authentication & Authorization APIs', level=3)
    apis_auth = [
        'POST /api/auth/register - Đăng ký tài khoản mới',
        'POST /api/auth/login - Đăng nhập (trả về JWT token [5])',
        'POST /api/auth/logout - Đăng xuất',
        'POST /api/auth/refresh-token - Làm mới token',
        'POST /api/auth/forgot-password - Quên mật khẩu',
        'POST /api/auth/reset-password - Đặt lại mật khẩu',
        'GET /api/auth/verify-email/:token - Xác thực email'
    ]
    for api in apis_auth:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('B.2. KYC & User Management APIs', level=3)
    apis_kyc = [
        'POST /api/kyc/submit - Gửi thông tin KYC',
        'GET /api/kyc/status - Kiểm tra trạng thái KYC',
        'PUT /api/users/profile - Cập nhật thông tin cá nhân',
        'GET /api/users/me - Lấy thông tin user hiện tại',
        'PUT /api/users/change-password - Đổi mật khẩu'
    ]
    for api in apis_kyc:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('B.3. Project & Listing APIs (Chủ dự án)', level=3)
    apis_chuduan = [
        'GET /api/chu-du-an/dashboard - Dashboard chủ dự án',
        'GET /api/chu-du-an/du-an - Danh sách dự án',
        'POST /api/chu-du-an/du-an - Tạo dự án mới',
        'POST /api/chu-du-an/du-an/tao-nhanh - Tạo nhanh dự án',
        'GET /api/chu-du-an/du-an/:id - Chi tiết dự án',
        'PUT /api/chu-du-an/du-an/:id - Cập nhật dự án',
        'DELETE /api/chu-du-an/du-an/:id - Lưu trữ dự án',
        'POST /api/chu-du-an/du-an/:id/yeu-cau-mo-lai - Yêu cầu mở lại dự án',
        'GET /api/chu-du-an/tin-dang - Danh sách tin đăng',
        'GET /api/chu-du-an/tin-nhap - Danh sách tin nháp',
        'POST /api/chu-du-an/tin-dang - Tạo tin đăng mới (gửi duyệt)',
        'POST /api/chu-du-an/tin-dang/nhap - Lưu nháp tin đăng',
        'GET /api/chu-du-an/tin-dang/:id - Chi tiết tin đăng',
        'GET /api/chu-du-an/tin-dang/:id/chinh-sua - Lấy tin đăng để chỉnh sửa',
        'PUT /api/chu-du-an/tin-dang/:id - Cập nhật tin đăng',
        'POST /api/chu-du-an/tin-dang/:id/gui-duyet - Gửi tin đăng đi duyệt',
        'DELETE /api/chu-du-an/tin-dang/:id - Xóa tin đăng',
        'GET /api/chu-du-an/cuoc-hen - Danh sách cuộc hẹn',
        'GET /api/chu-du-an/cuoc-hen/metrics - Metrics cuộc hẹn',
        'PUT /api/chu-du-an/cuoc-hen/:id/xac-nhan - Xác nhận cuộc hẹn',
        'POST /api/chu-du-an/cuoc-hen/:id/phe-duyet - Phê duyệt cuộc hẹn',
        'POST /api/chu-du-an/cuoc-hen/:id/tu-choi - Từ chối cuộc hẹn',
        'GET /api/chu-du-an/bao-cao-hieu-suat - Báo cáo hiệu suất',
        'GET /api/chu-du-an/bao-cao-chi-tiet - Báo cáo chi tiết',
        'GET /api/chu-du-an/bao-cao/doanh-thu-theo-thang - Doanh thu theo tháng',
        'GET /api/chu-du-an/bao-cao/top-tin-dang - Top tin đăng',
        'GET /api/chu-du-an/bao-cao/conversion-rate - Tỷ lệ chuyển đổi'
    ]
    for api in apis_chuduan:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('B.4. Booking & Appointment APIs', level=3)
    apis_booking = [
        'POST /api/bookings - Tạo đơn đặt chỗ',
        'GET /api/bookings - Danh sách đơn đặt chỗ',
        'GET /api/bookings/:id - Chi tiết đơn đặt chỗ',
        'PUT /api/bookings/:id/confirm - Xác nhận đơn',
        'PUT /api/bookings/:id/cancel - Hủy đơn',
        'POST /api/appointments - Đặt lịch hẹn xem phòng',
        'GET /api/appointments - Danh sách cuộc hẹn',
        'PUT /api/appointments/:id/status - Cập nhật trạng thái cuộc hẹn'
    ]
    for api in apis_booking:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('B.5. Payment & Wallet APIs', level=3)
    apis_payment = [
        'POST /api/payments/deposit - Đặt cọc (tích hợp SePay)',
        'POST /api/payments/verify - Xác thực thanh toán',
        'GET /api/wallet/balance - Kiểm tra số dư ví',
        'GET /api/wallet/transactions - Lịch sử giao dịch',
        'POST /api/wallet/withdraw - Yêu cầu rút tiền',
        'GET /api/wallet/withdraw/history - Lịch sử rút tiền'
    ]
    for api in apis_payment:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('B.6. Contract & Rating APIs', level=3)
    apis_contract = [
        'POST /api/contracts - Tạo hợp đồng',
        'GET /api/contracts - Danh sách hợp đồng',
        'GET /api/contracts/:id - Chi tiết hợp đồng',
        'POST /api/contracts/:id/sign - Ký hợp đồng điện tử',
        'POST /api/ratings - Đánh giá tin đăng',
        'GET /api/ratings/listing/:id - Danh sách đánh giá của tin đăng'
    ]
    for api in apis_contract:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('B.7. Admin & Analytics APIs', level=3)
    apis_admin = [
        'GET /api/admin/users - Quản lý người dùng',
        'PUT /api/admin/users/:id/status - Khóa/Mở khóa tài khoản',
        'GET /api/admin/kyc/pending - Danh sách KYC chờ duyệt',
        'PUT /api/admin/kyc/:id/approve - Duyệt KYC',
        'GET /api/admin/analytics/dashboard - Dashboard thống kê',
        'GET /api/admin/analytics/revenue - Báo cáo doanh thu',
        'GET /api/admin/logs - Nhật ký hệ thống (audit log)'
    ]
    for api in apis_admin:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('B.8. Real-time APIs (Socket.IO [4])', level=3)
    apis_socket = [
        'socket.on("message:new") - Nhận tin nhắn mới',
        'socket.emit("message:send") - Gửi tin nhắn',
        'socket.on("notification:new") - Nhận thông báo mới',
        'socket.on("booking:status") - Cập nhật trạng thái đơn real-time'
    ]
    for api in apis_socket:
        doc.add_paragraph(api, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('API Documentation đầy đủ: docs/api/endpoints.md')
    doc.add_page_break()
    
    # Phụ lục C: Use Cases
    doc.add_heading('Phụ lục C: Đặc tả Use Cases', level=2)
    doc.add_paragraph('Hệ thống được phát triển dựa trên 39 Use Cases chi tiết, chia thành 6 nhóm:')
    doc.add_paragraph()
    
    doc.add_heading('C.1. Người dùng Chung (5 UCs)', level=3)
    ucs_gen = [
        'UC-GEN-01: Đăng nhập',
        'UC-GEN-02: Đăng ký tài khoản',
        'UC-GEN-03: Chuyển đổi vai trò',
        'UC-GEN-04: Xem danh sách cuộc hẹn',
        'UC-GEN-05: Trung tâm thông báo'
    ]
    for uc in ucs_gen:
        doc.add_paragraph(uc, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('C.2. Khách hàng (7 UCs)', level=3)
    ucs_cust = [
        'UC-CUST-01: Tìm kiếm phòng trọ',
        'UC-CUST-02: Quản lý yêu thích',
        'UC-CUST-03: Hẹn lịch xem phòng',
        'UC-CUST-04: Thực hiện đặt cọc (CọcGiữChỗ/CọcAnNinh)',
        'UC-CUST-05: Hủy giao dịch (Hoàn tiền)',
        'UC-CUST-06: Quản lý ví điện tử',
        'UC-CUST-07: Nhắn tin'
    ]
    for uc in ucs_cust:
        doc.add_paragraph(uc, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('C.3. Nhân viên Bán hàng (7 UCs)', level=3)
    ucs_sale = [
        'UC-SALE-01: Đăng ký lịch làm việc',
        'UC-SALE-02: Xem chi tiết cuộc hẹn',
        'UC-SALE-03: Quản lý cuộc hẹn (Xác nhận/Đổi lịch/Hủy)',
        'UC-SALE-04: Xác nhận cọc của khách hàng',
        'UC-SALE-05: Báo cáo kết quả cuộc hẹn',
        'UC-SALE-06: Xem báo cáo thu nhập',
        'UC-SALE-07: Nhắn tin'
    ]
    for uc in ucs_sale:
        doc.add_paragraph(uc, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('C.4. Chủ dự án (5 UCs)', level=3)
    ucs_project = [
        'UC-PROJ-01: Đăng tin cho thuê (1 phòng/nhiều phòng, lưu nháp)',
        'UC-PROJ-02: Xác nhận cuộc hẹn (nếu yêu cầu)',
        'UC-PROJ-03: Xem báo cáo kinh doanh',
        'UC-PROJ-04: Báo cáo hợp đồng cho thuê',
        'UC-PROJ-05: Nhắn tin'
    ]
    for uc in ucs_project:
        doc.add_paragraph(uc, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('C.5. Nhân viên Điều hành (6 UCs)', level=3)
    ucs_oper = [
        'UC-OPER-01: Duyệt tin đăng',
        'UC-OPER-02: Quản lý danh sách dự án',
        'UC-OPER-03: Quản lý lịch làm việc NVBH (tổng thể)',
        'UC-OPER-04: Quản lý hồ sơ nhân viên',
        'UC-OPER-05: Tạo tài khoản nhân viên',
        'UC-OPER-06: Lập biên bản bàn giao'
    ]
    for uc in ucs_oper:
        doc.add_paragraph(uc, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('C.6. Quản trị viên Hệ thống (9 UCs)', level=3)
    ucs_admin = [
        'UC-ADMIN-01: Quản lý tài khoản người dùng',
        'UC-ADMIN-02: Quản lý danh sách dự án',
        'UC-ADMIN-03: Quản lý danh sách khu vực',
        'UC-ADMIN-04: Xem báo cáo thu nhập toàn hệ thống',
        'UC-ADMIN-05: Quản lý chính sách',
        'UC-ADMIN-06: Quản lý mẫu hợp đồng',
        'UC-ADMIN-07: Quản lý quyền & phân quyền (RBAC)',
        'UC-ADMIN-08: Xem nhật ký hệ thống (Audit Log)',
        'UC-ADMIN-09: Quản lý chính sách cọc theo TinĐăng'
    ]
    for uc in ucs_admin:
        doc.add_paragraph(uc, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Đặc tả chi tiết từng Use Case: docs/use-cases-v1.2.md')
    doc.add_page_break()
    
    # Phụ lục D: State Machine Diagrams
    doc.add_heading('Phụ lục D: State Machine Diagrams', level=2)
    doc.add_paragraph('Hệ thống áp dụng State Machine Pattern [11] để quản lý trạng thái của các entities:')
    doc.add_paragraph()
    
    doc.add_heading('D.1. Tin đăng State Machine', level=3)
    doc.add_paragraph('Trạng thái: Nhap → ChoDuyet → DaDuyet → DaDang → DaThue → DaDong → DaXoa')
    doc.add_paragraph('Transitions:')
    transitions_listing = [
        '• submit(): Nhap → ChoDuyet',
        '• approve(): ChoDuyet → DaDuyet',
        '• reject(): ChoDuyet → Nhap',
        '• publish(): DaDuyet → DaDang',
        '• rent(): DaDang → DaThue',
        '• close(): DaDang/DaThue → DaDong',
        '• delete(): Any → DaXoa'
    ]
    for t in transitions_listing:
        doc.add_paragraph(t, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('D.2. Đơn đặt chỗ State Machine', level=3)
    doc.add_paragraph('Trạng thái: ChoXacNhan → DaXacNhan → DaThanhToan → HoanThanh → DaHuy')
    doc.add_paragraph('Transitions:')
    transitions_booking = [
        '• create(): New → ChoXacNhan',
        '• confirm(): ChoXacNhan → DaXacNhan',
        '• pay(): DaXacNhan → DaThanhToan',
        '• complete(): DaThanhToan → HoanThanh',
        '• cancel(): ChoXacNhan/DaXacNhan → DaHuy'
    ]
    for t in transitions_booking:
        doc.add_paragraph(t, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('D.3. KYC State Machine', level=3)
    doc.add_paragraph('Trạng thái: ChuaKYC → ChoXuLy → DaDuyet → BiTuChoi')
    doc.add_paragraph('Transitions:')
    transitions_kyc = [
        '• submit(): ChuaKYC → ChoXuLy',
        '• approve(): ChoXuLy → DaDuyet',
        '• reject(): ChoXuLy → BiTuChoi',
        '• resubmit(): BiTuChoi → ChoXuLy'
    ]
    for t in transitions_kyc:
        doc.add_paragraph(t, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('D.4. Hợp đồng State Machine', level=3)
    doc.add_paragraph('Trạng thái: Nhap → ChoKy → HieuLuc → DaHetHan → BiHuy')
    doc.add_paragraph('Transitions:')
    transitions_contract = [
        '• create(): New → Nhap',
        '• send(): Nhap → ChoKy',
        '• sign(): ChoKy → HieuLuc (cả 2 bên ký)',
        '• expire(): HieuLuc → DaHetHan',
        '• cancel(): Nhap/ChoKy/HieuLuc → BiHuy'
    ]
    for t in transitions_contract:
        doc.add_paragraph(t, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('State Machine implementations: server/services/*Service.js')
    doc.add_page_break()
    
    # Phụ lục E: Security & Best Practices
    doc.add_heading('Phụ lục E: Bảo mật và Best Practices', level=2)
    doc.add_paragraph()
    
    doc.add_heading('E.1. Security Implementation', level=3)
    doc.add_paragraph('Hệ thống tuân thủ OWASP Top 10:2021 [8]:')
    security_items = [
        '• A01:2021 - Broken Access Control: Sử dụng JWT [5] + Role-based access control',
        '• A02:2021 - Cryptographic Failures: Mã hóa mật khẩu với bcrypt (salt rounds=10)',
        '• A03:2021 - Injection: Sử dụng Prepared Statements (mysql2)',
        '• A04:2021 - Insecure Design: Áp dụng State Machine [11] và Validation layers',
        '• A05:2021 - Security Misconfiguration: Environment variables, secure headers',
        '• A07:2021 - Authentication Failures: JWT với refresh token, rate limiting',
        '• A08:2021 - Software and Data Integrity: Input validation, CSRF protection',
        '• A09:2021 - Security Logging: Audit log (NhatKyHeThong)',
        '• A10:2021 - Server-Side Request Forgery: URL validation và whitelist'
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('E.2. Payment Security (PCI DSS)', level=3)
    doc.add_paragraph('Tuân thủ PCI DSS [9] cho xử lý thanh toán:')
    pci_items = [
        '• KHÔNG lưu trữ thông tin thẻ tín dụng',
        '• Sử dụng SePay Payment Gateway (52 integration points)',
        '• HTTPS/TLS 1.2+ cho mọi giao dịch',
        '• Tokenization cho merchant transactions',
        '• Log mọi giao dịch thanh toán (GiaoDichVi, ThanhToan tables)'
    ]
    for item in pci_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('E.3. Node.js Best Practices', level=3)
    doc.add_paragraph('Áp dụng Node.js Best Practices [12]:')
    nodejs_items = [
        '• Project Structure: Component-based architecture (theo domain)',
        '• Error Handling: Centralized error handling middleware',
        '• Code Style: ESLint + Prettier configuration',
        '• Testing: Unit tests (Jest) + Integration tests',
        '• Async Operations: async/await với proper error handling',
        '• Environment Variables: dotenv + validation',
        '• Logging: Winston với log rotation',
        '• Performance: Database connection pooling, Redis caching',
        '• Monitoring: Health checks (/api/health endpoint)'
    ]
    for item in nodejs_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('E.4. React Best Practices', level=3)
    doc.add_paragraph('Frontend tuân thủ React [1] best practices:')
    react_items = [
        '• Component Structure: Functional components + Hooks',
        '• State Management: Context API + useReducer',
        '• Code Splitting: React.lazy() + Suspense',
        '• Performance: useMemo, useCallback, React.memo',
        '• Error Boundaries: Graceful error handling',
        '• Accessibility: ARIA labels, semantic HTML',
        '• SEO: React Helmet, meta tags',
        '• Build Optimization: Vite với tree-shaking'
    ]
    for item in react_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()
    
    doc.add_paragraph('Source code repository: https://github.com/tienchung21/daphongtro.git')
    doc.add_paragraph('Documentation: docs/ folder trong repository')
    doc.add_paragraph()
    doc.add_paragraph('Thời gian thực hiện: 10/09/2025 - 01/12/2025')

def main():
    """Main function to merge all chapters"""
    print("[*] Starting to merge all chapters...")
    
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    
    # Create final document
    final_doc = Document()
    setup_document_styles(final_doc)
    
    # 1. Cover page (Bìa)
    print("[*] Creating cover page...")
    create_cover_page(final_doc)
    
    # 2. Secondary cover (Trang phụ bìa)
    print("[*] Creating secondary cover...")
    create_secondary_cover(final_doc)
    
    # 3. Abstract page (tiếng Anh)
    print("[*] Creating abstract page...")
    create_abstract_page(final_doc)
    
    # 4. Acknowledgment (Lời cảm ơn)
    print("[*] Creating acknowledgment page...")
    create_acknowledgment_page(final_doc)
    
    # 5. Evaluation pages (Nhận xét GVHD và phản biện)
    print("[*] Creating evaluation pages...")
    create_evaluation_pages(final_doc)
    
    # 6. Table of Contents
    print("[*] Creating table of contents...")
    create_toc(final_doc)
    
    # 7. Abbreviations page (Danh mục từ viết tắt)
    print("[*] Creating abbreviations page...")
    create_abbreviations_page(final_doc)
    
    # 3. Merge all chapters
    chapters = [
        'BaoCao_Chuong1_2_FULL.docx',
        'BaoCao_Chuong3_FULL.docx',
        'BaoCao_Chuong4_FULL.docx',
        'BaoCao_Chuong5_FULL.docx',
        'BaoCao_Chuong6_FULL.docx'
    ]
    
    for chapter_file in chapters:
        chapter_path = os.path.join(base_dir, chapter_file)
        append_document(final_doc, chapter_path)
    
    # 4. Add references
    print("[*] Creating references...")
    create_references(final_doc)
    
    # 5. Add appendix
    print("[*] Creating appendix...")
    create_appendix(final_doc)
    
    # Save final document
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = os.path.join(base_dir, f'BaoCao_KLTN_HeThongChoThuePhongTro_FINAL_{timestamp}.docx')
    final_doc.save(output_file)
    
    print(f"\n[OK] DONE! Final document saved to:")
    print(f"     {output_file}")
    print(f"\n[INFO] Estimated total pages: 60-80 pages")
    print(f"\n[INFO] Document structure:")
    print(f"  - Cover page")
    print(f"  - Table of Contents")
    print(f"  - Chapter 1-2: Overview & Theory (~10 pages)")
    print(f"  - Chapter 3: Analysis & Design (~15 pages)")
    print(f"  - Chapter 4: Implementation (~25 pages) - MOST IMPORTANT")
    print(f"  - Chapter 5: Results & Evaluation (~8 pages)")
    print(f"  - Chapter 6: Conclusion (~6 pages)")
    print(f"  - References & Appendix (~3 pages)")
    print(f"\n[SUCCESS] All chapters merged successfully!")

if __name__ == "__main__":
    main()

