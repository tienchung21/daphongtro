#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word Document Creator Helper
Tạo và chỉnh sửa file .docx với đầy đủ tính năng

Sử dụng:
    python scripts/create_docx.py --title "Tiêu đề" --content "Nội dung" --output output.docx
    python scripts/create_docx.py --json input.json --output output.docx
"""

import argparse
import json
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, url, text):
    """
    Thêm hyperlink vào paragraph
    """
    # Create the w:hyperlink tag and add needed values
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )
    
    # Create a new run object (text)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Set color and underline
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink


def create_document(config):
    """
    Tạo document từ config
    
    Config structure:
    {
        "title": "Tiêu đề chính",
        "sections": [
            {
                "type": "heading1", 
                "text": "Heading level 1"
            },
            {
                "type": "heading2",
                "text": "Heading level 2"
            },
            {
                "type": "paragraph",
                "text": "Đoạn văn bản",
                "bold": false,
                "italic": false,
                "font_size": 12,
                "color": "#000000",
                "align": "left"  // left, center, right, justify
            },
            {
                "type": "list",
                "items": ["Item 1", "Item 2", "Item 3"]
            },
            {
                "type": "table",
                "headers": ["Cột 1", "Cột 2", "Cột 3"],
                "rows": [
                    ["Dữ liệu 1", "Dữ liệu 2", "Dữ liệu 3"],
                    ["Dữ liệu 4", "Dữ liệu 5", "Dữ liệu 6"]
                ]
            },
            {
                "type": "image",
                "path": "path/to/image.png",
                "width": 6  // inches
            },
            {
                "type": "link",
                "text": "Click here",
                "url": "https://example.com"
            },
            {
                "type": "page_break"
            }
        ]
    }
    """
    doc = Document()
    
    # Add title if provided
    if 'title' in config and config['title']:
        title = doc.add_heading(config['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process sections
    for section in config.get('sections', []):
        section_type = section.get('type', 'paragraph')
        
        if section_type == 'heading1':
            doc.add_heading(section.get('text', ''), level=1)
            
        elif section_type == 'heading2':
            doc.add_heading(section.get('text', ''), level=2)
            
        elif section_type == 'heading3':
            doc.add_heading(section.get('text', ''), level=3)
            
        elif section_type == 'paragraph':
            p = doc.add_paragraph()
            run = p.add_run(section.get('text', ''))
            
            # Apply formatting
            if section.get('bold'):
                run.bold = True
            if section.get('italic'):
                run.italic = True
            if section.get('underline'):
                run.underline = True
                
            # Font size
            if 'font_size' in section:
                run.font.size = Pt(section['font_size'])
            
            # Color
            if 'color' in section:
                color = section['color'].lstrip('#')
                rgb = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
                run.font.color.rgb = RGBColor(*rgb)
            
            # Alignment
            align = section.get('align', 'left').upper()
            if align == 'CENTER':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'RIGHT':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'JUSTIFY':
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
        elif section_type == 'list':
            for item in section.get('items', []):
                doc.add_paragraph(item, style='List Bullet')
                
        elif section_type == 'numbered_list':
            for item in section.get('items', []):
                doc.add_paragraph(item, style='List Number')
                
        elif section_type == 'table':
            headers = section.get('headers', [])
            rows = section.get('rows', [])
            
            if headers and rows:
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Headers
                for i, header in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = header
                    # Bold header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Data rows
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx + 1].cells[col_idx].text = str(cell_data)
                        
        elif section_type == 'image':
            try:
                width = section.get('width', 6)
                doc.add_picture(section.get('path'), width=Inches(width))
            except Exception as e:
                print(f"[WARN] Khong the them anh: {e}", file=sys.stderr)
                
        elif section_type == 'link':
            p = doc.add_paragraph()
            add_hyperlink(p, section.get('url', ''), section.get('text', ''))
            
        elif section_type == 'page_break':
            doc.add_page_break()
    
    return doc


def main():
    parser = argparse.ArgumentParser(
        description='Tạo file Word .docx với đầy đủ tính năng',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Ví dụ sử dụng:

1. Tạo document đơn giản:
   python scripts/create_docx.py --title "Báo cáo" --content "Nội dung báo cáo" --output report.docx

2. Tạo document từ JSON file:
   python scripts/create_docx.py --json config.json --output output.docx

3. Tạo document từ JSON string:
   python scripts/create_docx.py --json '{"title":"Test","sections":[{"type":"paragraph","text":"Hello"}]}' --output test.docx
        """
    )
    
    parser.add_argument('--json', help='JSON config (file path hoặc JSON string)')
    parser.add_argument('--title', help='Tiêu đề document')
    parser.add_argument('--content', help='Nội dung đơn giản')
    parser.add_argument('--output', '-o', required=True, help='Tên file output (.docx)')
    
    args = parser.parse_args()
    
    # Build config
    config = {}
    
    if args.json:
        # Try to load as file first
        try:
            with open(args.json, 'r', encoding='utf-8') as f:
                config = json.load(f)
        except FileNotFoundError:
            # Try to parse as JSON string
            try:
                config = json.loads(args.json)
            except json.JSONDecodeError as e:
                print(f"[ERROR] Loi parse JSON: {e}", file=sys.stderr)
                sys.exit(1)
    else:
        # Simple mode with title and content
        if args.title:
            config['title'] = args.title
        
        if args.content:
            config['sections'] = [
                {
                    'type': 'paragraph',
                    'text': args.content
                }
            ]
    
    # Create document
    try:
        doc = create_document(config)
        doc.save(args.output)
        print(f"[OK] Da tao file: {args.output}")
    except Exception as e:
        print(f"[ERROR] Loi tao document: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()

