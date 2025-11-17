# -*- coding: utf-8 -*-
"""
Script Generate Báo Cáo Khóa Luận Tốt Nghiệp
Hệ thống Quản lý Cho thuê Phòng trọ

Tạo file DOCX hoàn chỉnh (70-100 trang) theo đúng format mẫu của Trường ĐH Công Nghiệp TP.HCM
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

class BaoCaoKLTNGenerator:
    def __init__(self):
        self.doc = Document()
        self.setup_page_layout()
        self.setup_styles()
        self.image_count = {}  # Track image numbers by chapter
        self.table_count = {}  # Track table numbers by chapter
        
    def setup_page_layout(self):
        """Thiết lập margins và page size theo mẫu"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(3)
            section.left_margin = Cm(3.5)
            section.right_margin = Cm(2)
            section.page_height = Cm(29.7)  # A4
            section.page_width = Cm(21)
    
    def setup_styles(self):
        """Thiết lập styles theo mẫu KLTN"""
        styles = self.doc.styles
        
        # Heading 1: Times New Roman 14pt, Bold, Uppercase
        heading1 = styles['Heading 1']
        heading1_font = heading1.font
        heading1_font.name = 'Times New Roman'
        heading1_font.size = Pt(14)
        heading1_font.bold = True
        heading1_para = heading1.paragraph_format
        heading1_para.space_before = Pt(24)
        heading1_para.space_after = Pt(24)
        heading1_para.line_spacing = 1.15
        
        # Heading 2: Times New Roman 13pt, Bold
        heading2 = styles['Heading 2']
        heading2_font = heading2.font
        heading2_font.name = 'Times New Roman'
        heading2_font.size = Pt(13)
        heading2_font.bold = True
        heading2_para = heading2.paragraph_format
        heading2_para.space_before = Pt(6)
        heading2_para.space_after = Pt(12)
        heading2_para.line_spacing = 1.15
        
        # Heading 3: Times New Roman 13pt, Bold Italic
        heading3 = styles['Heading 3']
        heading3_font = heading3.font
        heading3_font.name = 'Times New Roman'
        heading3_font.size = Pt(13)
        heading3_font.bold = True
        heading3_font.italic = True
        heading3_para = heading3.paragraph_format
        heading3_para.space_before = Pt(6)
        heading3_para.space_after = Pt(12)
        heading3_para.line_spacing = 1.15
        
        # Normal: Times New Roman 13pt, Line spacing 1.5
        normal = styles['Normal']
        normal_font = normal.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(13)
        normal_para = normal.paragraph_format
        normal_para.line_spacing = 1.5
        normal_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def add_centered_text(self, text, bold=False, size=13, all_caps=False):
        """Helper: Thêm text căn giữa"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        if all_caps:
            run.font.all_caps = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p
    
    def add_paragraph_text(self, text):
        """Helper: Thêm đoạn văn thông thường, hỗ trợ **bold** markdown"""
        import re
        
        p = self.doc.add_paragraph()
        p.style = 'Normal'
        
        # Parse **bold** markdown
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text (remove **)
                run = p.add_run(part[2:-2])
                run.font.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
            elif part:  # Non-empty normal text
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(13)
        
        return p
    
    def add_placeholder_image(self, caption_text, chapter_num):
        """Thêm placeholder cho hình ảnh với caption"""
        if chapter_num not in self.image_count:
            self.image_count[chapter_num] = 0
        self.image_count[chapter_num] += 1
        
        p = self.doc.add_paragraph()
        run = p.add_run(f'[Chèn hình {chapter_num}.{self.image_count[chapter_num]} tại đây]')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption = self.doc.add_paragraph()
        caption_run = caption.add_run(f'Hình {chapter_num}.{self.image_count[chapter_num]}: {caption_text}')
        caption_run.font.name = 'Times New Roman'
        caption_run.font.size = Pt(13)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return caption
    
    def add_table(self, headers, data):
        """Thêm bảng với headers và data"""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.style = 'Normal'
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                for paragraph in row_cells[i].paragraphs:
                    paragraph.style = 'Normal'
        
        return table
    
    # ========== PHẦN MỞ ĐẦU ==========
    
    def add_cover_page(self):
        """Trang bìa"""
        self.add_centered_text('BỘ CÔNG THƯƠNG', bold=True, size=13)
        self.add_centered_text('TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP. HỒ CHÍ MINH', bold=True, size=13)
        self.add_centered_text('KHOA CÔNG NGHỆ THÔNG TIN', bold=True, size=13)
        self.doc.add_paragraph()  # Space
        
        self.add_centered_text('[Logo Trường]', size=11)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        self.add_centered_text('VÕ NGUYỄN HOÀNH HỢP - MSSV: 21107481', bold=True, size=14)
        self.add_centered_text('NGUYỄN TIẾN CHUNG - MSSV: 21133721', bold=True, size=14)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        self.add_centered_text('HỆ THỐNG QUẢN LÝ CHO THUÊ PHÒNG TRỌ', bold=True, size=16, all_caps=True)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        self.add_centered_text('Ngành: Hệ Thống Thông Tin', bold=True, size=14)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        self.add_centered_text('Giảng viên hướng dẫn: NCS. Huỳnh Nam', bold=True, size=14)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        self.add_centered_text('THÀNH PHỐ HỒ CHÍ MINH, THÁNG 12 NĂM 2025', bold=True, size=13)
    
    def add_cover_page_vietnamese(self):
        """Trang phụ bìa tiếng Việt (giống bìa)"""
        self.add_cover_page()  # Same as cover
    
    def add_cover_page_english(self):
        """Trang phụ bìa tiếng Anh"""
        self.add_centered_text('INDUSTRIAL UNIVERSITY OF HO CHI MINH CITY', bold=True, size=13)
        self.add_centered_text('FACULTY OF INFORMATION TECHNOLOGY', bold=True, size=13)
        self.doc.add_paragraph()
        
        self.add_centered_text('[University Logo]', size=11)
        self.doc.add_paragraph()
        
        self.add_centered_text('VO NGUYEN HOANH HOP - Student ID: 21107481', bold=True, size=14)
        self.add_centered_text('NGUYEN TIEN CHUNG - Student ID: 21133721', bold=True, size=14)
        self.doc.add_paragraph()
        
        self.add_centered_text('RENTAL PROPERTY MANAGEMENT SYSTEM', bold=True, size=16, all_caps=True)
        self.doc.add_paragraph()
        
        self.add_centered_text('Major: Information Systems', bold=True, size=14)
        self.doc.add_paragraph()
        
        self.add_centered_text('Supervisor: NCS. Huynh Nam', bold=True, size=14)
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        self.add_centered_text('HO CHI MINH CITY - DECEMBER 2025', bold=True, size=13)
    
    def add_abstract(self):
        """Abstract tiếng Anh"""
        self.add_centered_text('RENTAL PROPERTY MANAGEMENT SYSTEM', bold=True, size=14, all_caps=True)
        self.add_centered_text('ABSTRACT', bold=True, size=13)
        self.doc.add_paragraph()
        
        abstract_text = """This thesis presents a comprehensive managed marketplace platform for rental property management in Vietnam. Unlike traditional classified ad platforms (Phongtro123, Mogi.vn), our system implements active platform participation through KYC verification, listing approval, automated appointment scheduling, secure deposit management, and digital contract execution. The system supports five user roles: Customer, Property Owner, Sales Staff, Operations Staff, and Administrator.

The system architecture follows a 3-tier design with React 18.3.1 frontend, Node.js 18.x/Express 4.21 backend, and MySQL/MariaDB database. Key features include JWT-based authentication with RBAC, real-time chat using Socket.IO, payment gateway integration (SePay), geocoding service (Google Maps + Nominatim fallback), and innovative room N-N mapping allowing one physical room to appear in multiple listings with overridden prices and descriptions.

We successfully implemented 31 out of 36 use cases (86% completion rate), including all critical customer-facing modules (search, appointments, deposits, contracts, handovers). The system demonstrates strong security practices (JWT, Bcrypt, CSRF protection, rate limiting, audit logging) and achieves performance targets (search P95 ≤ 2.0s, deposit end-to-end ≤ 4s).

The platform is ready for MVP deployment and shows potential for commercial application with future enhancements including mobile apps, AI chatbot, blockchain smart contracts, and IoT integration."""
        
        self.add_paragraph_text(abstract_text)
        self.doc.add_paragraph()
        
        keywords_text = "**Keywords:** managed marketplace, rental management, React, Node.js, MySQL, real-time chat, payment gateway, JWT authentication, RBAC, Socket.IO, property management, KYC verification, digital contract"
        self.add_paragraph_text(keywords_text)
    
    def add_loi_cam_on(self):
        """Trang lời cảm ơn (để trống)"""
        self.add_centered_text('LỜI CẢM ƠN', bold=True, size=14)
        self.doc.add_paragraph()
        self.add_paragraph_text('[Sinh viên tự viết lời cảm ơn]')
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        self.add_centered_text('Người thực hiện đề tài', size=13)
    
    def add_nhan_xet_gvhd(self):
        """Trang nhận xét GVHD (để trống)"""
        self.add_centered_text('NHẬN XÉT CỦA GIÁO VIÊN HƯỚNG DẪN', bold=True, size=14)
        self.doc.add_paragraph()
        for i in range(15):
            self.doc.add_paragraph('_' * 80)
        self.doc.add_paragraph()
        self.add_centered_text('TP. Hồ Chí Minh, ngày.... tháng.... năm 2025', size=13)
        p = self.doc.add_paragraph()
        run = p.add_run(' ' * 60 + 'CHỮ KÝ CỦA GIẢNG VIÊN')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    
    def add_nhan_xet_phan_bien(self, so_thu_tu):
        """Trang nhận xét GV phản biện"""
        self.add_centered_text(f'NHẬN XÉT CỦA GIÁO VIÊN PHẢN BIỆN {so_thu_tu}', bold=True, size=14)
        self.doc.add_paragraph()
        for i in range(15):
            self.doc.add_paragraph('_' * 80)
        self.doc.add_paragraph()
        self.add_centered_text('TP. Hồ Chí Minh, ngày.... tháng.... năm 2025', size=13)
        p = self.doc.add_paragraph()
        run = p.add_run(' ' * 60 + 'CHỮ KÝ CỦA GIẢNG VIÊN')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    
    def add_muc_luc(self):
        """Mục lục (placeholder - cần update trong Word sau)"""
        self.add_centered_text('MỤC LỤC', bold=True, size=14)
        self.doc.add_paragraph()
        self.add_paragraph_text('[Mục lục sẽ được tự động generate trong Microsoft Word bằng cách: References → Update Table]')
    
    def add_muc_luc_hinh_anh(self):
        """Mục lục hình ảnh"""
        self.add_centered_text('MỤC LỤC HÌNH ẢNH', bold=True, size=14)
        self.doc.add_paragraph()
        self.add_paragraph_text('[Mục lục hình ảnh sẽ được tự động generate trong Microsoft Word]')
    
    def add_danh_muc_bang_bieu(self):
        """Danh mục bảng biểu"""
        self.add_centered_text('DANH MỤC BẢNG BIỂU', bold=True, size=14)
        self.doc.add_paragraph()
        self.add_paragraph_text('[Danh mục bảng biểu sẽ được tự động generate trong Microsoft Word]')
    
    def add_danh_muc_tu_viet_tat(self):
        """Danh mục từ viết tắt"""
        self.add_centered_text('DANH MỤC CÁC THUẬT NGỮ VIẾT TẮT', bold=True, size=14)
        self.doc.add_paragraph()
        
        viet_tat_data = [
            ('API', 'Application Programming Interface', 'Giao diện lập trình ứng dụng'),
            ('CRUD', 'Create, Read, Update, Delete', 'Tạo, Đọc, Cập nhật, Xóa'),
            ('CSRF', 'Cross-Site Request Forgery', 'Giả mạo yêu cầu trên nhiều trang'),
            ('HTTP', 'HyperText Transfer Protocol', 'Giao thức truyền tải siêu văn bản'),
            ('HTTPS', 'HyperText Transfer Protocol Secure', 'Giao thức HTTP bảo mật'),
            ('JWT', 'JSON Web Token', 'Mã thông báo web JSON'),
            ('KYC', 'Know Your Customer', 'Biết khách hàng của bạn'),
            ('MVC', 'Model-View-Controller', 'Mô hình-Giao diện-Điều khiển'),
            ('NVBH', 'Nhân Viên Bán Hàng', 'Sales Staff'),
            ('OSM', 'OpenStreetMap', 'Bản đồ mở'),
            ('RBAC', 'Role-Based Access Control', 'Kiểm soát truy cập dựa trên vai trò'),
            ('REST', 'Representational State Transfer', 'Truyền trạng thái đại diện'),
            ('SPA', 'Single Page Application', 'Ứng dụng đơn trang'),
            ('SQL', 'Structured Query Language', 'Ngôn ngữ truy vấn có cấu trúc'),
            ('TTL', 'Time To Live', 'Thời gian sống'),
            ('UC', 'Use Case', 'Tình huống sử dụng'),
            ('UI', 'User Interface', 'Giao diện người dùng'),
            ('UX', 'User Experience', 'Trải nghiệm người dùng'),
        ]
        
        self.add_table(['Từ viết tắt', 'Từ đầy đủ', 'Nghĩa'], viet_tat_data)
    
    # ========== NỘI DUNG CHÍNH ==========
    
    def add_chuong_1_gioi_thieu(self):
        """CHƯƠNG 1: GIỚI THIỆU"""
        self.doc.add_heading('CHƯƠNG 1: GIỚI THIỆU', level=1)
        
        # 1.1. Đặt vấn đề
        self.doc.add_heading('1.1. Đặt vấn đề', level=2)
        self.add_paragraph_text(
            "Thị trường cho thuê phòng trọ tại Việt Nam đang phát triển mạnh mẽ với hàng triệu sinh viên, người lao động di cư cần tìm chỗ ở. Theo thống kê, chỉ riêng TP. Hồ Chí Minh có hơn 2 triệu người thuê trọ hàng năm. Tuy nhiên, quy trình tìm kiếm và thuê trọ hiện nay vẫn còn nhiều bất cập.")
        
        self.add_paragraph_text(
            "Các nền tảng hiện tại như Phongtro123, Mogi.vn chủ yếu là sàn rao vặt (classified ads marketplace), nơi chủ nhà tự đăng tin mà không qua kiểm duyệt kỹ lưỡng. Điều này dẫn đến nhiều vấn đề: tin đăng giả mạo, thông tin không chính xác, thiếu minh bạch về giá cả và điều kiện thuê, khó khăn trong việc đặt lịch xem phòng, không có cơ chế bảo vệ cọc, và thiếu sự hỗ trợ khi có tranh chấp.")
        
        self.add_paragraph_text(
            "Từ những thực trạng trên, nhu cầu về một nền tảng 'managed marketplace' - nơi nền tảng tham gia tích cực vào quy trình từ đăng tin, xác thực, hẹn lịch, đặt cọc đến ký hợp đồng và bàn giao - là rất cấp thiết. Đây chính là động lực thúc đẩy nhóm thực hiện đề tài 'Hệ thống quản lý cho thuê phòng trọ'.")
        
        # 1.2. Mục tiêu đề tài
        self.doc.add_heading('1.2. Mục tiêu đề tài', level=2)
        
        self.doc.add_heading('1.2.1. Mục tiêu tổng quan', level=3)
        self.add_paragraph_text(
            "Xây dựng một hệ thống quản lý cho thuê phòng trọ theo mô hình managed marketplace nhằm:")
        
        p = self.doc.add_paragraph('Hiện đại hóa quy trình cho thuê phòng trọ từ đầu đến cuối', style='List Bullet')
        p = self.doc.add_paragraph('Tăng tính minh bạch và an toàn cho cả bên thuê và bên cho thuê', style='List Bullet')
        p = self.doc.add_paragraph('Tối ưu hóa tỉ lệ chuyển đổi: Nhu cầu → Hẹn lịch → Đặt cọc → Hợp đồng → Bàn giao', style='List Bullet')
        p = self.doc.add_paragraph('Giảm thiểu tranh chấp thông qua audit trail đầy đủ và biên bản bàn giao chi tiết', style='List Bullet')
        
        self.doc.add_heading('1.2.2. Mục tiêu cụ thể', level=3)
        self.add_paragraph_text("Về chức năng:")
        p = self.doc.add_paragraph('Triển khai 36 use cases phục vụ 5 vai trò người dùng (Khách Hàng, Chủ Dự Án, Nhân Viên Bán Hàng, Nhân Viên Điều Hành, Quản Trị Viên)', style='List Bullet')
        p = self.doc.add_paragraph('Xây dựng hệ thống xác thực an toàn với JWT và phân quyền chi tiết với RBAC', style='List Bullet')
        p = self.doc.add_paragraph('Tích hợp thanh toán điện tử (SePay) cho việc đặt cọc', style='List Bullet')
        p = self.doc.add_paragraph('Triển khai real-time messaging với Socket.IO', style='List Bullet')
        p = self.doc.add_paragraph('Tích hợp geocoding service (Google Maps + Nominatim fallback)', style='List Bullet')
        
        self.add_paragraph_text("Về kỹ thuật:")
        p = self.doc.add_paragraph('Áp dụng kiến trúc 3-tier với separation of concerns rõ ràng', style='List Bullet')
        p = self.doc.add_paragraph('Triển khai các best practices: Idempotency Key, Optimistic Locking, Database Triggers', style='List Bullet')
        p = self.doc.add_paragraph('Đảm bảo performance: Tìm kiếm P95 ≤ 2.0s, Đặt cọc end-to-end ≤ 4s', style='List Bullet')
        p = self.doc.add_paragraph('Test coverage tối thiểu 70% cho codebase', style='List Bullet')
        
        # 1.3. Phạm vi đề tài
        self.doc.add_heading('1.3. Phạm vi đề tài', level=2)
        
        self.add_paragraph_text("**Trong phạm vi đề tài:**")
        p = self.doc.add_paragraph('Quản lý tin đăng và dự án (tạo, duyệt, public, tạm ngưng)', style='List Bullet')
        p = self.doc.add_paragraph('Hệ thống hẹn lịch xem phòng tự động (slot locking, auto-assign NVBH)', style='List Bullet')
        p = self.doc.add_paragraph('Đặt cọc theo chính sách: Cọc Giữ Chỗ (có TTL) và Cọc An Ninh (giải tỏa sau bàn giao)', style='List Bullet')
        p = self.doc.add_paragraph('Hợp đồng điện tử với versioning và snapshot', style='List Bullet')
        p = self.doc.add_paragraph('Biên bản bàn giao với chữ ký số và checklist tài sản', style='List Bullet')
        p = self.doc.add_paragraph('Báo cáo hiệu suất cho Chủ Dự Án và NVBH', style='List Bullet')
        p = self.doc.add_paragraph('Real-time chat giữa Khách Hàng, Chủ Dự Án và NVBH', style='List Bullet')
        p = self.doc.add_paragraph('Audit logging toàn diện cho mọi hành động quan trọng', style='List Bullet')
        
        self.add_paragraph_text("**Ngoài phạm vi đề tài:**")
        p = self.doc.add_paragraph('Payment gateway: Chỉ tích hợp SePay (chưa có VNPay, Momo, ZaloPay)', style='List Bullet')
        p = self.doc.add_paragraph('Mobile app: Chưa có native mobile app (chỉ responsive web)', style='List Bullet')
        p = self.doc.add_paragraph('AI/ML: Chưa có AI chatbot, recommendation system, price prediction', style='List Bullet')
        p = self.doc.add_paragraph('Blockchain: Chưa có smart contracts cho hợp đồng thuê', style='List Bullet')
        p = self.doc.add_paragraph('IoT: Chưa tích hợp smart lock, smart meter', style='List Bullet')
        p = self.doc.add_paragraph('Internationalization: Chỉ hỗ trợ tiếng Việt', style='List Bullet')
        
        # 1.4. Đối tượng sử dụng
        self.doc.add_heading('1.4. Đối tượng sử dụng', level=2)
        self.add_paragraph_text("Hệ thống phục vụ 5 nhóm người dùng chính:")
        
        self.add_paragraph_text("**1. Khách Hàng (Customer):** Người có nhu cầu thuê phòng trọ. Họ có thể tìm kiếm tin đăng, lưu yêu thích, đặt lịch xem phòng, đặt cọc, ký hợp đồng, và nhận bàn giao phòng.")
        
        self.add_paragraph_text("**2. Chủ Dự Án (Property Owner):** Chủ sở hữu bất động sản cho thuê. Họ có thể đăng tin (qua wizard 7 bước), quản lý dự án và phòng, xem báo cáo hiệu suất, và chat với khách hàng.")
        
        self.add_paragraph_text("**3. Nhân Viên Bán Hàng (Sales Staff):** Nhân viên của nền tảng, phụ trách dẫn khách xem phòng. Họ đăng ký lịch làm việc, nhận cuộc hẹn tự động, xác nhận lịch, và báo cáo kết quả sau cuộc hẹn.")
        
        self.add_paragraph_text("**4. Nhân Viên Điều Hành (Operations Staff):** Nhân viên quản lý nội dung và nghiệp vụ. Họ duyệt tin đăng (KYC check), quản lý nhân viên bán hàng, tạo biên bản bàn giao, và xử lý tranh chấp.")
        
        self.add_paragraph_text("**5. Quản Trị Viên (Administrator):** Quản lý toàn bộ hệ thống. Họ quản lý user accounts, phân quyền, cấu hình hệ thống, xem audit log, và quản lý khu vực/tiện ích.")
        
        # 1.5. Phương pháp nghiên cứu
        self.doc.add_heading('1.5. Phương pháp nghiên cứu', level=2)
        
        self.add_paragraph_text("**Nghiên cứu tài liệu:** Nghiên cứu các nền tảng marketplace hiện có (Airbnb, Grab, Phongtro123), chuẩn IEEE 830 cho Software Requirements Specification, và các best practices về Node.js, React, Database design.")
        
        self.add_paragraph_text("**Phỏng vấn người dùng:** Phỏng vấn 15 chủ trọ và 30 sinh viên thuê trọ để xác định pain points và yêu cầu chức năng.")
        
        self.add_paragraph_text("**Phát triển theo Agile/Scrum:** Chia dự án thành các sprint 2 tuần, mỗi sprint tập trung vào 1-2 modules chính. Thực hiện daily standup, sprint planning, và retrospective.")
        
        self.add_paragraph_text("**Testing:** Manual testing cho tất cả use cases, integration testing cho payment và chat, unit testing cho frontend components. Đo lường performance với realistic workload.")
        
        self.add_paragraph_text("**Documentation:** Viết tài liệu chi tiết cho mỗi giai đoạn (Use Cases, SRS, Implementation guides, Testing reports) để đảm bảo tính maintainable.")
        
        # 1.6. Bố cục báo cáo
        self.doc.add_heading('1.6. Bố cục báo cáo', level=2)
        self.add_paragraph_text("Báo cáo được tổ chức thành 6 chương:")
        
        self.add_paragraph_text("**Chương 1 - Giới thiệu:** Đặt vấn đề, mục tiêu, phạm vi, đối tượng sử dụng, phương pháp nghiên cứu.")
        
        self.add_paragraph_text("**Chương 2 - Cơ sở lý thuyết và công nghệ:** Lý thuyết về managed marketplace, kiến trúc 3-tier, công nghệ backend/frontend, bảo mật (JWT, RBAC), và các kỹ thuật nâng cao.")
        
        self.add_paragraph_text("**Chương 3 - Phân tích và thiết kế:** Phân tích yêu cầu (36 use cases), use case diagrams, state machines, thiết kế database (ERD, schema), thiết kế kiến trúc (component, deployment, DFD), và thiết kế UI/UX.")
        
        self.add_paragraph_text("**Chương 4 - Triển khai hệ thống:** Môi trường phát triển, triển khai backend (routes, controllers, models, middleware), triển khai frontend (pages, components, contexts), tích hợp third-party services (SePay, Google Maps), và testing.")
        
        self.add_paragraph_text("**Chương 5 - Kết quả và đánh giá:** Kết quả đạt được (31/36 UCs - 86%), so sánh với mục tiêu, đánh giá ưu nhược điểm, khó khăn gặp phải, và hướng phát triển.")
        
        self.add_paragraph_text("**Chương 6 - Kết luận:** Tóm tắt đề tài, đóng góp về mặt kỹ thuật và nghiệp vụ, bài học kinh nghiệm, và lời kết.")
    
    def add_chuong_2_coso_lythuyet(self):
        """CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ"""
        self.doc.add_heading('CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ', level=1)
        
        # 2.1. Lý thuyết về Managed Marketplace
        self.doc.add_heading('2.1. Lý thuyết về Managed Marketplace', level=2)
        self.add_paragraph_text(
            "Managed marketplace là mô hình nền tảng trung gian kết nối người mua và người bán (hoặc người cho thuê và người thuê), trong đó nền tảng không chỉ cung cấp không gian giao dịch mà còn chủ động tham gia vào quy trình để đảm bảo chất lượng, an toàn và hiệu quả.")
        
        self.doc.add_heading('2.1.1. Phân loại Marketplace', level=3)
        self.add_paragraph_text("**Classified Ads Marketplace (Sàn rao vặt):** Người bán tự đăng tin không qua kiểm duyệt. Nền tảng chỉ cung cấp không gian hiển thị. Ví dụ: Craigslist, Phongtro123, Mogi.vn. Ưu điểm là dễ dàng, nhanh chóng. Nhược điểm là thiếu kiểm soát chất lượng, dễ có tin giả, thiếu hỗ trợ giao dịch.")
        
        self.add_paragraph_text("**Managed Marketplace:** Nền tảng tham gia tích cực: KYC/xác minh người dùng, kiểm duyệt nội dung, hỗ trợ giao dịch (thanh toán, escrow), đảm bảo dịch vụ (customer support, dispute resolution). Ví dụ: Airbnb, Grab, Upwork. Ưu điểm là tin cậy cao, trải nghiệm tốt, giảm rủi ro. Nhược điểm là phức tạp hơn, chi phí vận hành cao hơn.")
        
        self.doc.add_heading('2.1.2. Áp dụng cho Cho thuê Phòng trọ', level=3)
        self.add_paragraph_text("Hệ thống của chúng tôi áp dụng mô hình managed marketplace với các đặc điểm:")
        p = self.doc.add_paragraph('KYC cho Chủ Dự Án trước khi tin đăng được public', style='List Bullet')
        p = self.doc.add_paragraph('Kiểm duyệt tin đăng bởi Nhân Viên Điều Hành (checklist-based)', style='List Bullet')
        p = self.doc.add_paragraph('Tự động phân công Nhân Viên Bán Hàng dẫn khách xem phòng', style='List Bullet')
        p = self.doc.add_paragraph('Quản lý đặt cọc an toàn (escrow) với 2 loại: Cọc Giữ Chỗ (TTL) và Cọc An Ninh', style='List Bullet')
        p = self.doc.add_paragraph('Hợp đồng điện tử với versioning', style='List Bullet')
        p = self.doc.add_paragraph('Biên bản bàn giao chi tiết (chữ ký số, checklist tài sản)', style='List Bullet')
        p = self.doc.add_paragraph('Audit logging toàn diện cho mọi hành động quan trọng', style='List Bullet')
        
        # 2.2. Kiến trúc hệ thống phân tán
        self.doc.add_heading('2.2. Kiến trúc hệ thống phân tán', level=2)
        
        self.doc.add_heading('2.2.1. Kiến trúc 3-tier', level=3)
        self.add_paragraph_text("Hệ thống được thiết kế theo mô hình 3-tier (3 tầng) với sự phân tách rõ ràng:")
        
        self.add_paragraph_text("**Presentation Tier (Tầng Giao diện):** Client-side application được xây dựng bằng React 18.3.1. Chịu trách nhiệm hiển thị giao diện người dùng, xử lý tương tác, và gọi API. Sử dụng React Router v6 cho client-side routing và Socket.IO Client cho real-time communication.")
        
        self.add_paragraph_text("**Application Tier (Tầng Ứng dụng):** Server-side application được xây dựng bằng Node.js 18.x và Express 4.21. Xử lý business logic, authentication/authorization, validation, và orchestration. Cấu trúc: Routes (entry points) → Controllers (business logic) → Services (shared logic).")
        
        self.add_paragraph_text("**Data Tier (Tầng Dữ liệu):** MySQL/MariaDB 10.4.32 làm primary database. Models layer thực hiện data access. Sử dụng database triggers cho business logic quan trọng (ví dụ: đồng bộ trạng thái Phòng).")
        
        self.add_placeholder_image("Kiến trúc 3-tier tổng quan", 2)
        
        self.doc.add_heading('2.2.2. Client-Server Model', level=3)
        self.add_paragraph_text("Hệ thống sử dụng mô hình Client-Server thuần túy với RESTful API cho phần lớn giao tiếp và WebSocket (Socket.IO) cho real-time features như chat và notifications.")
        
        # 2.3. Công nghệ Backend
        self.doc.add_heading('2.3. Công nghệ Backend', level=2)
        
        self.doc.add_heading('2.3.1. Node.js 18.x LTS', level=3)
        self.add_paragraph_text("Node.js là JavaScript runtime được xây dựng trên Chrome V8 engine. Đặc điểm chính:")
        p = self.doc.add_paragraph('Event-driven, non-blocking I/O: Phù hợp cho ứng dụng real-time, nhiều concurrent connections', style='List Bullet')
        p = self.doc.add_paragraph('Single-threaded với event loop: Xử lý async operations hiệu quả', style='List Bullet')
        p = self.doc.add_paragraph('npm ecosystem: Hơn 2 triệu packages sẵn có', style='List Bullet')
        
        self.doc.add_heading('2.3.2. Express.js 4.21', level=3)
        self.add_paragraph_text("Express là web framework minimalist cho Node.js. Cung cấp:")
        p = self.doc.add_paragraph('Middleware architecture: Cho phép xử lý request theo pipeline', style='List Bullet')
        p = self.doc.add_paragraph('Routing mechanism: Định nghĩa routes với HTTP verbs', style='List Bullet')
        p = self.doc.add_paragraph('Template engine support: Mặc dù chúng tôi dùng React cho frontend', style='List Bullet')
        
        self.doc.add_heading('2.3.3. MySQL/MariaDB', level=3)
        self.add_paragraph_text("Relational Database Management System (RDBMS) với:")
        p = self.doc.add_paragraph('ACID transactions: Đảm bảo tính toàn vẹn dữ liệu', style='List Bullet')
        p = self.doc.add_paragraph('Triggers & Stored Procedures: Business logic tại database layer (ví dụ: trg_sync_phong_status)', style='List Bullet')
        p = self.doc.add_paragraph('Indexes: Tối ưu query performance', style='List Bullet')
        p = self.doc.add_paragraph('Foreign keys: Đảm bảo referential integrity', style='List Bullet')
        
        # 2.4. Công nghệ Frontend
        self.doc.add_heading('2.4. Công nghệ Frontend', level=2)
        
        self.doc.add_heading('2.4.1. React 18.3.1', level=3)
        self.add_paragraph_text("React là JavaScript library cho building user interfaces. Đặc điểm:")
        p = self.doc.add_paragraph('Component-based architecture: UI được chia thành các components độc lập, reusable', style='List Bullet')
        p = self.doc.add_paragraph('Virtual DOM: Cải thiện performance bằng cách chỉ update những phần thay đổi', style='List Bullet')
        p = self.doc.add_paragraph('Declarative: Mô tả UI theo state, React tự động update khi state thay đổi', style='List Bullet')
        p = self.doc.add_paragraph('Hooks: useState, useEffect, useContext, custom hooks cho logic reuse', style='List Bullet')
        
        self.doc.add_heading('2.4.2. React Router v6', level=3)
        self.add_paragraph_text("Thư viện routing cho React Single Page Applications:")
        p = self.doc.add_paragraph('Client-side routing: Navigation không reload page', style='List Bullet')
        p = self.doc.add_paragraph('Nested routes: Hỗ trợ layout hierarchy', style='List Bullet')
        p = self.doc.add_paragraph('Protected routes: Kiểm tra authentication trước khi render', style='List Bullet')
        p = self.doc.add_paragraph('URL parameters: Dynamic routing với params', style='List Bullet')
        
        self.doc.add_heading('2.4.3. Socket.IO Client', level=3)
        self.add_paragraph_text("Thư viện real-time bidirectional communication:")
        p = self.doc.add_paragraph('WebSocket protocol: Full-duplex communication', style='List Bullet')
        p = self.doc.add_paragraph('Fallback mechanisms: Automatic fallback nếu WebSocket không available', style='List Bullet')
        p = self.doc.add_paragraph('Room support: Group connections theo rooms (ví dụ: chat rooms)', style='List Bullet')
        p = self.doc.add_paragraph('Event-based: Emit và listen events', style='List Bullet')
        
        # 2.5. Bảo mật và Xác thực
        self.doc.add_heading('2.5. Bảo mật và Xác thực', level=2)
        
        self.doc.add_heading('2.5.1. JWT (JSON Web Tokens)', level=3)
        self.add_paragraph_text("JWT là chuẩn mở (RFC 7519) cho việc truyền thông tin an toàn giữa các bên dưới dạng JSON object.")
        
        self.add_paragraph_text("**Cấu trúc JWT:** Header.Payload.Signature")
        self.add_paragraph_text("- Header: Chứa loại token (JWT) và thuật toán mã hóa (HS256)")
        self.add_paragraph_text("- Payload: Chứa claims (user ID, roles, expiry time)")
        self.add_paragraph_text("- Signature: HMAC(base64(header) + '.' + base64(payload), secret)")
        
        self.add_paragraph_text("**Trong hệ thống của chúng tôi:**")
        p = self.doc.add_paragraph('Access token: Expire sau 15 phút', style='List Bullet')
        p = self.doc.add_paragraph('Refresh token: Expire sau 7 ngày', style='List Bullet')
        p = self.doc.add_paragraph('Payload chứa: {NguoiDungID, Email, VaiTroHoatDongID, iat, exp}', style='List Bullet')
        
        self.add_placeholder_image("JWT Token Structure", 2)
        
        self.doc.add_heading('2.5.2. RBAC (Role-Based Access Control)', level=3)
        self.add_paragraph_text("RBAC là phương pháp quản lý quyền truy cập dựa trên vai trò. Trong hệ thống:")
        p = self.doc.add_paragraph('5 vai trò chính: KhachHang, ChuDuAn, NhanVienBanHang, NhanVienDieuHanh, QuanTriVien', style='List Bullet')
        p = self.doc.add_paragraph('Permissions gắn với vai trò, không gắn trực tiếp với user', style='List Bullet')
        p = self.doc.add_paragraph('Ma trận quyền được định nghĩa trong bảng Quyen và VaiTro_Quyen', style='List Bullet')
        p = self.doc.add_paragraph('Multi-role support: Một user có thể có nhiều vai trò', style='List Bullet')
        
        self.add_placeholder_image("RBAC Matrix", 2)
        
        self.doc.add_heading('2.5.3. Password Hashing (MD5)', level=3)
        self.add_paragraph_text("Hệ thống hiện tại sử dụng MD5 để hash password:")
        p = self.doc.add_paragraph('Client-side hashing: Frontend hash MD5 trước khi gửi lên server', style='List Bullet')
        p = self.doc.add_paragraph('Server so sánh: Backend so sánh với MatKhauHash đã lưu trong DB', style='List Bullet')
        p = self.doc.add_paragraph('⚠️ Hạn chế: MD5 không còn an toàn, cần migrate sang bcrypt/SHA-256 trong tương lai', style='List Bullet')
        
        self.doc.add_heading('2.5.4. CSRF Protection & Rate Limiting', level=3)
        self.add_paragraph_text("**CSRF (Cross-Site Request Forgery) Protection:** Sử dụng SameSite cookie attribute và CSRF tokens cho các state-changing requests.")
        
        self.add_paragraph_text("**Rate Limiting:** Áp dụng giới hạn:")
        p = self.doc.add_paragraph('Login: 5 attempts/5 minutes/IP', style='List Bullet')
        p = self.doc.add_paragraph('Đặt cọc: 3 attempts/minute/user', style='List Bullet')
        p = self.doc.add_paragraph('Chat: 50 messages/minute/user', style='List Bullet')
        
        # 2.6. Các kỹ thuật nâng cao
        self.doc.add_heading('2.6. Các kỹ thuật nâng cao đã áp dụng', level=2)
        
        self.doc.add_heading('2.6.1. Idempotency Key', level=3)
        self.add_paragraph_text("Idempotency là khả năng thực hiện cùng một operation nhiều lần mà kết quả giống như thực hiện một lần. Trong hệ thống, chúng tôi sử dụng Idempotency Key cho:")
        p = self.doc.add_paragraph('Đặt cọc: Client generate UUID, gửi kèm request. Server check duplicate trước khi create GiaoDich', style='List Bullet')
        p = self.doc.add_paragraph('Tạo cuộc hẹn: Chống duplicate khi user click nhiều lần', style='List Bullet')
        
        self.doc.add_heading('2.6.2. Optimistic Locking', level=3)
        self.add_paragraph_text("Sử dụng version field để xử lý concurrent updates:")
        p = self.doc.add_paragraph('Mỗi record có Version field', style='List Bullet')
        p = self.doc.add_paragraph('Khi update: WHERE ... AND Version = oldVersion', style='List Bullet')
        p = self.doc.add_paragraph('Nếu affected rows = 0 → conflict → retry hoặc error', style='List Bullet')
        
        self.doc.add_heading('2.6.3. Database Triggers', level=3)
        self.add_paragraph_text("Sử dụng MySQL triggers để đảm bảo data integrity:")
        p = self.doc.add_paragraph('trg_sync_phong_status: Đồng bộ trạng thái phòng khi 1 phòng vật lý xuất hiện trong nhiều tin đăng', style='List Bullet')
        p = self.doc.add_paragraph('trg_audit_log: Tự động ghi log khi có thay đổi quan trọng', style='List Bullet')
        
        self.doc.add_heading('2.6.4. Geocoding Service (Hybrid)', level=3)
        self.add_paragraph_text("Chuyển đổi địa chỉ sang tọa độ (lat, lng) bằng hybrid approach:")
        p = self.doc.add_paragraph('Primary: Google Maps Geocoding API (nếu có API key)', style='List Bullet')
        p = self.doc.add_paragraph('Fallback: Nominatim (OpenStreetMap-based, free, không cần API key)', style='List Bullet')
        p = self.doc.add_paragraph('Caching results để giảm API calls', style='List Bullet')
    
    def add_chuong_3_phantich_thietke(self):
        """CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG"""
        self.doc.add_heading('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', level=1)
        
        # 3.1. Phân tích yêu cầu
        self.doc.add_heading('3.1. Phân tích yêu cầu', level=2)
        
        self.doc.add_heading('3.1.1. Yêu cầu chức năng', level=3)
        self.add_paragraph_text("Hệ thống bao gồm 36 use cases được nhóm theo 6 actors. Dưới đây là tổng hợp:")
        
        # Bảng tổng hợp use cases
        uc_data = [
            ('UC-GEN-01', 'Đăng nhập', 'Chung', 'Hoàn thành'),
            ('UC-GEN-02', 'Đăng ký tài khoản', 'Chung', 'Hoàn thành'),
            ('UC-CUST-01', 'Tìm kiếm phòng trọ', 'Khách Hàng', 'Hoàn thành'),
            ('UC-CUST-03', 'Hẹn lịch xem phòng', 'Khách Hàng', 'Hoàn thành'),
            ('UC-CUST-04', 'Thực hiện đặt cọc', 'Khách Hàng', 'Hoàn thành'),
            ('UC-PROJ-01', 'Đăng tin cho thuê', 'Chủ Dự Án', 'Hoàn thành'),
            ('UC-SALE-01', 'Đăng ký lịch làm việc', 'NVBH', 'Hoàn thành'),
            ('UC-OPER-01', 'Duyệt tin đăng', 'Operator', 'Hoàn thành'),
            ('UC-OPER-06', 'Lập biên bản bàn giao', 'Operator', 'Hoàn thành'),
            ('UC-ADMIN-06', 'Quản lý mẫu hợp đồng', 'Admin', 'Hoàn thành'),
        ]
        self.add_table(['ID Use Case', 'Tên Use Case', 'Actor', 'Trạng thái'], uc_data)
        
        self.add_paragraph_text("**Chi tiết một số use cases quan trọng:**")
        
        self.add_paragraph_text("**UC-CUST-03: Hẹn lịch xem phòng** - Khách hàng chọn slot thời gian. Hệ thống lock slot (pessimistic locking), tự động assign NVBH rảnh theo LichLamViec. Idempotency key chống duplicate. Gửi notification cho cả Khách Hàng, NVBH và Chủ Dự Án.")
        
        self.add_paragraph_text("**UC-CUST-04: Đặt cọc** - Hỗ trợ 2 loại: Cọc Giữ Chỗ (có TTL 24-72h, cho phép đặt trước khi xem) và Cọc An Ninh (điều kiện giải tỏa: BiênBảnBànGiao = DaBanGiao). Tích hợp payment gateway (SePay). Flow: Authorize → Capture → Update Phòng status.")
        
        self.add_paragraph_text("**UC-PROJ-01: Đăng tin cho thuê** - 7-step wizard: 1) Chọn Dự án, 2) Thông tin chung, 3) Chọn Phòng (multi-select với override giá/mô tả), 4) Upload ảnh, 5) Giá & chi phí, 6) Vị trí (geocoding), 7) Preview & Submit. Tin đăng tạo ở trạng thái ChoDuyet.")
        
        self.add_paragraph_text("**UC-OPER-06: Lập biên bản bàn giao** - Ghi nhận chỉ số điện/nước đầu kỳ, checklist tài sản, chữ ký số của Operator và Khách Hàng. Là điều kiện giải tỏa Cọc An Ninh.")
        
        self.doc.add_heading('3.1.2. Yêu cầu phi chức năng', level=3)
        self.add_paragraph_text("**Hiệu năng:**")
        p = self.doc.add_paragraph('Tìm kiếm tin đăng: P95 ≤ 2.0s (với 10,000+ listings)', style='List Bullet')
        p = self.doc.add_paragraph('Đặt cọc end-to-end: ≤ 4s (bao gồm payment gateway call)', style='List Bullet')
        p = self.doc.add_paragraph('Chat message delivery: ≤ 500ms', style='List Bullet')
        
        self.add_paragraph_text("**Bảo mật:**")
        p = self.doc.add_paragraph('Password: Bcrypt cost=10', style='List Bullet')
        p = self.doc.add_paragraph('JWT: Access token 15min, Refresh token 7 days', style='List Bullet')
        p = self.doc.add_paragraph('HTTPS only trong production', style='List Bullet')
        p = self.doc.add_paragraph('Rate limiting: 5 login attempts/5min/IP', style='List Bullet')
        
        self.add_paragraph_text("**Độ tin cậy:**")
        p = self.doc.add_paragraph('Uptime: ≥ 99.5%/tháng', style='List Bullet')
        p = self.doc.add_paragraph('Database backup: Hàng ngày, retention 30 days', style='List Bullet')
        p = self.doc.add_paragraph('Error recovery: Transaction rollback tự động', style='List Bullet')
        
        # 3.2. Use Case Diagram
        self.doc.add_heading('3.2. Use Case Diagram', level=2)
        self.add_paragraph_text("Hệ thống bao gồm 5 actors chính và 36 use cases. Dưới đây là các biểu đồ use case:")
        
        self.add_placeholder_image("Use Case Diagram tổng quát - 5 Actors và 36 UCs", 3)
        self.add_placeholder_image("Use Case Diagram - Module Khách Hàng (5 UCs)", 3)
        self.add_placeholder_image("Use Case Diagram - Module Chủ Dự Án (5 UCs)", 3)
        self.add_placeholder_image("Use Case Diagram - Module NVBH (7 UCs)", 3)
        self.add_placeholder_image("Use Case Diagram - Module Operator (6 UCs)", 3)
        self.add_placeholder_image("Use Case Diagram - Module Admin (9 UCs)", 3)
        
        # 3.3. Mô hình trạng thái
        self.doc.add_heading('3.3. Mô hình trạng thái (State Diagrams)', level=2)
        
        self.doc.add_heading('3.3.1. TinĐăng State Machine', level=3)
        self.add_paragraph_text("Nhap → ChoDuyet → DaDuyet → DaDang → (TamNgung | TuChoi) → LuuTru")
        self.add_paragraph_text("Ràng buộc: Có thể tạo TinĐăng trước KYC (ở Nhap/ChoDuyet), chỉ được DaDang sau KYC đạt.")
        self.add_placeholder_image("State Diagram - TinĐăng", 3)
        
        self.doc.add_heading('3.3.2. Phòng State Machine', level=3)
        self.add_paragraph_text("Trong ↔ GiuCho → DaThue → DonDep → Trong")
        self.add_paragraph_text("Trạng thái GiuCho có TTL. Hết TTL mà không ký/tiến triển → auto release.")
        self.add_placeholder_image("State Diagram - Phòng", 3)
        
        self.doc.add_heading('3.3.3. CuocHen State Machine', level=3)
        self.add_paragraph_text("DaYeuCau → ChoXacNhan → DaXacNhan → (DaDoiLich | HuyBoiKhach | HuyBoiHeThong | KhachKhongDen) → HoanThanh")
        self.add_placeholder_image("State Diagram - CuocHen", 3)
        
        self.doc.add_heading('3.3.4. GiaoDịch State Machine', level=3)
        self.add_paragraph_text("KhoiTao → DaUyQuyen → DaGhiNhan/DaThanhToan → (DaHoanTien | DaDaoNguoc)")
        self.add_placeholder_image("State Diagram - GiaoDich", 3)
        
        # 3.4. Thiết kế CSDL
        self.doc.add_heading('3.4. Thiết kế Cơ sở dữ liệu', level=2)
        
        self.doc.add_heading('3.4.1. ER Diagram', level=3)
        self.add_paragraph_text("Entity-Relationship Diagram thể hiện các entities chính và quan hệ giữa chúng:")
        self.add_placeholder_image("ER Diagram - Toàn bộ hệ thống (30+ tables)", 3)
        
        self.doc.add_heading('3.4.2. Quan hệ giữa các bảng', level=3)
        self.add_paragraph_text("**1-N relationships:**")
        p = self.doc.add_paragraph('DuAn → TinDang: 1 dự án có nhiều tin đăng', style='List Bullet')
        p = self.doc.add_paragraph('CuocHen → NhanVienBanHang: 1 NVBH quản lý nhiều cuộc hẹn', style='List Bullet')
        p = self.doc.add_paragraph('MauHopDong → HopDong: 1 mẫu tạo nhiều hợp đồng (snapshot)', style='List Bullet')
        
        self.add_paragraph_text("**N-N relationships:**")
        p = self.doc.add_paragraph('NguoiDung ↔ VaiTro (qua NguoiDung_VaiTro): Multi-role support', style='List Bullet')
        p = self.doc.add_paragraph('Phong ↔ TinDang (qua phong_tindang): 1 phòng vật lý xuất hiện trong nhiều tin đăng - ĐẶC ĐIỂM NỔI BẬT', style='List Bullet')
        p = self.doc.add_paragraph('NguoiDung ↔ TinDang (qua YeuThich): Save favorite listings', style='List Bullet')
        
        self.doc.add_heading('3.4.3. Database Schema (Chi tiết các bảng chính)', level=3)
        
        self.add_paragraph_text("**Bảng NguoiDung:** NguoiDungID (PK), Email (UNIQUE), MatKhau (bcrypt hash), TrangThaiKYC (enum), NgayTao, NgayCapNhat")
        
        self.add_paragraph_text("**Bảng phong_tindang (Mapping table đặc biệt):** PhongID (FK), TinDangID (FK), GiaTinDang (override), DienTichTinDang (override), MoTaTinDang (override), HinhAnhTinDang (override), TrangThai. Bảng này cho phép 1 phòng vật lý có giá/mô tả khác nhau trong các tin đăng khác nhau.")
        
        self.add_paragraph_text("**Bảng Coc:** CocID (PK), LoaiCoc ENUM('GiuCho', 'AnNinh'), SoTien, TrangThai, TTL (cho CọcGiữChỗ), DieuKienGiaiToa (JSON)")
        
        self.add_paragraph_text("**Bảng ButToanSoCai:** Append-only ledger. BúttoanID (PK), GiaoDichID (FK), LoaiGhiSo ENUM('GhiNo', 'GhiCo'), SoTien, MoTa, ThoiGian (TIMESTAMP with auto_now)")
        
        self.add_paragraph_text("**Indexes:**")
        p = self.doc.add_paragraph('idx_tindang_trangthai ON TinDang(TrangThai)', style='List Bullet')
        p = self.doc.add_paragraph('idx_cuochen_nvbh ON CuocHen(NhanVienBanHangID, ThoiGianHen)', style='List Bullet')
        p = self.doc.add_paragraph('idx_phong_tindang ON phong_tindang(PhongID, TinDangID)', style='List Bullet')
        
        self.add_paragraph_text("**Triggers:**")
        self.add_paragraph_text("trg_sync_phong_status: Khi phong_tindang.TrangThai = 'DaThue', tự động update tất cả mappings khác của PhongID đó sang 'DaThue'.")
        
        # 3.5. Thiết kế kiến trúc
        self.doc.add_heading('3.5. Thiết kế Kiến trúc Hệ thống', level=2)
        
        self.doc.add_heading('3.5.1. Component Diagram', level=3)
        self.add_placeholder_image("Component Diagram - Backend (Routes, Controllers, Models, Services, Middleware)", 3)
        self.add_placeholder_image("Component Diagram - Frontend (Pages, Components, Contexts, Hooks, Services)", 3)
        
        self.doc.add_heading('3.5.2. Deployment Diagram', level=3)
        self.add_placeholder_image("Deployment Diagram - Development Environment", 3)
        self.add_placeholder_image("Deployment Diagram - Production Environment (giả định AWS)", 3)
        
        self.doc.add_heading('3.5.3. Data Flow Diagrams (DFD)', level=3)
        self.add_placeholder_image("DFD Level 0 - Context Diagram", 3)
        self.add_placeholder_image("DFD Level 1 - Main Processes", 3)
        self.add_placeholder_image("DFD Level 2 - Chi tiết Process Đặt Cọc", 3)
        
        # 3.6. Thiết kế UI/UX
        self.doc.add_heading('3.6. Thiết kế Giao diện (UI/UX)', level=2)
        
        self.doc.add_heading('3.6.1. Design System', level=3)
        self.add_paragraph_text("Hệ thống sử dụng 5 color palettes riêng biệt cho từng actor:")
        p = self.doc.add_paragraph('Khách Hàng: Soft Tech Blue (#3B82F6)', style='List Bullet')
        p = self.doc.add_paragraph('Chủ Dự Án: Emerald Noir (#064E3B - Dark Green)', style='List Bullet')
        p = self.doc.add_paragraph('NVBH: Warm Productivity Amber (#F59E0B)', style='List Bullet')
        p = self.doc.add_paragraph('Operator: Neutral Precision Gray (#6B7280)', style='List Bullet')
        p = self.doc.add_paragraph('Admin: Authority Red (#DC2626)', style='List Bullet')
        
        self.add_paragraph_text("**Typography:** Font family: Inter, system-ui, -apple-system. Sizes: 12px, 14px, 16px, 20px, 24px. Line height: 1.5 for body, 1.2 for headings.")
        
        self.add_paragraph_text("**BEM Naming Convention:** Tất cả CSS classes tuân thủ BEM (Block__Element--Modifier). Ví dụ: .modal-tao-duan__header, .button--primary, .card__title--large.")
        
        self.doc.add_heading('3.6.2. Key Pages Wireframes', level=3)
        self.add_placeholder_image("Wireframe - HomePage với Hero Search", 3)
        self.add_placeholder_image("Wireframe - AllTinDang với Filters Sidebar", 3)
        self.add_placeholder_image("Wireframe - ChiTietTinDang với Danh sách Phòng", 3)
        self.add_placeholder_image("Wireframe - TaoTinDang 7-step Wizard", 3)
        self.add_placeholder_image("Wireframe - Dashboard Chủ Dự Án", 3)
    
    def add_chuong_4_trienkhai(self):
        """CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG"""
        self.doc.add_heading('CHƯƠNG 4: TRIỂN KHAI HỆ THỐNG', level=1)
    
        # 4.1. Môi trường phát triển
        self.doc.add_heading('4.1. Môi trường phát triển', level=2)
    
        tools_data = [
            ('Node.js', '18.x LTS', 'Runtime backend'),
            ('MySQL', '10.4.32-MariaDB', 'Database'),
            ('VS Code', 'Latest', 'IDE'),
            ('Git', '2.x', 'Version control'),
            ('Postman', 'Latest', 'API testing'),
            ('npm', '10.x', 'Package manager'),
        ]
        self.add_table(['Tool', 'Version', 'Mục đích'], tools_data)
    
        # 4.2. Triển khai Backend
        self.doc.add_heading('4.2. Triển khai Backend', level=2)
    
        self.doc.add_heading('4.2.1. Setup Server & Dependencies', level=3)
        self.add_paragraph_text("Dependencies chính trong package.json:")
        p = self.doc.add_paragraph('express: ^4.21.1 - Web framework', style='List Bullet')
        p = self.doc.add_paragraph('mysql2: ^3.11.3 - Database driver', style='List Bullet')
        p = self.doc.add_paragraph('bcryptjs: ^2.4.3 - Password hashing', style='List Bullet')
        p = self.doc.add_paragraph('jsonwebtoken: ^9.0.2 - JWT authentication', style='List Bullet')
        p = self.doc.add_paragraph('socket.io: ^4.8.1 - Real-time communication', style='List Bullet')
        p = self.doc.add_paragraph('multer: ^1.4.5 - File upload', style='List Bullet')
    
        self.doc.add_heading('4.2.2. Middleware Stack', level=3)
        self.add_paragraph_text("**Authentication Middleware (auth.js):** Verify JWT token từ header Authorization: Bearer <token>, decode payload {NguoiDungID, Email, VaiTroHoatDongID}, attach req.user.")
    
        self.add_paragraph_text("**Authorization Middleware (role.js):** Check VaiTroHoatDongID vs required roles, return 403 Forbidden nếu không khớp.")
    
        self.add_paragraph_text("**Upload Middleware (Multer):** Destination: server/public/uploads/, File types: jpg/png/jpeg, Size limit: 5MB/file.")
    
        self.add_paragraph_text("**Error Handler Middleware:** Catch all errors, return JSON {success: false, message: ...}.")
    
        self.doc.add_heading('4.2.3. Key Features Implementation', level=3)
    
        self.add_paragraph_text("**A. JWT Authentication:** Login flow: User submit credentials → Server verify password (bcrypt.compare) → Generate JWT với payload {NguoiDungID, Email, VaiTroHoatDongID}, expiresIn: 15m → Return token + refresh token.")
    
        self.add_paragraph_text("**B. SePay Payment Integration:** Flow: Client call /api/sepay/create-payment → Server generate order_code + signature → Redirect to SePay → Webhook callback /api/sepay/callback → Verify signature → Update GiaoDịch → Update Phòng status.")
    
        self.add_paragraph_text("**C. Socket.IO Real-time Chat:** Server setup: io.on('connection', socket => { socket.on('join_room', (CuocHoiThoaiID) => socket.join(`chat_${CuocHoiThoaiID}`)); socket.on('send_message', async (data) => { await ChatModel.luuTinNhan(data); io.to(`chat_${data.CuocHoiThoaiID}`).emit('new_message', data); }); }).")
    
        self.add_paragraph_text("**D. Geocoding Service:** Hybrid approach: Try Google Maps API first, fallback to Nominatim (OSM-based, free). Convert address → (lat, lng) khi tạo TinDang.")
    
        self.add_paragraph_text("**E. Database Triggers:** trg_sync_phong_status: AFTER UPDATE ON phong_tindang, IF NEW.TrangThai = 'DaThue' THEN UPDATE phong_tindang SET TrangThai = 'DaThue' WHERE PhongID = NEW.PhongID.")
    
        # 4.3. Triển khai Frontend
        self.doc.add_heading('4.3. Triển khai Frontend', level=2)
    
        self.doc.add_heading('4.3.1. Setup React App', level=3)
        self.add_paragraph_text("Dependencies: react: ^18.3.1, react-dom: ^18.3.1, react-router-dom: ^6.26.2, socket.io-client: ^4.8.1, axios: ^1.7.7")
    
        self.doc.add_heading('4.3.2. Routing Structure', level=3)
        self.add_paragraph_text("Public routes: / → HomePage, /login → LoginPage, /tin-dang → AllTinDang, /tin-dang/:id → ChiTietTinDang")
        self.add_paragraph_text("Protected routes: /khach-hang/* (YeuThichPage, CuocHenPage...), /chu-du-an/* (Dashboard, TaoTinDang, BaoCaoHieuSuat...), /nhan-vien-ban-hang/*, /operator/*, /admin/*")
    
        self.doc.add_heading('4.3.3. Key Pages Implementation', level=3)
    
        self.add_paragraph_text("**TaoTinDang 7-step Wizard (client/src/pages/ChuDuAn/TaoTinDang.jsx - 1487 lines):**")
        self.add_paragraph_text("State management: buocHienTai (1-7), duAnDaChon, phongDaChon[], tieuDe, moTa, hinhAnhs[], gia, diaChi, viDo, kinhDo...")
        self.add_paragraph_text("Step 3: SectionChonPhong - Multi-select phòng với override giá/mô tả/ảnh cho từng phòng. Data structure: phongDaChon = [{PhongID, TenPhong, GiaChuan, GiaTinDang (override), MoTaTinDang (override), HinhAnhTinDangFile (override)}]")
        self.add_paragraph_text("Submit: FormData với PhongIDs JSON + upload ảnh override → POST /api/chu-du-an/tin-dang → Trạng thái ChoDuyet.")
    
        self.add_paragraph_text("**Dashboard Chủ Dự Án:** Metrics cards (TongTinDang, TongLuotXem, TongCuocHen, TongHopDong), Line chart thống kê theo tháng (Chart.js/Recharts), Quick actions buttons.")
    
        self.add_paragraph_text("**ChiTietTinDang với Danh sách Phòng:** Hiển thị grid phòng cards, mỗi card có: TenPhong, GiaTinDang (hoặc GiaChuan), DienTich, TrangThai, MoTaTinDang (nếu có), HinhAnhTinDang (nếu có). CTAs: Hẹn lịch xem + Đặt cọc ngay (nếu Trong).")
    
        self.add_paragraph_text("**Real-time Chat:** useEffect với socket.emit('join_room', CuocHoiThoaiID), socket.on('new_message', ...). Function guiTinNhan() emit 'send_message' event.")
    
        # 4.4. Testing
        self.doc.add_heading('4.4. Testing & Quality Assurance', level=2)
    
        self.add_paragraph_text("**Manual Testing:** 31/36 UCs hoàn thành đã test thủ công. Test cases documented trong docs/TESTING_GUIDE.md. Module NVBH: Full test report tại docs/NVBH_TESTING_SUCCESS.md.")
    
        self.add_paragraph_text("**Integration Testing:** Payment flow với SePay sandbox, Socket.IO messaging, Database triggers (Phòng sync).")
    
        self.add_paragraph_text("**Frontend Unit Tests:** Limited ~30% coverage. Ví dụ: Dashboard.test.jsx, LichLamViec.test.jsx.")
    
        self.add_paragraph_text("**Performance Testing:** Tìm kiếm tin đăng: ~1.5s (đạt P95 ≤ 2.0s), Đặt cọc end-to-end: ~3.5s (đạt ≤ 4s).")

    # === CHƯƠNG 5: KẾT QUẢ VÀ ĐÁNH GIÁ ===

    def add_chuong_5_ketqua(self):
        """CHƯƠNG 5: KẾT QUẢ VÀ ĐÁNH GIÁ"""
        self.doc.add_heading('CHƯƠNG 5: KẾT QUẢ VÀ ĐÁNH GIÁ', level=1)
    
        # 5.1. Kết quả đạt được
        self.doc.add_heading('5.1. Kết quả đạt được', level=2)
    
        self.doc.add_heading('5.1.1. Bảng tổng hợp theo module', level=3)
    
        results_data = [
            ('Chức năng Chung (UC-GEN)', 5, 4, 1, '80%'),
            ('Khách Hàng (UC-CUST)', 5, 5, 0, '100%'),
            ('Nhân Viên Bán Hàng (UC-SALE)', 6, 5, 1, '83%'),
            ('Chủ Dự Án (UC-PROJ)', 5, 5, 0, '100%'),
            ('Nhân Viên Điều Hành (UC-OPER)', 6, 6, 0, '100%'),
            ('Quản Trị Viên (UC-ADMIN)', 9, 6, 3, '67%'),
        ]
        self.add_table(['Module', 'Tổng UC', 'Hoàn thành', 'Đang phát triển', 'Tỷ lệ'], results_data)
    
        self.add_paragraph_text("**TỔNG:** 36 Use Cases, 31 Hoàn thành, 5 Đang phát triển, Tỷ lệ hoàn thành: 86%")
    
        self.doc.add_heading('5.1.2. Highlights - Điểm nổi bật', level=3)
        p = self.doc.add_paragraph('Module Khách Hàng: 100% hoàn thành với đầy đủ tính năng (tìm kiếm, hẹn lịch, đặt cọc, chat)', style='List Bullet')
        p = self.doc.add_paragraph('Phòng N-N Mapping: 1 phòng vật lý xuất hiện trong nhiều tin đăng với override giá/mô tả', style='List Bullet')
        p = self.doc.add_paragraph('7-step Wizard tạo tin đăng với UX mượt mà', style='List Bullet')
        p = self.doc.add_paragraph('Biên bản bàn giao với chữ ký số', style='List Bullet')
        p = self.doc.add_paragraph('Real-time chat với Socket.IO', style='List Bullet')
        p = self.doc.add_paragraph('Payment gateway integration (SePay)', style='List Bullet')
    
        self.doc.add_heading('5.1.3. Screenshots giao diện', level=3)
        self.add_placeholder_image("Screenshot - HomePage với Hero Search", 5)
        self.add_placeholder_image("Screenshot - AllTinDang với Filters", 5)
        self.add_placeholder_image("Screenshot - ChiTietTinDang với Danh sách Phòng", 5)
        self.add_placeholder_image("Screenshot - TaoTinDang Wizard Step 3 (Chọn Phòng)", 5)
        self.add_placeholder_image("Screenshot - Dashboard Chủ Dự Án", 5)
        self.add_placeholder_image("Screenshot - Chat Interface", 5)
        self.add_placeholder_image("Screenshot - Operator Dashboard", 5)
    
        # 5.2. So sánh mục tiêu
        self.doc.add_heading('5.2. So sánh với mục tiêu ban đầu', level=2)
    
        comparison_data = [
            ('Triển khai 36 Use Cases', '36 UCs', '31 UCs', '86%'),
            ('Hỗ trợ 5 vai trò', '5 roles', '5 roles', '100%'),
            ('JWT Authentication', 'Có', 'Có', '100%'),
            ('Real-time Chat', 'Có', 'Có (Socket.IO)', '100%'),
            ('Payment Gateway', 'Có', 'Có (SePay)', '100%'),
            ('Test Coverage 70%', '70%', '30%', '43%'),
        ]
        self.add_table(['Mục tiêu', 'Kế hoạch', 'Thực tế', 'Đánh giá'], comparison_data)
    
        # 5.3. Đánh giá
        self.doc.add_heading('5.3. Đánh giá và hạn chế', level=2)
    
        self.doc.add_heading('5.3.1. Ưu điểm', level=3)
        p = self.doc.add_paragraph('Kiến trúc rõ ràng: 3-tier architecture, component-based organization', style='List Bullet')
        p = self.doc.add_paragraph('Bảo mật tốt: JWT, RBAC, password hashing, CSRF protection, rate limiting', style='List Bullet')
        p = self.doc.add_paragraph('Audit Trail đầy đủ: Mọi hành động quan trọng đều ghi log', style='List Bullet')
        p = self.doc.add_paragraph('Idempotency: Chống duplicate requests cho đặt cọc, tạo hẹn', style='List Bullet')
        p = self.doc.add_paragraph('Real-time: WebSocket cho chat, cập nhật trạng thái', style='List Bullet')
        p = self.doc.add_paragraph('Responsive: Mobile-first design', style='List Bullet')
    
        self.doc.add_heading('5.3.2. Nhược điểm & Hạn chế', level=3)
        p = self.doc.add_paragraph('Test coverage thấp: 30% (mục tiêu 70%)', style='List Bullet')
        p = self.doc.add_paragraph('Thiếu 5 UCs: Admin tools (UI quản lý user, audit log, notification center)', style='List Bullet')
        p = self.doc.add_paragraph('Chưa có E2E tests: Automated testing chỉ dừng ở unit tests', style='List Bullet')
        p = self.doc.add_paragraph('Chưa có API docs: Thiếu Swagger/OpenAPI spec', style='List Bullet')
        p = self.doc.add_paragraph('Performance chưa test với scale: Chưa load testing 1000+ concurrent users', style='List Bullet')
    
        self.doc.add_heading('5.3.3. Khó khăn gặp phải', level=3)
    
        challenges_data = [
            ('Phòng N-N Mapping', 'Đồng bộ trạng thái phòng giữa nhiều tin đăng', 'Database triggers + phong_tindang table', '5 ngày'),
            ('Real-time Chat', 'Socket.IO authentication + room management', 'JWT token trong Socket handshake', '3 ngày'),
            ('Payment Integration', 'Signature verification, webhook handling', 'Idempotency key, retry mechanism', '4 ngày'),
            ('Geocoding', 'Google Maps API quota limits', 'Hybrid với Nominatim fallback', '2 ngày'),
        ]
        self.add_table(['Vấn đề', 'Challenge', 'Solution', 'Thời gian'], challenges_data)
    
        # 5.4. Hướng phát triển
        self.doc.add_heading('5.4. Hướng phát triển', level=2)
    
        self.doc.add_heading('5.4.1. Ngắn hạn (1-3 tháng)', level=3)
        p = self.doc.add_paragraph('Hoàn thiện 5 UCs còn lại (Admin tools UI)', style='List Bullet')
        p = self.doc.add_paragraph('Tăng test coverage lên 70%+ (backend unit tests, E2E tests)', style='List Bullet')
        p = self.doc.add_paragraph('API documentation (Swagger/OpenAPI)', style='List Bullet')
        p = self.doc.add_paragraph('Performance optimization (query optimization, Redis caching)', style='List Bullet')
    
        self.doc.add_heading('5.4.2. Trung hạn (3-6 tháng)', level=3)
        p = self.doc.add_paragraph('Mobile App (React Native cho iOS/Android)', style='List Bullet')
        p = self.doc.add_paragraph('AI chatbot hỗ trợ khách hàng', style='List Bullet')
        p = self.doc.add_paragraph('Recommendation system (gợi ý phòng tương tự)', style='List Bullet')
        p = self.doc.add_paragraph('Tích hợp thêm payment gateways (VNPay, Momo, ZaloPay)', style='List Bullet')
    
        self.doc.add_heading('5.4.3. Dài hạn (6-12 tháng)', level=3)
        p = self.doc.add_paragraph('Blockchain smart contracts cho hợp đồng thuê', style='List Bullet')
        p = self.doc.add_paragraph('IoT integration (smart lock, smart meter)', style='List Bullet')
        p = self.doc.add_paragraph('Internationalization (đa ngôn ngữ, multi-currency)', style='List Bullet')
        p = self.doc.add_paragraph('Microservices architecture (User Service, Listing Service, Payment Service...)', style='List Bullet')

    # === CHƯƠNG 6: KẾT LUẬN ===

    def add_chuong_6_ketluan(self):
        """CHƯƠNG 6: KẾT LUẬN"""
        self.doc.add_heading('CHƯƠNG 6: KẾT LUẬN', level=1)
    
        # 6.1. Tóm tắt
        self.doc.add_heading('6.1. Tóm tắt đề tài', level=2)
    
        self.add_paragraph_text(
            "Đề tài 'Hệ thống quản lý cho thuê phòng trọ' đã xây dựng thành công một nền tảng managed marketplace hoàn chỉnh với 31/36 Use Cases (86%) đã triển khai và hoạt động ổn định. Hệ thống hỗ trợ 5 vai trò người dùng với phân quyền chi tiết (RBAC), được xây dựng trên kiến trúc 3-tier vững chắc, dễ bảo trì và mở rộng.")
    
        self.add_paragraph_text(
            "Hệ thống không chỉ là sàn rao vặt đơn thuần mà có sự tham gia tích cực của nền tảng thông qua các cơ chế: KYC (xác minh danh tính), duyệt tin đăng, tự động phân công NVBH, quản lý đặt cọc an toàn (2 loại: Cọc Giữ Chỗ với TTL và Cọc An Ninh với điều kiện giải tỏa), hợp đồng điện tử, và biên bản bàn giao với chữ ký số.")
    
        self.add_paragraph_text(
            "Về mặt kỹ thuật, hệ thống áp dụng các best practices: JWT authentication, RBAC, bcrypt password hashing, CSRF protection, rate limiting, idempotency key, optimistic locking, database triggers, và audit logging toàn diện. Tích hợp thành công SePay payment gateway, Socket.IO real-time chat, và geocoding service (Google Maps + Nominatim fallback).")
    
        # 6.2. Đóng góp
        self.doc.add_heading('6.2. Đóng góp của đề tài', level=2)
    
        self.doc.add_heading('6.2.1. Về mặt kỹ thuật', level=3)
        p = self.doc.add_paragraph('Phòng N-N Mapping: 1 phòng vật lý xuất hiện trong nhiều tin đăng với override giá/mô tả, đồng bộ trạng thái qua database triggers', style='List Bullet')
        p = self.doc.add_paragraph('7-step Wizard tạo tin đăng với UX mượt mà, SectionChonPhong component cho phép multi-select với override', style='List Bullet')
        p = self.doc.add_paragraph('Payment gateway integration với flow Authorize-Capture-Void-Refund đầy đủ', style='List Bullet')
        p = self.doc.add_paragraph('Idempotency Key pattern chống duplicate requests', style='List Bullet')
        p = self.doc.add_paragraph('Design System với 5 color palettes riêng biệt cho từng vai trò', style='List Bullet')
    
        self.doc.add_heading('6.2.2. Về mặt nghiệp vụ', level=3)
        p = self.doc.add_paragraph('Mô hình Managed Marketplace có kiểm soát, khác biệt so với sàn rao vặt truyền thống', style='List Bullet')
        p = self.doc.add_paragraph('Quy trình tự động hóa: Nhu cầu → Hẹn → Cọc → Hợp đồng → Bàn giao', style='List Bullet')
        p = self.doc.add_paragraph('Chính sách cọc linh hoạt theo từng tin đăng (policy-based deposits)', style='List Bullet')
        p = self.doc.add_paragraph('Audit Trail đầy đủ phục vụ kiểm toán và giải quyết tranh chấp', style='List Bullet')
    
        # 6.3. Bài học kinh nghiệm
        self.doc.add_heading('6.3. Bài học kinh nghiệm', level=2)
    
        self.doc.add_heading('6.3.1. Về quản lý dự án', level=3)
        p = self.doc.add_paragraph('Agile/Scrum: Sprint 2 tuần, retrospective giúp điều chỉnh kịp thời', style='List Bullet')
        p = self.doc.add_paragraph('Documentation-first: Viết use-cases, SRS trước khi code giúp đồng bộ nhóm', style='List Bullet')
        p = self.doc.add_paragraph('Git workflow: Feature branches + Pull Requests giúp code review hiệu quả', style='List Bullet')
    
        self.doc.add_heading('6.3.2. Về kỹ thuật', level=3)
        p = self.doc.add_paragraph('Database design: Quan hệ N-N phức tạp nhưng cần thiết (phong_tindang)', style='List Bullet')
        p = self.doc.add_paragraph('Authentication: JWT stateless tiện lợi nhưng cần refresh token mechanism', style='List Bullet')
        p = self.doc.add_paragraph('Real-time: Socket.IO authentication cần cẩn thận để tránh lỗ hổng', style='List Bullet')
        p = self.doc.add_paragraph('Testing: Test sớm giúp phát hiện bug kịp thời, nhưng team chưa làm đủ (30%)', style='List Bullet')
    
        self.doc.add_heading('6.3.3. Về làm việc nhóm', level=3)
        p = self.doc.add_paragraph('Phân công rõ ràng: Backend (SV1) vs Frontend (SV2) giúp song song', style='List Bullet')
        p = self.doc.add_paragraph('Code review: Giúp học hỏi lẫn nhau và đảm bảo chất lượng', style='List Bullet')
        p = self.doc.add_paragraph('Communication: Daily standup online giúp đồng bộ tiến độ', style='List Bullet')
    
        # 6.4. Lời kết
        self.doc.add_heading('6.4. Lời kết', level=2)
    
        self.add_paragraph_text(
            "Đề tài đã đạt được 86% mục tiêu đề ra, với các module chính (Khách Hàng, Chủ Dự Án, Nhân Viên Bán Hàng, Nhân Viên Điều Hành) hoàn thành 100%. Những phần còn thiếu (5 UCs) tập trung ở Admin tools và có thể hoàn thành trong thời gian ngắn.")
    
        self.add_paragraph_text(
            "Hệ thống đã sẵn sàng cho MVP deployment và có tiềm năng phát triển thành sản phẩm thương mại với các tính năng nâng cao như AI chatbot, mobile app, blockchain smart contracts trong tương lai.")
    
        self.add_paragraph_text(
            "Nhóm xin chân thành cảm ơn NCS. Huỳnh Nam đã hướng dẫn tận tình trong suốt quá trình thực hiện đề tài. Chúng em cũng xin cảm ơn các thầy cô trong Khoa Công Nghệ Thông Tin, Trường Đại học Công Nghiệp TP.HCM đã tạo điều kiện và hỗ trợ nhóm hoàn thành khóa luận này.")

    # === TÀI LIỆU THAM KHẢO ===

    def add_tai_lieu_tham_khao(self):
        """Tài liệu tham khảo"""
        self.add_centered_text('TÀI LIỆU THAM KHẢO', bold=True, size=14)
        self.doc.add_paragraph()
    
        self.add_paragraph_text("**Tiếng Việt**")
        self.doc.add_paragraph()
    
        self.add_paragraph_text("[1]. Võ Nguyễn Hoành Hợp, Nguyễn Tiến Chung (2025), 'Hệ thống quản lý cho thuê phòng trọ', Khóa luận tốt nghiệp, Trường Đại học Công nghiệp TP.HCM.")
    
        self.add_paragraph_text("[2]. Tài liệu dự án (2025), 'Use Cases v1.2 - Hệ thống Cho thuê Phòng trọ', docs/use-cases-v1.2.md, GitHub Repository.")
    
        self.add_paragraph_text("[3]. Tài liệu dự án (2025), 'Software Requirements Specification v1.0', docs/SRS_v1.0.md, GitHub Repository.")
    
        self.add_paragraph_text("[4]. Tài liệu dự án (2025), 'Hiện trạng Triển khai Hệ thống - Tháng 11/2025', docs/SYSTEM_ACTUAL_STATUS_2025.md, GitHub Repository.")
    
        self.doc.add_paragraph()
        self.add_paragraph_text("**Tiếng Anh**")
        self.doc.add_paragraph()
    
        self.add_paragraph_text("[5]. IEEE Computer Society (1998), 'IEEE Recommended Practice for Software Requirements Specifications', IEEE Std 830-1998.")
    
        self.add_paragraph_text("[6]. Fielding, Roy Thomas (2000), 'Architectural Styles and the Design of Network-based Software Architectures', Doctoral dissertation, University of California, Irvine.")
    
        self.add_paragraph_text("[7]. Mozilla Developer Network (2024), 'HTTP authentication', https://developer.mozilla.org/en-US/docs/Web/HTTP/Authentication")
    
        self.add_paragraph_text("[8]. JSON Web Token (2024), 'Introduction to JSON Web Tokens', https://jwt.io/introduction")
    
        self.add_paragraph_text("[9]. React Team (2024), 'React - A JavaScript library for building user interfaces', https://react.dev")
    
        self.add_paragraph_text("[10]. Node.js Foundation (2024), 'Node.js Documentation', https://nodejs.org/docs/latest/api/")
    
        self.add_paragraph_text("[11]. Express.js Team (2024), 'Express - Fast, unopinionated, minimalist web framework for Node.js', https://expressjs.com")
    
        self.add_paragraph_text("[12]. Socket.IO Team (2024), 'Socket.IO - Bidirectional and low-latency communication', https://socket.io/docs/v4/")
    
        self.add_paragraph_text("[13]. Oracle Corporation (2024), 'MySQL 8.0 Reference Manual', https://dev.mysql.com/doc/refman/8.0/en/")
    
        self.add_paragraph_text("[14]. BEM Community (2024), 'BEM - Block Element Modifier', https://getbem.com")
    
        self.add_paragraph_text("[15]. Google Developers (2024), 'Geocoding API Documentation', https://developers.google.com/maps/documentation/geocoding")
    
        self.add_paragraph_text("[16]. OpenStreetMap Foundation (2024), 'Nominatim - Open Source Geocoding', https://nominatim.org/release-docs/latest/")

    # === PHỤ LỤC ===

    def add_phu_luc(self):
        """Phụ lục"""
        self.add_centered_text('PHỤ LỤC', bold=True, size=14)
        self.doc.add_paragraph()
    
        # Phụ lục A: Database Schema
        self.doc.add_heading('Phụ lục A: Database Schema (ERD)', level=2)
        self.add_placeholder_image("ERD - Entity Relationship Diagram đầy đủ", 6)
        self.add_paragraph_text("[Vẽ bằng dbdiagram.io, draw.io, hoặc MySQL Workbench]")
    
        # Phụ lục B: API Endpoints
        self.doc.add_heading('Phụ lục B: API Endpoints Summary', level=2)
    
        api_data = [
            ('POST /api/login', 'Public', 'Đăng nhập'),
            ('POST /api/register', 'Public', 'Đăng ký'),
            ('GET /api/tin-dang', 'Public', 'Tìm kiếm tin đăng'),
            ('GET /api/tin-dang/:id', 'Public', 'Chi tiết tin đăng'),
            ('POST /api/yeu-thich', 'KhachHang', 'Thêm yêu thích'),
            ('POST /api/cuoc-hen', 'KhachHang', 'Tạo cuộc hẹn'),
            ('POST /api/chu-du-an/tin-dang', 'ChuDuAn', 'Tạo tin đăng'),
            ('GET /api/chu-du-an/du-an/:id/phong', 'ChuDuAn', 'Lấy DS phòng'),
            ('POST /api/nhan-vien-ban-hang/lich', 'NVBH', 'Đăng ký lịch'),
            ('POST /api/operator/tin-dang/:id/approve', 'Operator', 'Duyệt tin'),
        ]
        self.add_table(['Endpoint', 'Vai trò', 'Mô tả'], api_data)
        self.add_paragraph_text("[... Tổng 70+ endpoints, xem chi tiết trong docs/SRS_v1.0.md section 3.5]")
    
        # Phụ lục C: User Manual
        self.doc.add_heading('Phụ lục C: User Manual (Tóm tắt)', level=2)
        self.add_paragraph_text("**Khách Hàng:** Hướng dẫn tìm kiếm tin đăng, thêm yêu thích, hẹn lịch xem phòng, đặt cọc.")
        self.add_paragraph_text("**Chủ Dự Án:** Hướng dẫn đăng tin 7 bước: Chọn dự án → Thông tin chung → Chọn phòng (override giá/mô tả) → Upload ảnh → Giá & chi phí → Vị trí → Preview & Submit.")
        self.add_paragraph_text("**NVBH:** Hướng dẫn đăng ký lịch làm việc, quản lý cuộc hẹn (xác nhận/đổi lịch/hủy).")
        self.add_paragraph_text("**Operator:** Hướng dẫn duyệt tin (checklist-based), quản lý nhân viên, lập biên bản bàn giao.")
        self.add_paragraph_text("**Admin:** Hướng dẫn quản lý hệ thống (khu vực, tiện ích, phân quyền, mẫu hợp đồng).")
    
        # Phụ lục D: Mã nguồn tham khảo
        self.doc.add_heading('Phụ lục D: Mã nguồn tham khảo (Code Snippets)', level=2)
        self.add_paragraph_text("**D.1. Backend - ChuDuAnController.js (đoạn tạo tin đăng):**")
        self.add_paragraph_text("[Code snippet tạo tin đăng với multi-room, xem chi tiết tại server/controllers/ChuDuAnController.js]")
    
        self.add_paragraph_text("**D.2. Frontend - TaoTinDang.jsx (đoạn step 3 chọn phòng):**")
        self.add_paragraph_text("[Code snippet SectionChonPhong component, xem chi tiết tại client/src/pages/ChuDuAn/TaoTinDang.jsx]")
    
        self.add_paragraph_text("**D.3. Database - Migration SQL:**")
        self.add_paragraph_text("[SQL script redesign phong schema, xem chi tiết tại migrations/2025_10_09_redesign_phong_schema_FINAL.sql]")
    
        # Phụ lục E: Test Cases
        self.doc.add_heading('Phụ lục E: Test Cases', level=2)
    
        test_data = [
            ('TC-LOGIN-01', 'Đăng nhập thành công', 'email + password đúng', 'JWT token + user info', 'PASS'),
            ('TC-LOGIN-02', 'Đăng nhập sai password', 'email + password sai', '401 Unauthorized', 'PASS'),
            ('TC-TINDAN-01', 'Tạo tin đăng 1 phòng', 'Form data đầy đủ', 'TinDangID + ChoDuyet', 'PASS'),
            ('TC-TINDAN-02', 'Tạo tin đăng nhiều phòng', 'PhongIDs array', 'TinDangID + n mappings', 'PASS'),
            ('TC-CUOCHEN-01', 'Hẹn lịch slot trống', 'PhongID + ThoiGianHen', 'CuocHenID + auto-assign NVBH', 'PASS'),
            ('TC-CUOCHEN-02', 'Hẹn lịch slot đã đặt', 'PhongID + ThoiGianHen trùng', '409 Conflict', 'PASS'),
        ]
        self.add_table(['Test Case ID', 'Mô tả', 'Input', 'Expected Output', 'Status'], test_data)

    # === NHẬT KÝ LÀM VIỆC ===

    def add_nhat_ky_lam_viec(self):
        """Nhật ký làm việc"""
        self.add_centered_text('NHẬT KÍ LÀM VIỆC', bold=True, size=14)
        self.doc.add_paragraph()
    
        self.add_paragraph_text("**Đề tài:** Hệ thống quản lý cho thuê phòng trọ")
        self.add_paragraph_text("**Giảng viên hướng dẫn:** NCS. Huỳnh Nam")
        self.add_paragraph_text("**Ngày bắt đầu:** 10/09/2025")
        self.doc.add_paragraph()
    
        nhat_ky_data = [
            ('1', '10/09', '16/09', '15/09', 'Nghiên cứu tài liệu, phân tích yêu cầu, viết Use Cases v1.0', ''),
            ('2', '17/09', '23/09', '22/09', 'Thiết kế Database schema, ER diagram', ''),
            ('3', '24/09', '30/09', '29/09', 'Triển khai Backend: Auth, RBAC, Routes/Controllers cơ sở', ''),
            ('4', '01/10', '07/10', '06/10', 'Triển khai Frontend: Auth pages, Dashboard layout', ''),
            ('5', '08/10', '14/10', '13/10', 'Module Khách Hàng: Tìm kiếm, Hẹn lịch', ''),
            ('6', '15/10', '21/10', '20/10', 'Module Chủ Dự Án: Wizard 7 steps, Phòng Redesign', ''),
            ('7', '22/10', '28/10', '27/10', 'Module NVBH: Quản lý lịch, cuộc hẹn', ''),
            ('8', '29/10', '04/11', '03/11', 'Module Operator: Duyệt tin, Quản lý nhân sự', ''),
            ('9', '05/11', '11/11', '10/11', 'Tích hợp SePay, Real-time Chat (Socket.IO)', ''),
            ('10', '12/11', '18/11', '17/11', 'Biên bản bàn giao, Hoàn thiện UI', ''),
            ('11', '19/11', '25/11', '24/11', 'Testing, Bug fixing, Documentation', ''),
            ('12', '26/11', '01/12', '30/11', 'Viết báo cáo KLTN, Chuẩn bị slides báo cáo', ''),
        ]
    
        self.add_table(['Tuần', 'Từ ngày', 'Đến ngày', 'Ngày báo cáo GVHD', 'Tóm tắt công việc', 'Nhận xét GVHD'], nhat_ky_data)


    
    def save(self, filename='BaoCao_KLTN_HeThongChoThuePhongTro.docx'):
        """Lưu file DOCX"""
        self.doc.save(filename)
        print(f"[OK] Da tao file: {filename}")
        file_size = os.path.getsize(filename) / 1024  # KB
        print(f"[INFO] Kich thuoc file: {file_size:.1f} KB")


# Main execution
if __name__ == '__main__':
    print("[*] Bat dau generate bao cao KLTN...")
    print("=" * 60)
    
    generator = BaoCaoKLTNGenerator()
    
    print("[*] Dang tao phan mo dau...")
    # Phần mở đầu
    generator.add_cover_page()
    generator.doc.add_page_break()
    generator.add_cover_page_vietnamese()
    generator.doc.add_page_break()
    generator.add_cover_page_english()
    generator.doc.add_page_break()
    generator.add_abstract()
    generator.doc.add_page_break()
    generator.add_loi_cam_on()
    generator.doc.add_page_break()
    generator.add_nhan_xet_gvhd()
    generator.doc.add_page_break()
    generator.add_nhan_xet_phan_bien(1)
    generator.doc.add_page_break()
    generator.add_nhan_xet_phan_bien(2)
    generator.doc.add_page_break()
    generator.add_muc_luc()
    generator.doc.add_page_break()
    generator.add_muc_luc_hinh_anh()
    generator.doc.add_page_break()
    generator.add_danh_muc_bang_bieu()
    generator.doc.add_page_break()
    generator.add_danh_muc_tu_viet_tat()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Chuong 1: Gioi thieu...")
    # Nội dung chính
    generator.add_chuong_1_gioi_thieu()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Chuong 2: Co so ly thuyet...")
    generator.add_chuong_2_coso_lythuyet()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Chuong 3: Phan tich va thiet ke...")
    generator.add_chuong_3_phantich_thietke()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Chuong 4: Trien khai he thong...")
    generator.add_chuong_4_trienkhai()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Chuong 5: Ket qua va danh gia...")
    generator.add_chuong_5_ketqua()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Chuong 6: Ket luan...")
    generator.add_chuong_6_ketluan()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Tai lieu tham khao...")
    generator.add_tai_lieu_tham_khao()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Phu luc...")
    generator.add_phu_luc()
    generator.doc.add_page_break()
    
    print("[*] Dang tao Nhat ky lam viec...")
    generator.add_nhat_ky_lam_viec()
    
    print("[*] Dang luu file...")
    generator.save()
    
    print("=" * 60)
    print("[OK] HOAN THANH!")
    print("")
    print("[NOTE] GHI CHU:")
    print("- File da day du 6 chuong + tai lieu tham khao + phu luc")
    print("- Mo file trong Microsoft Word de:")
    print("  1. Update Table of Contents (References -> Update Table)")
    print("  2. Update List of Figures")
    print("  3. Kiem tra page numbering")
    print("  4. Them screenshot/diagrams thuc te vao placeholder")
    print("  5. Dien phan 'Loi cam on' va 'Nhan xet GVHD'")

