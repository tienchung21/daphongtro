"""
Script xuất báo cáo KLTN từ baocao_data_full.py ra file DOCX
Sử dụng: python generate_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from baocao_data_full import CHUONG_1, CHUONG_2, CHUONG_3, CHUONG_4

def create_styles(doc):
    """Tạo các styles cho document"""
    # Title style
    styles = doc.styles
    
    # Heading 1 - Chương
    heading1 = styles['Heading 1']
    heading1.font.name = 'Times New Roman'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2 - Section
    heading2 = styles['Heading 2']
    heading2.font.name = 'Times New Roman'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    
    # Heading 3 - Subsection
    heading3 = styles['Heading 3']
    heading3.font.name = 'Times New Roman'
    heading3.font.size = Pt(13)
    heading3.font.bold = True
    
    # Normal text
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(13)
    
    # Code style
    try:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    except:
        code_style = styles['Code']
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)

def parse_markdown_bold(text):
    """Parse **text** thành bold (và loại bỏ ** markers)"""
    parts = []
    current = 0
    for match in re.finditer(r'\*\*(.+?)\*\*', text):
        if match.start() > current:
            parts.append(('normal', text[current:match.start()]))
        parts.append(('bold', match.group(1)))  # Chỉ lấy nội dung, bỏ **
        current = match.end()
    if current < len(text):
        parts.append(('normal', text[current:]))
    return parts if parts else [('normal', text)]

def add_formatted_paragraph(doc, text, style='Normal'):
    """Thêm paragraph với formatting (bold, italic)"""
    if text.strip().startswith('```'):
        # Code block
        lines = text.strip().split('\n')
        for line in lines:
            if line.strip() and not line.strip().startswith('```'):
                p = doc.add_paragraph(line, style='Code')
        return
    
    # Check if heading
    if text.startswith('**') and text.endswith('**') and '\n' not in text:
        # Heading 3 style
        clean_text = text.strip('*').strip()
        doc.add_heading(clean_text, level=3)
        return
    
    # Normal paragraph with bold support
    p = doc.add_paragraph(style=style)
    parts = parse_markdown_bold(text)
    for part_type, part_text in parts:
        run = p.add_run(part_text)
        if part_type == 'bold':
            run.bold = True

def process_content(doc, content):
    """Xử lý content với markdown formatting"""
    lines = content.strip().split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Code block
        if line.strip().startswith('```'):
            code_lines = [line]
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                code_lines.append(lines[i])
            
            code_text = '\n'.join(code_lines)
            add_formatted_paragraph(doc, code_text)
            i += 1
            continue
        
        # Table (simple detection)
        if '|' in line and line.count('|') > 2:
            # Skip table for now (complex formatting)
            while i < len(lines) and '|' in lines[i]:
                i += 1
            continue
        
        # Bullet list (chỉ nếu bắt đầu bằng - hoặc •)
        if line.strip().startswith('- ') or line.strip().startswith('• '):
            text = line.strip()[2:]
            # Parse bold trong bullet text
            p = doc.add_paragraph(style='List Bullet')
            parts = parse_markdown_bold(text)
            for part_type, part_text in parts:
                run = p.add_run(part_text)
                if part_type == 'bold':
                    run.bold = True
            i += 1
            continue
        
        # Numbered list - KHÔNG parse nếu có ** (là heading)
        match = re.match(r'^\s*(\d+)\.\s+(.+)$', line)
        if match and '**' not in line:
            text = match.group(2)
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Special: Dòng "28. **Title:**" → Xử lý như paragraph với bold
        if match and '**' in line:
            # Không tạo numbered list, xử lý như normal paragraph
            p = doc.add_paragraph(style='Normal')
            parts = parse_markdown_bold(line)
            for part_type, part_text in parts:
                run = p.add_run(part_text)
                if part_type == 'bold':
                    run.bold = True
            i += 1
            continue
        
        # Heading 3 (starts with **)
        if line.strip().startswith('**') and line.strip().endswith('**'):
            heading_text = line.strip('*').strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # Normal paragraph
        add_formatted_paragraph(doc, line)
        i += 1

def add_chapter(doc, chapter_data):
    """Thêm một chương vào document"""
    # Add chapter title
    doc.add_heading(chapter_data['tieu_de'], level=1)
    doc.add_paragraph()  # Spacing
    
    # Add sections
    for section in chapter_data['sections']:
        # Section title
        section_title = f"{section['id']}. {section['title']}"
        doc.add_heading(section_title, level=2)
        
        # Section content
        process_content(doc, section['content'])
        
        # Add spacing after section
        doc.add_paragraph()

def generate_docx(output_file='BaoCao_KLTN_VideoCall_AI_Translation.docx'):
    """Tạo file DOCX từ dữ liệu"""
    print("🚀 Bắt đầu tạo file DOCX...")
    
    # Create document
    doc = Document()
    
    # Set up styles
    create_styles(doc)
    
    # Add cover page
    title = doc.add_heading('BÁO CÁO KHÓA LUẬN TỐT NGHIỆP', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle = doc.add_heading('HỆ THỐNG QUẢN LÝ CHO THUÊ PHÒNG TRỌ', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle2 = doc.add_paragraph('Tích hợp Video Call với AI Translation Real-time')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(14)
    subtitle2.runs[0].bold = True
    
    # Add page break
    doc.add_page_break()
    
    # Add chapters
    print("📄 Đang xử lý Chương 1...")
    add_chapter(doc, CHUONG_1)
    doc.add_page_break()
    
    print("📄 Đang xử lý Chương 2...")
    add_chapter(doc, CHUONG_2)
    doc.add_page_break()
    
    print("📄 Đang xử lý Chương 3...")
    add_chapter(doc, CHUONG_3)
    doc.add_page_break()
    
    print("📄 Đang xử lý Chương 4...")
    add_chapter(doc, CHUONG_4)
    
    # Add bibliography (citations)
    doc.add_page_break()
    doc.add_heading('TÀI LIỆU THAM KHẢO', level=1)
    
    references = [
        "[1] Sherpa-ONNX Documentation, \"Streaming Speech Recognition,\" k2-fsa, 2024. [Online]. Available: https://k2-fsa.github.io/sherpa/onnx/. [Accessed: Nov. 19, 2025].",
        "[2] VinAI Research, \"VinAI Translate v2 Model Card,\" VinAI, 2024. [Online]. Available: https://github.com/VinAIResearch/. [Accessed: Nov. 19, 2025].",
        "[3] Piper TTS, \"Fast, local neural text to speech system,\" GitHub, 2024. [Online]. Available: https://github.com/rhasspy/piper. [Accessed: Nov. 19, 2025].",
        "[4] Google Cloud, \"Cloud Translation Pricing,\" Google, 2024. [Online]. Available: https://cloud.google.com/translate/pricing. [Accessed: Nov. 19, 2025].",
        "[5] MediaSoup, \"Cutting Edge WebRTC Video Conferencing,\" MediaSoup v3, 2024. [Online]. Available: https://mediasoup.org/. [Accessed: Nov. 19, 2025].",
        "[6] W3C, \"WebRTC 1.0: Real-Time Communication Between Browsers,\" W3C Recommendation, 2024. [Online]. Available: https://www.w3.org/TR/webrtc/. [Accessed: Nov. 19, 2025].",
        "[7] Google Cloud, \"Cloud Speech-to-Text Pricing,\" Google, 2024. [Online]. Available: https://cloud.google.com/speech-to-text/pricing. [Accessed: Nov. 19, 2025].",
        "[8] Meta AI, \"No Language Left Behind: Scaling Human-Centered Machine Translation,\" arXiv:2207.04672, 2022. [Online]. Available: https://arxiv.org/abs/2207.04672. [Accessed: Nov. 19, 2025].",
        "[9] AWS, \"Amazon Transcribe Pricing,\" Amazon Web Services, 2024. [Online]. Available: https://aws.amazon.com/transcribe/pricing/. [Accessed: Nov. 19, 2025]."
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref, style='Normal')
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.left_indent = Inches(0.5)
    
    # Save document
    doc.save(output_file)
    print(f"✅ Hoàn thành! File đã được lưu: {output_file}")
    print(f"📊 Tổng số trang: ~{len(doc.sections)} sections")
    print(f"📝 Tổng số paragraphs: {len(doc.paragraphs)}")

if __name__ == '__main__':
    try:
        generate_docx()
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()
