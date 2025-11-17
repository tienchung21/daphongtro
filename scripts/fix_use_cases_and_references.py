#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script sửa lại Use Cases và thêm trích dẫn chuẩn IEEE
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def fix_use_cases_table(doc):
    """Sửa lại bảng Use Cases với thông tin chính xác"""
    
    # Tìm và xóa bảng Use Cases cũ (nếu có)
    # Thêm bảng mới với thông tin chính xác
    
    print("Đang sửa bảng Use Cases...")
    
    # Thông tin Use Cases CHÍNH XÁC từ use-cases-v1.2.md
    use_cases_data = [
        ("Mã nhóm", "Tên nhóm", "Số lượng", "Người dùng chính"),
        ("UC-GEN", "Chức năng Chung", "5 UCs", "Tất cả người dùng"),
        ("UC-CUST", "Khách hàng", "7 UCs", "Khách hàng"),
        ("UC-SALE", "Nhân viên Bán hàng", "7 UCs", "Nhân viên BH"),
        ("UC-PROJ", "Chủ dự án", "5 UCs", "Chủ dự án"),
        ("UC-OPER", "Nhân viên Điều hành", "6 UCs", "Điều hành"),
        ("UC-ADMIN", "Quản trị Hệ thống", "9 UCs", "Admin"),
        ("", "TỔNG CỘNG", "39 UCs", ""),
    ]
    
    return use_cases_data

def add_ieee_references(doc):
    """Thêm trích dẫn chuẩn IEEE thực tế"""
    
    print("Đang thêm trích dẫn IEEE...")
    
    # Tìm section "Tài liệu tham khảo"
    # Thay thế bằng các trích dẫn chuẩn IEEE thực tế
    
    ieee_references = [
        # React & Frontend
        "[1] Facebook Inc., \"React - A JavaScript library for building user interfaces,\" Meta Open Source, 2023. [Online]. Available: https://react.dev/. [Accessed: Nov. 7, 2025].",
        
        # Node.js & Backend
        "[2] OpenJS Foundation, \"Node.js - JavaScript runtime,\" Node.js Foundation, 2024. [Online]. Available: https://nodejs.org/. [Accessed: Nov. 7, 2025].",
        
        # MySQL
        "[3] Oracle Corporation, \"MySQL 8.0 Reference Manual,\" Oracle, 2024. [Online]. Available: https://dev.mysql.com/doc/refman/8.0/en/. [Accessed: Nov. 7, 2025].",
        
        # Socket.IO
        "[4] G. Rauch, \"Socket.IO - Realtime application framework,\" Socket.IO, 2024. [Online]. Available: https://socket.io/. [Accessed: Nov. 7, 2025].",
        
        # JWT Authentication
        "[5] M. Jones, J. Bradley, and N. Sakimura, \"JSON Web Token (JWT),\" RFC 7519, May 2015. [Online]. Available: https://tools.ietf.org/html/rfc7519. [Accessed: Nov. 7, 2025].",
        
        # RESTful API
        "[6] R. Fielding, \"Architectural Styles and the Design of Network-based Software Architectures,\" Ph.D. dissertation, Univ. California, Irvine, 2000. [Online]. Available: https://www.ics.uci.edu/~fielding/pubs/dissertation/top.htm. [Accessed: Nov. 7, 2025].",
        
        # Marketplace platforms
        "[7] A. Hagiu and J. Wright, \"Multi-Sided Platforms,\" International Journal of Industrial Organization, vol. 43, pp. 162-174, Nov. 2015, doi: 10.1016/j.ijindorg.2015.03.003.",
        
        # Web Security
        "[8] OWASP Foundation, \"OWASP Top Ten 2021,\" Open Web Application Security Project, 2021. [Online]. Available: https://owasp.org/Top10/. [Accessed: Nov. 7, 2025].",
        
        # Payment Gateway
        "[9] PCI Security Standards Council, \"Payment Card Industry Data Security Standard (PCI DSS) v4.0,\" PCI SSC, 2022. [Online]. Available: https://www.pcisecuritystandards.org/. [Accessed: Nov. 7, 2025].",
        
        # Thị trường BĐS Việt Nam
        "[10] Savills Vietnam, \"Vietnam Property Market Report 2023-2024,\" Savills Research, Ho Chi Minh City, Vietnam, 2024. [Online]. Available: https://www.savills.com.vn/research/. [Accessed: Nov. 7, 2025].",
        
        # Software Architecture
        "[11] M. Fowler, \"Patterns of Enterprise Application Architecture,\" Boston, MA: Addison-Wesley, 2002.",
        
        # Node.js Best Practices
        "[12] Y. Goldberg et al., \"Node.js Best Practices,\" GitHub Repository, 2024. [Online]. Available: https://github.com/goldbergyoni/nodebestpractices. [Accessed: Nov. 7, 2025].",
    ]
    
    return ieee_references

def create_corrected_chapter_3(output_file):
    """Tạo lại Chương 3 với thông tin chính xác"""
    
    print("Đang tạo Chương 3 đã sửa...")
    
    doc = Document()
    
    # === CHƯƠNG 3 ===
    heading = doc.add_heading('CHƯƠNG 3', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    heading2 = doc.add_heading('PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', level=1)
    heading2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # 3.1 Phân tích yêu cầu
    doc.add_heading('3.1. Phân tích yêu cầu', level=2)
    
    # 3.1.1 Use Cases
    doc.add_heading('3.1.1. Đặc tả Use Cases', level=3)
    
    p = doc.add_paragraph()
    p.add_run(
        'Hệ thống được phân tích thành 39 Use Cases, được tổ chức thành 6 nhóm chức năng chính '
        'theo vai trò người dùng. Dưới đây là bảng tổng hợp:'
    )
    
    doc.add_paragraph()
    
    # Tạo bảng Use Cases CHÍNH XÁC
    use_cases_data = fix_use_cases_table(doc)
    
    table = doc.add_table(rows=len(use_cases_data), cols=4)
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(use_cases_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            
            # Header row hoặc Total row: in đậm
            if i == 0 or i == len(use_cases_data) - 1:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()
    
    # Chi tiết từng nhóm Use Cases
    doc.add_heading('3.1.1.1. UC-GEN - Chức năng Chung (5 UCs)', level=4)
    p = doc.add_paragraph()
    p.add_run('Nhóm này bao gồm các chức năng cơ bản mà tất cả người dùng đều sử dụng:')
    
    use_case_gen = [
        'UC-GEN-01: Đăng nhập - Xác thực và tạo phiên truy cập an toàn',
        'UC-GEN-02: Đăng ký tài khoản - Cho phép người dùng mới tạo tài khoản',
        'UC-GEN-03: Chuyển đổi vai trò - Chuyển đổi giữa các vai trò (multi-role)',
        'UC-GEN-04: Xem danh sách cuộc hẹn - Quản lý lịch hẹn theo vai trò',
        'UC-GEN-05: Trung tâm thông báo - Xem và quản lý thông báo',
    ]
    
    for uc in use_case_gen:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.1.1.2. UC-CUST - Khách hàng (7 UCs)', level=4)
    p = doc.add_paragraph()
    p.add_run('Nhóm chức năng dành cho khách hàng tìm kiếm và thuê phòng:')
    
    use_case_cust = [
        'UC-CUST-01: Tìm kiếm phòng trọ - Tìm tin đăng với bộ lọc nâng cao',
        'UC-CUST-02: Quản lý yêu thích - Lưu tin đăng quan tâm',
        'UC-CUST-03: Hẹn lịch xem phòng - Đặt lịch hẹn xem phòng',
        'UC-CUST-04: Thực hiện đặt cọc - Đặt cọc giữ chỗ hoặc cọc an ninh',
        'UC-CUST-05: Hủy giao dịch - Yêu cầu hoàn tiền theo chính sách',
        'UC-CUST-06: Quản lý ví điện tử - Quản lý số dư và lịch sử giao dịch',
        'UC-CUST-07: Nhắn tin - Chat với chủ dự án/nhân viên',
    ]
    
    for uc in use_case_cust:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.1.1.3. UC-SALE - Nhân viên Bán hàng (7 UCs)', level=4)
    p = doc.add_paragraph()
    p.add_run('Nhóm chức năng hỗ trợ nhân viên bán hàng quản lý cuộc hẹn:')
    
    use_case_sale = [
        'UC-SALE-01: Đăng ký lịch làm việc - Thiết lập ca làm việc',
        'UC-SALE-02: Xem chi tiết cuộc hẹn - Xem thông tin cuộc hẹn được giao',
        'UC-SALE-03: Quản lý cuộc hẹn - Xác nhận/Đổi lịch/Hủy',
        'UC-SALE-04: Xác nhận cọc - Xác nhận khoản đặt cọc của khách',
        'UC-SALE-05: Báo cáo kết quả cuộc hẹn - Ghi nhận kết quả sau khi hẹn',
        'UC-SALE-06: Xem báo cáo thu nhập - Xem hoa hồng và thống kê',
        'UC-SALE-07: Nhắn tin - Chat với khách hàng',
    ]
    
    for uc in use_case_sale:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.1.1.4. UC-PROJ - Chủ dự án (5 UCs)', level=4)
    p = doc.add_paragraph()
    p.add_run('Nhóm chức năng cho chủ dự án quản lý tin đăng và báo cáo:')
    
    use_case_proj = [
        'UC-PROJ-01: Đăng tin cho thuê - Tạo và quản lý tin đăng (CRUD đầy đủ)',
        'UC-PROJ-02: Xác nhận cuộc hẹn - Duyệt lịch hẹn (nếu yêu cầu)',
        'UC-PROJ-03: Xem báo cáo kinh doanh - Dashboard và thống kê',
        'UC-PROJ-04: Báo cáo hợp đồng cho thuê - Quản lý hợp đồng',
        'UC-PROJ-05: Nhắn tin - Chat với khách hàng/nhân viên',
    ]
    
    for uc in use_case_proj:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.1.1.5. UC-OPER - Nhân viên Điều hành (6 UCs)', level=4)
    p = doc.add_paragraph()
    p.add_run('Nhóm chức năng vận hành và quản lý nghiệp vụ:')
    
    use_case_oper = [
        'UC-OPER-01: Duyệt tin đăng - Kiểm duyệt tin đăng mới',
        'UC-OPER-02: Quản lý danh sách dự án - CRUD dự án',
        'UC-OPER-03: Quản lý lịch làm việc NVBH - Điều phối nhân sự',
        'UC-OPER-04: Quản lý hồ sơ nhân viên - Xem và cập nhật hồ sơ',
        'UC-OPER-05: Tạo tài khoản nhân viên - Onboarding nhân viên mới',
        'UC-OPER-06: Lập biên bản bàn giao - Ghi nhận bàn giao phòng',
    ]
    
    for uc in use_case_oper:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('3.1.1.6. UC-ADMIN - Quản trị Hệ thống (9 UCs)', level=4)
    p = doc.add_paragraph()
    p.add_run('Nhóm chức năng quản trị toàn hệ thống:')
    
    use_case_admin = [
        'UC-ADMIN-01: Quản lý tài khoản người dùng - CRUD users, lock/unlock',
        'UC-ADMIN-02: Quản lý danh sách dự án - Quản lý tất cả dự án',
        'UC-ADMIN-03: Quản lý danh sách khu vực - Quản lý địa điểm',
        'UC-ADMIN-04: Xem báo cáo thu nhập toàn hệ thống - Analytics tổng thể',
        'UC-ADMIN-05: Quản lý chính sách - Cấu hình policies',
        'UC-ADMIN-06: Quản lý mẫu hợp đồng - Version control mẫu hợp đồng',
        'UC-ADMIN-07: Quản lý quyền & Phân quyền - RBAC configuration',
        'UC-ADMIN-08: Xem nhật ký hệ thống - Audit logs',
        'UC-ADMIN-09: Quản lý chính sách cọc - Cấu hình chính sách cọc',
    ]
    
    for uc in use_case_admin:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph()
    
    # Lưu file
    doc.save(output_file)
    print(f"✅ Đã tạo file: {output_file}")

def create_corrected_references(output_file):
    """Tạo file Tài liệu tham khảo với trích dẫn IEEE chuẩn"""
    
    print("Đang tạo Tài liệu tham khảo...")
    
    doc = Document()
    
    heading = doc.add_heading('TÀI LIỆU THAM KHẢO', level=1)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # Thêm trích dẫn IEEE
    ieee_refs = add_ieee_references(doc)
    
    for ref in ieee_refs:
        p = doc.add_paragraph()
        p.add_run(ref)
        # Hanging indent cho citation
        p.paragraph_format.left_indent = Pt(36)
        p.paragraph_format.first_line_indent = Pt(-36)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    
    # Lưu file
    doc.save(output_file)
    print(f"✅ Đã tạo file: {output_file}")

if __name__ == "__main__":
    print("="*80)
    print("SỬA LẠI USE CASES VÀ THÊM TRÍCH DẪN IEEE")
    print("="*80)
    print()
    
    # Tạo Chương 3 đã sửa
    create_corrected_chapter_3("BaoCao_Chuong3_CORRECTED.docx")
    
    print()
    
    # Tạo Tài liệu tham khảo với IEEE citations
    create_corrected_references("TaiLieuThamKhao_IEEE.docx")
    
    print()
    print("="*80)
    print("✅ HOÀN THÀNH!")
    print("="*80)
    print()
    print("📄 Files đã tạo:")
    print("   - BaoCao_Chuong3_CORRECTED.docx")
    print("   - TaiLieuThamKhao_IEEE.docx")
    print()
    print("⚠️  NEXT STEPS:")
    print("   1. Mở BaoCao_Chuong3_CORRECTED.docx để kiểm tra")
    print("   2. Copy nội dung vào file FINAL")
    print("   3. Copy Tài liệu tham khảo vào file FINAL")
    print()





