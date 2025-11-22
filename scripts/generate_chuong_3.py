#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Generator for Chuong 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
- Use Cases (39 UCs từ docs/use-cases-v1.2.md)
- Database Schema (32 tables, 52 foreign keys)
- API Design (146 endpoints)
- State Machines
- UI/UX Design
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def setup_document_styles(doc):
    """Setup document styles"""
    styles = doc.styles
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(13)
    h3_style.font.bold = True
    
    # Normal style
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(13)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.space_after = Pt(6)

def add_paragraph(doc, text, bold=False, italic=False):
    """Add paragraph with formatting"""
    p = doc.add_paragraph(text)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    return p

def add_bullet_list(doc, items):
    """Add bullet list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

def add_numbered_list(doc, items):
    """Add numbered list"""
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)

def add_table_with_data(doc, headers, rows):
    """Add table with data"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def generate_section_3_1(doc):
    """3.1. Phân tích yêu cầu hệ thống"""
    doc.add_heading('3.1. Phân tích yêu cầu hệ thống', level=2)
    
    # 3.1.1. Yêu cầu chức năng
    doc.add_heading('3.1.1. Yêu cầu chức năng', level=3)
    
    add_paragraph(doc, """Hệ thống được xây dựng dựa trên mô hình Managed Marketplace [7] với 39 Use Cases 
chính được phân loại theo 6 nhóm chức năng lớn theo vai trò người dùng [11]:""")
    
    # Table: Use Cases Overview (CORRECTED - từ use-cases-v1.2.md)
    use_case_groups = [
        ['UC-GEN', 'Chức năng Chung', '5 UCs', 'Tất cả người dùng'],
        ['UC-CUST', 'Khách hàng', '7 UCs', 'Khách hàng'],
        ['UC-SALE', 'Nhân viên Bán hàng', '7 UCs', 'Nhân viên BH'],
        ['UC-PROJ', 'Chủ dự án', '5 UCs', 'Chủ dự án'],
        ['UC-OPER', 'Nhân viên Điều hành', '6 UCs', 'Điều hành'],
        ['UC-ADMIN', 'Quản trị Hệ thống', '9 UCs', 'Admin']
    ]
    
    add_table_with_data(doc, 
        ['Mã nhóm', 'Tên nhóm', 'Số lượng', 'Người dùng chính'],
        use_case_groups
    )
    
    doc.add_paragraph()
    
    # 3.1.2. Yêu cầu phi chức năng
    doc.add_heading('3.1.2. Yêu cầu phi chức năng', level=3)
    
    nfr_list = [
        "Performance: Thời gian phản hồi API < 500ms cho 95% requests",
        "Scalability: Hỗ trợ 10,000+ concurrent users",
        "Security: JWT authentication, role-based authorization, SQL injection prevention",
        "Reliability: Uptime 99.9%, auto-retry cho failed transactions",
        "Usability: Responsive design, Mobile-first approach",
        "Maintainability: Modular architecture, comprehensive logging"
    ]
    
    add_bullet_list(doc, nfr_list)

def generate_section_3_2(doc):
    """3.2. Thiết kế Use Cases chi tiết"""
    doc.add_heading('3.2. Thiết kế Use Cases chi tiết', level=2)
    
    add_paragraph(doc, """Dưới đây là các Use Cases quan trọng nhất của hệ thống, được thiết kế theo 
chuẩn UML Use Case Specification:""")
    
    # 3.2.1. UC-PROJ-01: Đăng tin Cho thuê
    doc.add_heading('3.2.1. UC-PROJ-01: Đăng tin Cho thuê', level=3)
    
    uc_proj_01 = [
        ['Use Case ID', 'UC-PROJ-01'],
        ['Tên', 'Đăng tin Cho thuê'],
        ['Actor chính', 'Chủ dự án (ChuDuAn)'],
        ['Tiền điều kiện', 'ChuDuAn đã đăng nhập, đã KYC hoặc cho phép tạo trước KYC'],
        ['Luồng chính', '1. ChuDuAn chọn "Đăng tin mới"\n2. Hệ thống hiển thị form đăng tin\n3. ChuDuAn nhập thông tin: Tiêu đề, Mô tả, Giá, Diện tích, Địa chỉ\n4. ChuDuAn upload ảnh (tối thiểu 1 ảnh)\n5. ChuDuAn chọn thuộc tính và tiện ích\n6. Hệ thống validate dữ liệu\n7. Hệ thống tạo TinDang với trạng thái "ChoDuyet"\n8. Hệ thống gửi thông báo đến Admin'],
        ['Hậu điều kiện', 'TinDang được tạo với trạng thái "ChoDuyet"'],
        ['Luồng thay thế', '6a. Dữ liệu không hợp lệ: Hiển thị lỗi validation']
    ]
    
    add_table_with_data(doc, ['Field', 'Value'], uc_proj_01)
    
    doc.add_paragraph()
    
    # 3.2.2. UC-TEN-01: Đặt lịch xem phòng
    doc.add_heading('3.2.2. UC-TEN-01: Đặt lịch xem phòng', level=3)
    
    uc_ten_01 = [
        ['Use Case ID', 'UC-TEN-01'],
        ['Tên', 'Đặt lịch xem phòng'],
        ['Actor chính', 'Người thuê (NguoiThue)'],
        ['Tiền điều kiện', 'NguoiThue đã đăng nhập, TinDang ở trạng thái "DaDang"'],
        ['Luồng chính', '1. NguoiThue xem chi tiết TinDang\n2. NguoiThue chọn "Đặt lịch xem phòng"\n3. Hệ thống hiển thị lịch khả dụng\n4. NguoiThue chọn ngày giờ xem phòng\n5. NguoiThue nhập ghi chú (optional)\n6. Hệ thống tạo CuocHen với trạng thái "ChoXacNhan"\n7. Hệ thống gửi thông báo đến ChuDuAn\n8. Hệ thống gửi real-time notification qua Socket.IO'],
        ['Hậu điều kiện', 'CuocHen được tạo với trạng thái "ChoXacNhan"'],
        ['Luồng thay thế', '5a. Lịch không khả dụng: Đề xuất lịch khác']
    ]
    
    add_table_with_data(doc, ['Field', 'Value'], uc_ten_01)
    
    doc.add_paragraph()
    
    # 3.2.3. UC-DEPOSIT-01: Đặt cọc phòng
    doc.add_heading('3.2.3. UC-DEPOSIT-01: Đặt cọc phòng', level=3)
    
    uc_deposit_01 = [
        ['Use Case ID', 'UC-DEPOSIT-01'],
        ['Tên', 'Đặt cọc phòng'],
        ['Actor chính', 'Người thuê (NguoiThue)'],
        ['Tiền điều kiện', 'CuocHen ở trạng thái "DaXong", NguoiThue đã KYC'],
        ['Luồng chính', '1. NguoiThue chọn "Đặt cọc phòng"\n2. Hệ thống hiển thị thông tin cọc theo ChinhSachCoc\n3. NguoiThue xác nhận điều khoản\n4. Hệ thống tạo Coc với trạng thái "ChoThanhToan"\n5. Hệ thống tích hợp SePay để tạo payment link\n6. NguoiThue thực hiện thanh toán\n7. SePay webhook callback xác nhận thanh toán\n8. Hệ thống cập nhật Coc sang "DaThanhToan"\n9. Hệ thống cập nhật TinDang sang "DaCoc"'],
        ['Hậu điều kiện', 'Coc được tạo với trạng thái "DaThanhToan", TinDang chuyển sang "DaCoc"'],
        ['Luồng thay thế', '7a. Thanh toán thất bại: Giữ nguyên trạng thái "ChoThanhToan"']
    ]
    
    add_table_with_data(doc, ['Field', 'Value'], uc_deposit_01)
    
    doc.add_paragraph()
    
    # 3.2.4. UC-CONTRACT-01: Tạo hợp đồng
    doc.add_heading('3.2.4. UC-CONTRACT-01: Tạo hợp đồng', level=3)
    
    uc_contract_01 = [
        ['Use Case ID', 'UC-CONTRACT-01'],
        ['Tên', 'Tạo hợp đồng thuê'],
        ['Actor chính', 'Chủ dự án (ChuDuAn)'],
        ['Tiền điều kiện', 'Coc ở trạng thái "DaThanhToan"'],
        ['Luồng chính', '1. ChuDuAn chọn "Tạo hợp đồng"\n2. Hệ thống auto-fill thông tin từ TinDang và Coc\n3. ChuDuAn nhập thông tin hợp đồng: NgayBatDau, NgayKetThuc, GiaThue, TienCoc\n4. ChuDuAn chọn các dịch vụ đi kèm\n5. Hệ thống generate PDF hợp đồng\n6. Hệ thống tạo HopDong với trạng thái "ChoKyBenA"\n7. Hệ thống gửi thông báo đến NguoiThue'],
        ['Hậu điều kiện', 'HopDong được tạo với trạng thái "ChoKyBenA"'],
        ['Luồng thay thế', '5a. Lỗi generate PDF: Hiển thị thông báo lỗi']
    ]
    
    add_table_with_data(doc, ['Field', 'Value'], uc_contract_01)
    
    doc.add_paragraph()
    
    add_paragraph(doc, """Tổng cộng có 39 Use Cases được thiết kế chi tiết theo đặc tả use-cases-v1.2.md, bao gồm:""")
    
    all_use_cases = [
        "5 UCs Chức năng Chung (UC-GEN-01 đến UC-GEN-05): Đăng nhập, Đăng ký, Chuyển vai trò, Xem cuộc hẹn, Trung tâm thông báo",
        "7 UCs Khách hàng (UC-CUST-01 đến UC-CUST-07): Tìm kiếm, Yêu thích, Hẹn xem phòng, Đặt cọc, Hủy giao dịch, Quản lý ví, Chat",
        "7 UCs Nhân viên Bán hàng (UC-SALE-01 đến UC-SALE-07): Đăng ký lịch làm việc, Quản lý cuộc hẹn, Xác nhận cọc, Báo cáo kết quả, Xem báo cáo thu nhập",
        "5 UCs Chủ dự án (UC-PROJ-01 đến UC-PROJ-05): Đăng tin cho thuê, Xác nhận cuộc hẹn, Xem báo cáo kinh doanh, Báo cáo hợp đồng, Chat",
        "6 UCs Nhân viên Điều hành (UC-OPER-01 đến UC-OPER-06): Duyệt tin, Quản lý dự án, Quản lý lịch NVBH, Quản lý hồ sơ, Tạo tài khoản NV, Lập biên bản bàn giao",
        "9 UCs Quản trị Hệ thống (UC-ADMIN-01 đến UC-ADMIN-09): Quản lý users, Quản lý dự án, Quản lý khu vực, Báo cáo thu nhập, Quản lý chính sách, Mẫu HĐ, RBAC, Nhật ký, Chính sách cọc"
    ]
    
    add_bullet_list(doc, all_use_cases)

def generate_section_3_3(doc):
    """3.3. Thiết kế cơ sở dữ liệu"""
    doc.add_heading('3.3. Thiết kế cơ sở dữ liệu', level=2)
    
    # 3.3.1. Database Schema Overview
    doc.add_heading('3.3.1. Tổng quan Database Schema', level=3)
    
    add_paragraph(doc, """Hệ thống sử dụng MySQL 8.0 [3] làm RDBMS với 32 tables được tổ chức thành 8 nhóm chức năng:""")
    
    db_groups = [
        ['Quản lý người dùng', '5 tables', 'NguoiDung, ChuDuAn, Admin, KhachThue, NhatKyHeThong'],
        ['Quản lý dự án', '4 tables', 'DuAn, TinDang, HinhAnh, YeuThich'],
        ['Thuộc tính & Tiện ích', '4 tables', 'ThuocTinh, TienIch, TinDang_ThuocTinh, TinDang_TienIch'],
        ['Cuộc hẹn & Chat', '3 tables', 'CuocHen, TinNhan, TinNhan_HinhAnh'],
        ['Cọc & Thanh toán', '3 tables', 'ChinhSachCoc, Coc, LichSuThanhToan'],
        ['Hợp đồng', '3 tables', 'HopDong, DichVu, HopDong_DichVu'],
        ['Quản lý phòng', '4 tables', 'Phong, HoaDon, HoaDon_DichVu, YeuCauBaoTri'],
        ['Thông báo & Báo cáo', '6 tables', 'ThongBao, BaoCaoHieuSuat, DanhGia, CauHoi, TraLoi, LichSuThayDoi']
    ]
    
    add_table_with_data(doc, ['Nhóm chức năng', 'Số tables', 'Tables chính'], db_groups)
    
    doc.add_paragraph()
    
    # 3.3.2. Core Tables
    doc.add_heading('3.3.2. Chi tiết Core Tables', level=3)
    
    add_paragraph(doc, bold=True, text="Table: NguoiDung")
    add_paragraph(doc, "Table quản lý thông tin người dùng chung, là parent table cho tất cả các loại người dùng.")
    
    nguoidung_fields = [
        ['NguoiDungID', 'INT', 'PRIMARY KEY', 'ID người dùng'],
        ['Email', 'VARCHAR(255)', 'UNIQUE, NOT NULL', 'Email đăng nhập'],
        ['MatKhau', 'VARCHAR(255)', 'NOT NULL', 'Mật khẩu đã hash (bcrypt)'],
        ['HoTen', 'VARCHAR(255)', 'NOT NULL', 'Họ tên đầy đủ'],
        ['SoDienThoai', 'VARCHAR(20)', 'NULL', 'Số điện thoại'],
        ['DiaChi', 'TEXT', 'NULL', 'Địa chỉ'],
        ['AnhDaiDien', 'TEXT', 'NULL', 'URL ảnh đại diện'],
        ['LoaiNguoiDung', 'ENUM', 'NOT NULL', 'Admin|ChuDuAn|KhachThue'],
        ['TrangThaiKYC', 'ENUM', 'DEFAULT ChuaXacMinh', 'ChuaXacMinh|DangXuLy|DaXacMinh'],
        ['NgayTao', 'DATETIME', 'DEFAULT NOW()', 'Ngày tạo tài khoản']
    ]
    
    add_table_with_data(doc, ['Field', 'Type', 'Constraint', 'Description'], nguoidung_fields)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Table: TinDang")
    add_paragraph(doc, "Table quản lý tin đăng cho thuê phòng trọ.")
    
    tindang_fields = [
        ['TinDangID', 'INT', 'PRIMARY KEY', 'ID tin đăng'],
        ['ChuDuAnID', 'INT', 'FOREIGN KEY', 'Reference đến ChuDuAn'],
        ['DuAnID', 'INT', 'FOREIGN KEY', 'Reference đến DuAn (NULL nếu phòng lẻ)'],
        ['TieuDe', 'VARCHAR(500)', 'NOT NULL', 'Tiêu đề tin đăng'],
        ['MoTa', 'TEXT', 'NOT NULL', 'Mô tả chi tiết'],
        ['DienTich', 'DECIMAL(10,2)', 'NOT NULL', 'Diện tích (m²)'],
        ['GiaThue', 'DECIMAL(15,2)', 'NOT NULL', 'Giá thuê (VNĐ/tháng)'],
        ['DiaChi', 'TEXT', 'NOT NULL', 'Địa chỉ đầy đủ'],
        ['Tinh', 'VARCHAR(100)', 'NOT NULL', 'Tỉnh/Thành phố'],
        ['Huyen', 'VARCHAR(100)', 'NOT NULL', 'Quận/Huyện'],
        ['Xa', 'VARCHAR(100)', 'NOT NULL', 'Phường/Xã'],
        ['TrangThai', 'ENUM', 'DEFAULT Nhap', 'Nhap|ChoDuyet|DaDuyet|DaDang|DaCoc|DaThue|DaKhoa'],
        ['SoLuotXem', 'INT', 'DEFAULT 0', 'Số lượt xem'],
        ['NgayTao', 'DATETIME', 'DEFAULT NOW()', 'Ngày tạo'],
        ['NgayCapNhat', 'DATETIME', 'ON UPDATE NOW()', 'Ngày cập nhật']
    ]
    
    add_table_with_data(doc, ['Field', 'Type', 'Constraint', 'Description'], tindang_fields)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Table: CuocHen")
    add_paragraph(doc, "Table quản lý cuộc hẹn xem phòng.")
    
    cuochen_fields = [
        ['CuocHenID', 'INT', 'PRIMARY KEY', 'ID cuộc hẹn'],
        ['TinDangID', 'INT', 'FOREIGN KEY', 'Reference đến TinDang'],
        ['NguoiThueID', 'INT', 'FOREIGN KEY', 'Reference đến KhachThue'],
        ['ChuDuAnID', 'INT', 'FOREIGN KEY', 'Reference đến ChuDuAn'],
        ['NgayHen', 'DATETIME', 'NOT NULL', 'Ngày giờ hẹn'],
        ['GhiChu', 'TEXT', 'NULL', 'Ghi chú của người thuê'],
        ['TrangThai', 'ENUM', 'DEFAULT ChoXacNhan', 'ChoXacNhan|DaXacNhan|DaHuy|DaXong'],
        ['LyDoHuy', 'TEXT', 'NULL', 'Lý do hủy (nếu có)'],
        ['NgayTao', 'DATETIME', 'DEFAULT NOW()', 'Ngày tạo']
    ]
    
    add_table_with_data(doc, ['Field', 'Type', 'Constraint', 'Description'], cuochen_fields)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Table: Coc")
    add_paragraph(doc, "Table quản lý thông tin cọc phòng.")
    
    coc_fields = [
        ['CocID', 'INT', 'PRIMARY KEY', 'ID cọc'],
        ['TinDangID', 'INT', 'FOREIGN KEY', 'Reference đến TinDang'],
        ['NguoiThueID', 'INT', 'FOREIGN KEY', 'Reference đến KhachThue'],
        ['ChuDuAnID', 'INT', 'FOREIGN KEY', 'Reference đến ChuDuAn'],
        ['ChinhSachCocID', 'INT', 'FOREIGN KEY', 'Reference đến ChinhSachCoc'],
        ['SoTienCoc', 'DECIMAL(15,2)', 'NOT NULL', 'Số tiền cọc'],
        ['TrangThai', 'ENUM', 'DEFAULT ChoThanhToan', 'ChoThanhToan|DaThanhToan|DaHoan|DaBiMat'],
        ['NgayThanhToan', 'DATETIME', 'NULL', 'Ngày thanh toán'],
        ['NgayHetHan', 'DATETIME', 'NOT NULL', 'Ngày hết hạn giữ cọc'],
        ['MaGiaoDich', 'VARCHAR(100)', 'NULL', 'Mã giao dịch từ payment gateway'],
        ['NgayTao', 'DATETIME', 'DEFAULT NOW()', 'Ngày tạo']
    ]
    
    add_table_with_data(doc, ['Field', 'Type', 'Constraint', 'Description'], coc_fields)
    
    doc.add_paragraph()
    
    # 3.3.3. Relationships
    doc.add_heading('3.3.3. Database Relationships', level=3)
    
    add_paragraph(doc, """Các mối quan hệ chính trong database:""")
    
    relationships = [
        "NguoiDung 1:1 ChuDuAn/Admin/KhachThue (Inheritance)",
        "ChuDuAn 1:N DuAn (Một chủ dự án quản lý nhiều dự án)",
        "ChuDuAn 1:N TinDang (Một chủ dự án có nhiều tin đăng)",
        "DuAn 1:N TinDang (Một dự án có nhiều tin đăng phòng)",
        "TinDang 1:N HinhAnh (Một tin đăng có nhiều ảnh)",
        "TinDang N:M ThuocTinh (Many-to-Many qua TinDang_ThuocTinh)",
        "TinDang N:M TienIch (Many-to-Many qua TinDang_TienIch)",
        "TinDang 1:N CuocHen (Một tin đăng có nhiều cuộc hẹn)",
        "TinDang 1:N Coc (Một tin đăng có nhiều lần cọc trong lịch sử)",
        "Coc 1:1 HopDong (Một cọc dẫn đến một hợp đồng)",
        "HopDong 1:1 Phong (Một hợp đồng tạo một phòng thuê)",
        "Phong 1:N HoaDon (Một phòng có nhiều hóa đơn theo tháng)"
    ]
    
    add_bullet_list(doc, relationships)
    
    doc.add_paragraph()
    
    # 3.3.4. Indexing Strategy
    doc.add_heading('3.3.4. Chiến lược Indexing', level=3)
    
    add_paragraph(doc, """Để tối ưu hiệu suất truy vấn, các index sau được tạo:""")
    
    indexes = [
        "PRIMARY KEY indexes trên tất cả ID fields",
        "UNIQUE index trên NguoiDung.Email",
        "INDEX trên TinDang.TrangThai (thường xuyên filter)",
        "INDEX trên TinDang.ChuDuAnID (thường xuyên join)",
        "COMPOSITE INDEX trên (Tinh, Huyen, Xa) cho location search",
        "INDEX trên CuocHen.NgayHen (cho calendar queries)",
        "INDEX trên Coc.TrangThai, Coc.NgayHetHan (cho expired deposit check)",
        "FULLTEXT INDEX trên TinDang.TieuDe, TinDang.MoTa (cho text search)"
    ]
    
    add_bullet_list(doc, indexes)

def generate_section_3_4(doc):
    """3.4. Thiết kế API"""
    doc.add_heading('3.4. Thiết kế API', level=2)
    
    # 3.4.1. API Architecture
    doc.add_heading('3.4.1. Kiến trúc API', level=3)
    
    add_paragraph(doc, """Hệ thống cung cấp RESTful API [6] với 146 endpoints được tổ chức theo domain:""")
    
    api_groups = [
        ['Authentication', '/api/auth/*', '5 endpoints', 'Login, Register, Logout, Refresh Token, Profile'],
        ['Chủ dự án', '/api/chu-du-an/*', '25 endpoints', 'CRUD Dự án, Tin đăng, Cuộc hẹn, Cọc, Hợp đồng'],
        ['Người thuê', '/api/nguoi-thue/*', '20 endpoints', 'Tìm kiếm, Đặt lịch, Cọc, Hợp đồng, Thanh toán'],
        ['Admin', '/api/admin/*', '15 endpoints', 'Duyệt tin, KYC, Báo cáo, Quản lý người dùng'],
        ['Chat & Notification', '/api/chat/*', '10 endpoints', 'Tin nhắn, Thông báo real-time'],
        ['Payment', '/api/payment/*', '8 endpoints', 'SePay integration, Webhooks'],
        ['Utilities', '/api/utils/*', '7 endpoints', 'Upload, Location, Suggest']
    ]
    
    add_table_with_data(doc, ['Domain', 'Base Path', 'Số endpoints', 'Chức năng chính'], api_groups)
    
    doc.add_paragraph()
    
    # 3.4.2. Core API Endpoints
    doc.add_heading('3.4.2. Chi tiết Core API Endpoints', level=3)
    
    add_paragraph(doc, bold=True, text="Authentication APIs")
    
    auth_apis = [
        ['POST', '/api/auth/login', 'Login', 'email, password', '200: {token, user}'],
        ['POST', '/api/auth/register', 'Register', 'user data', '201: {user}'],
        ['POST', '/api/auth/refresh', 'Refresh token', 'refreshToken', '200: {token}'],
        ['GET', '/api/auth/profile', 'Get profile', 'JWT token', '200: {user}'],
        ['POST', '/api/auth/logout', 'Logout', 'JWT token', '200: {message}']
    ]
    
    add_table_with_data(doc, ['Method', 'Endpoint', 'Description', 'Request', 'Response'], auth_apis)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Chủ dự án - Quản lý Tin đăng APIs")
    
    tindang_apis = [
        ['GET', '/api/chu-du-an/tin-dang', 'Lấy danh sách tin', 'filters, pagination', '200: {data, total}'],
        ['GET', '/api/chu-du-an/tin-dang/:id', 'Chi tiết tin', 'id', '200: {tinDang}'],
        ['POST', '/api/chu-du-an/tin-dang', 'Tạo tin mới', 'tinDang data', '201: {tinDang}'],
        ['PUT', '/api/chu-du-an/tin-dang/:id', 'Cập nhật tin', 'id, updated data', '200: {tinDang}'],
        ['DELETE', '/api/chu-du-an/tin-dang/:id', 'Xóa tin', 'id', '200: {message}'],
        ['POST', '/api/chu-du-an/tin-dang/:id/images', 'Upload ảnh', 'id, images', '200: {images}'],
        ['PUT', '/api/chu-du-an/tin-dang/:id/status', 'Đổi trạng thái', 'id, status', '200: {tinDang}']
    ]
    
    add_table_with_data(doc, ['Method', 'Endpoint', 'Description', 'Request', 'Response'], tindang_apis)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Người thuê - Tìm kiếm & Đặt lịch APIs")
    
    nguoithue_apis = [
        ['GET', '/api/nguoi-thue/tim-kiem', 'Tìm kiếm phòng', 'filters, location', '200: {results}'],
        ['GET', '/api/nguoi-thue/tin-dang/:id', 'Chi tiết tin', 'id', '200: {tinDang}'],
        ['POST', '/api/nguoi-thue/cuoc-hen', 'Đặt lịch xem', 'tinDangId, ngayHen', '201: {cuocHen}'],
        ['GET', '/api/nguoi-thue/cuoc-hen', 'DS cuộc hẹn', 'filters', '200: {cuocHen}'],
        ['PUT', '/api/nguoi-thue/cuoc-hen/:id', 'Hủy cuộc hẹn', 'id, lyDoHuy', '200: {cuocHen}'],
        ['POST', '/api/nguoi-thue/coc', 'Đặt cọc phòng', 'tinDangId, cocData', '201: {coc, paymentUrl}'],
        ['GET', '/api/nguoi-thue/yeu-thich', 'DS yêu thích', '', '200: {yeuThich}']
    ]
    
    add_table_with_data(doc, ['Method', 'Endpoint', 'Description', 'Request', 'Response'], nguoithue_apis)
    
    doc.add_paragraph()
    
    # 3.4.3. API Security
    doc.add_heading('3.4.3. API Security', level=3)
    
    add_paragraph(doc, """Các biện pháp bảo mật API tuân thủ OWASP Top 10 [8]:""")
    
    security_measures = [
        "JWT Authentication [5]: Access token (15 phút) + Refresh token (7 ngày)",
        "Role-based Authorization: Middleware kiểm tra quyền theo role",
        "Input Validation: Sử dụng express-validator cho tất cả inputs",
        "SQL Injection Prevention [8]: Sử dụng parameterized queries",
        "Rate Limiting: 100 requests/15 phút cho mỗi IP",
        "CORS Configuration: Whitelist specific origins",
        "HTTPS Only: Force HTTPS trong production",
        "Audit Logging: Ghi log tất cả critical operations"
    ]
    
    add_bullet_list(doc, security_measures)
    
    doc.add_paragraph()
    
    # 3.4.4. API Response Format
    doc.add_heading('3.4.4. Chuẩn API Response Format', level=3)
    
    add_paragraph(doc, """Tất cả API responses tuân thủ format chuẩn:""")
    
    add_paragraph(doc, bold=True, text="Success Response:")
    code_p = doc.add_paragraph("""
{
  "success": true,
  "data": { ... },
  "message": "Success message",
  "timestamp": "2025-01-15T10:30:00Z"
}
""")
    code_p.runs[0].font.name = 'Courier New'
    code_p.runs[0].font.size = Pt(10)
    
    add_paragraph(doc, bold=True, text="Error Response:")
    code_p2 = doc.add_paragraph("""
{
  "success": false,
  "error": {
    "code": "ERROR_CODE",
    "message": "Error message",
    "details": [ ... ]
  },
  "timestamp": "2025-01-15T10:30:00Z"
}
""")
    code_p2.runs[0].font.name = 'Courier New'
    code_p2.runs[0].font.size = Pt(10)

def generate_section_3_5(doc):
    """3.5. Thiết kế State Machines"""
    doc.add_heading('3.5. Thiết kế State Machines', level=2)
    
    add_paragraph(doc, """Hệ thống sử dụng State Machines [11] để quản lý các quy trình nghiệp vụ phức tạp:""")
    
    # 3.5.1. TinDang State Machine
    doc.add_heading('3.5.1. TinDang State Machine', level=3)
    
    add_paragraph(doc, """State machine quản lý vòng đời của Tin đăng:""")
    
    tindang_states = [
        ['Nhap', 'Tin đang soạn thảo', 'ChuDuAn chưa gửi duyệt'],
        ['ChoDuyet', 'Chờ Admin duyệt', 'ChuDuAn đã gửi, Admin chưa duyệt'],
        ['DaDuyet', 'Admin đã duyệt', 'Tin hợp lệ nhưng chưa đăng'],
        ['DaDang', 'Đang đăng công khai', 'Hiển thị trên trang chủ'],
        ['DaCoc', 'Đã có người cọc', 'Đang chờ ký hợp đồng'],
        ['DaThue', 'Đã cho thuê', 'Có hợp đồng hiệu lực'],
        ['DaKhoa', 'Đã khóa', 'Admin khóa do vi phạm']
    ]
    
    add_table_with_data(doc, ['State', 'Mô tả', 'Điều kiện'], tindang_states)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="State Transitions:")
    
    tindang_transitions = [
        "Nhap → ChoDuyet: ChuDuAn gửi duyệt",
        "ChoDuyet → DaDuyet: Admin duyệt",
        "ChoDuyet → Nhap: Admin từ chối (ChuDuAn sửa lại)",
        "DaDuyet → DaDang: ChuDuAn chọn đăng tin",
        "DaDang → DaCoc: NguoiThue đặt cọc thành công",
        "DaCoc → DaThue: Ký hợp đồng thành công",
        "DaCoc → DaDang: Hết hạn cọc, tin quay lại đăng",
        "Any State → DaKhoa: Admin khóa tin"
    ]
    
    add_bullet_list(doc, tindang_transitions)
    
    doc.add_paragraph()
    
    # 3.5.2. CuocHen State Machine
    doc.add_heading('3.5.2. CuocHen State Machine', level=3)
    
    cuochen_states = [
        ['ChoXacNhan', 'Chờ ChuDuAn xác nhận', 'NguoiThue vừa đặt lịch'],
        ['DaXacNhan', 'ChuDuAn đã xác nhận', 'Cuộc hẹn được confirm'],
        ['DaHuy', 'Cuộc hẹn bị hủy', 'Một trong 2 bên hủy'],
        ['DaXong', 'Cuộc hẹn hoàn thành', 'Đã xem phòng xong']
    ]
    
    add_table_with_data(doc, ['State', 'Mô tả', 'Điều kiện'], cuochen_states)
    
    doc.add_paragraph()
    
    # 3.5.3. Coc State Machine
    doc.add_heading('3.5.3. Coc State Machine', level=3)
    
    coc_states = [
        ['ChoThanhToan', 'Chờ thanh toán', 'NguoiThue vừa tạo cọc'],
        ['DaThanhToan', 'Đã thanh toán', 'Payment gateway confirm'],
        ['DaHoan', 'Đã hoàn cọc', 'ChuDuAn hoàn tiền'],
        ['DaBiMat', 'Cọc bị mất', 'NguoiThue vi phạm']
    ]
    
    add_table_with_data(doc, ['State', 'Mô tả', 'Điều kiện'], coc_states)
    
    doc.add_paragraph()
    
    # 3.5.4. HopDong State Machine
    doc.add_heading('3.5.4. HopDong State Machine', level=3)
    
    hopdong_states = [
        ['ChoKyBenA', 'Chờ NguoiThue ký', 'ChuDuAn vừa tạo HĐ'],
        ['ChoKyBenB', 'Chờ ChuDuAn ký', 'NguoiThue đã ký'],
        ['HieuLuc', 'Hợp đồng hiệu lực', 'Cả 2 bên đã ký'],
        ['HetHan', 'Hợp đồng hết hạn', 'Qua ngày kết thúc'],
        ['DaHuy', 'Hợp đồng bị hủy', 'Một bên hủy trước hiệu lực']
    ]
    
    add_table_with_data(doc, ['State', 'Mô tả', 'Điều kiện'], hopdong_states)

def generate_section_3_6(doc):
    """3.6. Thiết kế giao diện người dùng"""
    doc.add_heading('3.6. Thiết kế giao diện người dùng (UI/UX)', level=2)
    
    # 3.6.1. Design Principles
    doc.add_heading('3.6.1. Nguyên tắc thiết kế', level=3)
    
    add_paragraph(doc, """Giao diện được thiết kế sử dụng React [1] tuân theo các nguyên tắc:""")
    
    design_principles = [
        "Mobile-First: Thiết kế ưu tiên mobile, responsive trên mọi thiết bị",
        "Minimalist: Giao diện tối giản, tập trung vào nội dung",
        "Consistent: Nhất quán về màu sắc, typography, component",
        "Accessible: Tuân thủ WCAG 2.1 AA cho accessibility",
        "Fast: Optimize performance, lazy loading, code splitting"
    ]
    
    add_bullet_list(doc, design_principles)
    
    doc.add_paragraph()
    
    # 3.6.2. Component Architecture
    doc.add_heading('3.6.2. Kiến trúc Component', level=3)
    
    add_paragraph(doc, """Frontend sử dụng React [1] với component-based architecture theo Node.js best practices [12]:""")
    
    component_structure = [
        ['Layout Components', '5 components', 'Navbar, Sidebar, Footer, Container'],
        ['Page Components', '15 pages', 'Dashboard, TinDang, CuocHen, Coc, HopDong'],
        ['Feature Components', '30+ components', 'ModalTaoDuAn, TableTinDang, FormCuocHen'],
        ['Shared Components', '20+ components', 'Button, Input, Modal, Table, Card'],
        ['Utility Components', '10+ components', 'Loading, Error, Empty, Toast']
    ]
    
    add_table_with_data(doc, ['Loại Component', 'Số lượng', 'Ví dụ'], component_structure)
    
    doc.add_paragraph()
    
    # 3.6.3. Key UI Screens
    doc.add_heading('3.6.3. Màn hình chính', level=3)
    
    add_paragraph(doc, bold=True, text="1. Dashboard - Chủ dự án")
    add_paragraph(doc, """Hiển thị tổng quan quản lý:
- Thống kê: Tổng dự án, tin đăng, cuộc hẹn, doanh thu
- Biểu đồ: Hiệu suất theo thời gian
- Quick actions: Tạo dự án mới, Đăng tin, Xem cuộc hẹn
- Recent activities: 10 hoạt động gần nhất""")
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="2. Quản lý Tin đăng")
    add_paragraph(doc, """Bảng danh sách tin đăng với:
- Filters: Trạng thái, Dự án, Khoảng giá, Ngày tạo
- Search: Full-text search
- Actions: Sửa, Xóa, Đổi trạng thái, Upload ảnh
- Pagination: Load more khi scroll""")
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="3. Chi tiết Tin đăng - Người thuê")
    add_paragraph(doc, """Hiển thị đầy đủ thông tin:
- Gallery ảnh với lightbox
- Thông tin phòng: Giá, diện tích, địa chỉ
- Thuộc tính & tiện ích
- Map tích hợp Google Maps
- Contact ChuDuAn: Chat, Đặt lịch xem
- Related listings""")
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="4. Quản lý Cuộc hẹn")
    add_paragraph(doc, """Calendar view với:
- Month/Week/Day view
- Drag & drop để reschedule
- Color coding theo trạng thái
- Quick actions: Xác nhận, Hủy, Hoàn thành""")
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="5. Chat & Notifications")
    add_paragraph(doc, """Real-time chat:
- List conversations
- Message thread với ảnh
- Typing indicator
- Read receipts
- Push notifications""")
    
    doc.add_paragraph()
    
    # 3.6.4. Design System
    doc.add_heading('3.6.4. Design System', level=3)
    
    add_paragraph(doc, bold=True, text="Color Palette:")
    
    colors = [
        ['Primary', '#8b5cf6', 'Purple', 'Main brand color'],
        ['Secondary', '#64748b', 'Gray', 'Secondary actions'],
        ['Success', '#10b981', 'Green', 'Success states'],
        ['Warning', '#f59e0b', 'Orange', 'Warning states'],
        ['Danger', '#ef4444', 'Red', 'Error states'],
        ['Info', '#3b82f6', 'Blue', 'Info messages']
    ]
    
    add_table_with_data(doc, ['Name', 'Hex', 'Color', 'Usage'], colors)
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Typography:")
    add_paragraph(doc, """
- Primary Font: Inter (sans-serif)
- Heading 1: 32px, Bold
- Heading 2: 24px, Bold
- Heading 3: 20px, SemiBold
- Body: 16px, Regular
- Small: 14px, Regular
""")
    
    doc.add_paragraph()
    
    add_paragraph(doc, bold=True, text="Spacing System:")
    add_paragraph(doc, "Sử dụng 8px base unit: 8px, 16px, 24px, 32px, 40px, 48px, 64px")

def generate_chuong_3(doc):
    """Generate full Chuong 3"""
    # Chapter title
    doc.add_heading('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', level=1)
    doc.add_paragraph()
    
    # Generate all sections
    generate_section_3_1(doc)
    doc.add_page_break()
    
    generate_section_3_2(doc)
    doc.add_page_break()
    
    generate_section_3_3(doc)
    doc.add_page_break()
    
    generate_section_3_4(doc)
    doc.add_page_break()
    
    generate_section_3_5(doc)
    doc.add_page_break()
    
    generate_section_3_6(doc)

def main():
    """Main function"""
    print("[*] Generating Chuong 3: PHAN TICH VA THIET KE HE THONG...")
    
    # Create document
    doc = Document()
    
    # Setup styles
    setup_document_styles(doc)
    
    # Add title page
    title = doc.add_heading('BÁO CÁO KHÓA LUẬN TỐT NGHIỆP', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph('HỆ THỐNG QUẢN LÝ CHO THUÊ PHÒNG TRỌ')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    
    subtitle2 = doc.add_paragraph('CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].italic = True
    subtitle2.runs[0].font.size = Pt(13)
    
    doc.add_page_break()
    
    # Generate chapter
    generate_chuong_3(doc)
    
    # Save document
    output_dir = os.path.join(os.path.dirname(__file__), '..')
    output_file = os.path.join(output_dir, 'BaoCao_Chuong3_FULL.docx')
    
    doc.save(output_file)
    
    print(f"[OK] Done! File saved to: {output_file}")
    print(f"[INFO] Total pages: ~20-25 pages (Chuong 3)")

if __name__ == "__main__":
    main()

